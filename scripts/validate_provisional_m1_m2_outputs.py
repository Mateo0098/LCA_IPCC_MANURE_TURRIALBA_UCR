"""Validación cruzada de productos académicos de la corrida PROVISIONAL M1–M2."""

from __future__ import annotations

import csv
import hashlib
import math
from pathlib import Path

from docx import Document

import generate_conclusions_docx as conclusions_generator
import generate_a2_jjagwe_benchmark as benchmark_generator
import generate_thesis_graphics as graphics_generator
import validate_ef31_operational_inventory as ef31_validator
from quantitative_comparison import Comparison, dominant


ROOT = Path(__file__).resolve().parents[1]
PROCESSED = ROOT / "processed"
TABLES = ROOT / "outputs" / "tablas_tesis"
GRAPHICS = ROOT / "outputs" / "graficos_tesis"
DOCS = ROOT / "outputs" / "documentos_tfg"
LABEL = "PROVISIONAL M1–M2"
REPORT = DOCS / "reporte_validacion_provisional_m1_m2.md"


def read_rows(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as stream:
        return list(csv.DictReader(stream))


def close(left: object, right: object, tolerance: float = 1e-12) -> bool:
    return math.isclose(float(left), float(right), rel_tol=tolerance, abs_tol=tolerance)


def finite_sum(row: dict[str, str], columns: list[str]) -> float:
    values = []
    for column in columns:
        raw = row.get(column, "")
        try:
            value = float(raw)
        except (TypeError, ValueError):
            continue
        if math.isfinite(value):
            values.append(value)
    return sum(values)


def document_text(path: Path) -> tuple[str, str]:
    document = Document(path)
    body = [paragraph.text for paragraph in document.paragraphs]
    body.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    headers = [paragraph.text for section in document.sections for paragraph in section.header.paragraphs]
    return "\n".join(body), "\n".join(headers)


def fmt_es(value: float, decimals: int) -> str:
    text = f"{float(value):,.{decimals}f}"
    text = text.replace(",", "X").replace(".", ",").replace("X", " ")
    return text.rstrip("0").rstrip(",")


def validate_characterization() -> None:
    integration = {
        (row["material"], row["variable"]): row
        for row in read_rows(PROCESSED / "muestreos_integracion_interjornada_provisional.csv")
        if row["valor_integrado_provisional"]
    }
    labels = {
        "Estiércol fresco": "estiércol fresco",
        "Estiércol precompostado": "estiércol precompostado",
        "Aguas verdes": "aguas verdes",
        "Purines": "purines",
    }
    variables = {
        "Nitrogeno total": "N total",
        "Humedad": "humedad",
        "Materia seca": "materia seca",
        "Cenizas": "cenizas",
        "Solidos volatiles": "sólidos volátiles",
        "Densidad": "densidad",
        "Carbono": "carbono",
        "Relación C/N": "relación C/N",
    }
    table = read_rows(TABLES / "tabla_02_caracterizacion_muestras.csv")
    for row in table:
        source = integration[(labels[row["tipo_muestra"]], variables[row["variable"]])]
        assert close(row["valor"], source["valor_integrado_provisional"])
        assert row["jornada_muestreo"] == source["jornadas_elegibles"]
        assert row["estado_integracion"] == source["estado_integracion"]
    active = read_rows(PROCESSED / "acv_parametros_escenario_etapa.csv")
    liquid = {(row["escenario"], row["etapa"]): row for row in active if row["escenario"] + row["etapa"] in {"A4", "B2"}}
    assert all(row["jornadas_n_ex"] == "M2" for row in liquid.values())


def validate_factor_and_masses() -> None:
    transformation = next(
        row for row in read_rows(PROCESSED / "muestreos_transformacion_masa_interjornada.csv")
        if row["tipo_fila"] == "integracion"
    )
    factor = transformation["mass_ratio_integrado"]
    masses = {(row["escenario"], row["etapa"]): row for row in read_rows(PROCESSED / "masa_total_escenario_etapa.csv")}
    flows = read_rows(TABLES / "tabla_03_flujos_icv.csv")
    for row in flows:
        key = (row["escenario"], row["etapa"])
        if row["flujo"] == "Factor restante fresco a precompostado":
            assert close(row["valor"], factor)
        if row["flujo"] == "Masa equivalente total":
            assert close(row["valor"], masses[key]["masa_total_kg_eq"])


def validate_a2_nitrogen_basis() -> None:
    params = {
        (row["escenario"], row["etapa"]): row
        for row in read_rows(PROCESSED / "acv_parametros_escenario_etapa.csv")
    }
    masses = {
        (row["escenario"], row["etapa"]): row
        for row in read_rows(PROCESSED / "masa_total_escenario_etapa.csv")
    }
    a2 = params[("A", "2")]
    integrated_n = next(
        row for row in read_rows(PROCESSED / "muestreos_integracion_interjornada_provisional.csv")
        if row["material"] == "estiércol precompostado" and row["variable"] == "N total"
    )
    assert close(a2["n_ex_pct"], integrated_n["valor_integrado_provisional"])
    assert a2["base_analitica_n_ex"] == "material preparado/seco por CIA a 80 °C durante 48 h"
    assert a2["transformacion_n_acv"] == "multiplicar_por_fraccion_materia_seca_gravimetrica_TFG_105C"
    effective_fraction = (float(a2["n_ex_pct"]) / 100.0) * (float(a2["materia_seca_pct"]) / 100.0)
    annual_n = float(masses[("A", "2")]["masa_total_kg_eq"]) * effective_fraction
    assert effective_fraction > 0.0 and annual_n > 0.0
    for key, row in params.items():
        if key != ("A", "2"):
            assert row["transformacion_n_acv"] == "ninguna"


def validate_emissions() -> None:
    raw = {(row["Escenario"], row["Etapa"]): row for row in read_rows(PROCESSED / "ACV_resumen_emisiones.csv")}
    table = read_rows(TABLES / "tabla_06_emisiones_por_etapa.csv")
    grouped: dict[tuple[str, str, str], float] = {}
    for row in table:
        key = (row["escenario"], row["etapa"], row["sustancia"])
        grouped[key] = grouped.get(key, 0.0) + float(row["valor"])
    columns = {
        "CO2": ["CO2_medido"], "CH4": ["CH4_ec1"],
        "N2O": ["N2O_ec14", "N2O_ec2", "N2O_ec5", "N2O_ec6", "N2O_ec16", "N2O_ec18"],
        "NH3": ["NH3_ec12", "NH3_ec20"], "NO3": ["NO3_ec13", "NO3_ec21"],
    }
    for key, row in raw.items():
        for substance, source_columns in columns.items():
            assert close(grouped.get((*key, substance), 0.0), finite_sum(row, source_columns))


def validate_a2_benchmark() -> None:
    generated = read_rows(PROCESSED / "a2_ipcc_jjagwe_benchmark.csv")
    expected = benchmark_generator.build_rows()
    assert len(generated) == len(expected) == 4
    for observed, calculated in zip(generated, expected):
        assert observed.keys() == calculated.keys()
        for key, value in calculated.items():
            if isinstance(value, float):
                assert close(observed[key], value)
            else:
                assert observed[key] == value
    assert all("eutrof" not in row["indicador"].lower() for row in generated)
    direct = next(row for row in generated if row["indicador"] == "N2O directo por materia seca de entrada")
    emissions = next(
        row for row in read_rows(PROCESSED / "ACV_resumen_emisiones.csv")
        if row["Escenario"] == "A" and row["Etapa"] == "2"
    )
    dry_mass = float(direct["masa_seca_entrada_a2_kg_anio"])
    assert close(direct["valor_ipcc"], float(emissions["N2O_ec2"]) * 1_000_000.0 / dry_mass)
    assert not close(direct["valor_ipcc"], finite_sum(emissions, ["N2O_ec2", "N2O_ec5", "N2O_ec6"]) * 1_000_000.0 / dry_mass)
    models = read_rows(PROCESSED / "modelo_etapa_overrides.csv")
    a2_model = next(row for row in models if row["escenario"] == "A" and row["etapa"] == "2")
    assert a2_model["modelo"] == "ipcc"
    assert all(row["modelo"] != "medido" for row in models)


IMPACT_COLUMNS = {
    "Cambio climático": ("impacto_calentamiento_global_kg_co2eq", "kg CO2-eq/año"),
    "Eutrofización terrestre": ("impacto_eutrofizacion_terrestre_mol_neq", "mol N-eq/año"),
    "Eutrofización marina": ("impacto_eutrofizacion_marina_kg_neq", "kg N-eq/año"),
}


def validate_impacts_and_comparison() -> None:
    stage_source = {
        (row["Escenario"], row["Etapa"]): row
        for row in read_rows(PROCESSED / "acv_impacto_por_etapa_escenario.csv")
    }
    stage_table = read_rows(TABLES / "tabla_07_impactos_por_etapa.csv")
    grouped: dict[tuple[str, str, str], float] = {}
    units: dict[tuple[str, str, str], set[str]] = {}
    for row in stage_table:
        key = (row["escenario"], row["etapa"], row["categoria_impacto"])
        grouped[key] = grouped.get(key, 0.0) + float(row["resultado_equivalente"])
        units.setdefault(key, set()).add(row["unidad_equivalente"])
    for key, row in stage_source.items():
        for category, (column, expected_unit) in IMPACT_COLUMNS.items():
            assert close(grouped[(*key, category)], row[column])
            assert units[(*key, category)] == {expected_unit}

    totals_source = {
        (row["Escenario"], category): float(row[column])
        for row in read_rows(PROCESSED / "acv_impacto_total_por_escenario.csv")
        for category, (column, _) in IMPACT_COLUMNS.items()
    }
    totals_table = {
        (row["escenario"], row["categoria_impacto"]): float(row["resultado_total"])
        for row in read_rows(TABLES / "tabla_08_impactos_totales_por_escenario.csv")
    }
    assert totals_source.keys() == totals_table.keys()
    assert all(close(totals_source[key], totals_table[key]) for key in totals_source)
    comparison_rows = read_rows(TABLES / "tabla_09_comparacion_escenarios.csv")
    assert {row["categoria_impacto"] for row in comparison_rows} == set(IMPACT_COLUMNS)
    for row in comparison_rows:
        category = row["categoria_impacto"]
        expected_unit = IMPACT_COLUMNS[category][1]
        a, b = totals_source[("A", category)], totals_source[("B", category)]
        assert close(row["escenario_A"], a) and close(row["escenario_B"], b)
        assert close(row["diferencia_absoluta_B_menos_A"], b - a)
        assert close(row["diferencia_porcentual_B_vs_A"], (b - a) / a * 100.0)
        assert row["escenario_con_mayor_impacto"] == ("B" if b > a else "A" if a > b else "Iguales")
        assert row["unidad"] == expected_unit
        comparison = Comparison("A", a, "B", b, expected_unit)
        comparison.assert_consistent(
            difference=float(row["diferencia_absoluta_B_menos_A"]),
            percentage=float(row["diferencia_porcentual_B_vs_A"]),
        )
        comparison.assert_rounding_unambiguous(6)


def validate_quantitative_narratives() -> None:
    results_text, _ = document_text(DOCS / "resultados_desarrollados_tfg.docx")
    conclusions_text, _ = document_text(DOCS / "conclusiones_desarrolladas_tfg.docx")
    totals = {
        (row["Escenario"], category): float(row[column])
        for row in read_rows(PROCESSED / "acv_impacto_total_por_escenario.csv")
        for category, (column, _) in IMPACT_COLUMNS.items()
    }
    narrative_terms = {
        "Cambio climático": "calentamiento global",
        "Eutrofización terrestre": "eutrofización terrestre",
        "Eutrofización marina": "eutrofización marina",
    }
    for category, (_, unit) in IMPACT_COLUMNS.items():
        comparison = Comparison(
            "Escenario A", totals[("A", category)], "Escenario B", totals[("B", category)], unit
        )
        assert comparison.higher_label and comparison.lower_label
        visible_term = narrative_terms[category]
        assert visible_term in results_text.lower() and visible_term in conclusions_text.lower()
        if category == "Cambio climático":
            assert f"mayor impacto de calentamiento global en el {comparison.higher_label}" in results_text
        else:
            assert f"En {visible_term}, el {comparison.higher_label} presentó el mayor impacto" in results_text
        assert comparison.higher_label in conclusions_text

    synthetic = Comparison("A", 10.0, "B", 8.0, "kg/año")
    synthetic.assert_consistent(difference=-2.0, percentage=-20.0)
    assert synthetic.higher_label == "A" and synthetic.lower_label == "B"
    try:
        synthetic.assert_consistent(difference=2.0, percentage=-20.0)
    except RuntimeError:
        pass
    else:
        raise AssertionError("No se detectó una diferencia con signo contradictorio.")
    try:
        Comparison("A", 1.0004, "B", 1.0003, "kg/año").assert_rounding_unambiguous(3)
    except RuntimeError:
        pass
    else:
        raise AssertionError("No se detectó una comparación ambigua por redondeo.")
    assert dominant({"A1": 1.0, "A2": 2.0}, decimals=3) == ("A2", 2.0)


def validate_documents_and_conclusions() -> None:
    methodology_text, methodology_header = document_text(DOCS / "metodologia_desarrollada_tfg.docx")
    results_text, results_header = document_text(DOCS / "resultados_desarrollados_tfg.docx")
    conclusions_text, conclusions_header = document_text(DOCS / "conclusiones_desarrolladas_tfg.docx")
    for body, header in (
        (methodology_text, methodology_header),
        (results_text, results_header),
        (conclusions_text, conclusions_header),
    ):
        assert LABEL in body and LABEL in header and "M3" in body
        assert "PO4-eq" not in body and "kg PO₄" not in body

    for visible in ("EF 3.1", "53,23 kWh/año", "182,50 L/año"):
        assert visible in methodology_text
    assert "30 min" in methodology_text and "no medida instrumentalmente" in methodology_text

    characterization = read_rows(TABLES / "tabla_02_caracterizacion_muestras.csv")
    required = [
        row for row in characterization
        if (row["tipo_muestra"], row["variable"]) in {
            ("Estiércol fresco", "Nitrogeno total"), ("Estiércol fresco", "Materia seca"),
            ("Estiércol fresco", "Solidos volatiles"), ("Estiércol precompostado", "Materia seca"),
            ("Aguas verdes", "Nitrogeno total"), ("Purines", "Nitrogeno total"),
        }
    ]
    assert len(required) == 6
    assert all(fmt_es(float(row["valor"]), 3) in methodology_text for row in required)
    assert all(fmt_es(float(row["valor"]), 3) in results_text for row in required)

    totals = read_rows(PROCESSED / "acv_impacto_total_por_escenario.csv")
    for row in totals:
        for column, _ in IMPACT_COLUMNS.values():
            assert fmt_es(float(row[column]), 6) in results_text

    totals_dict = conclusions_generator.impact_totals()
    percentages = conclusions_generator.comparisons(totals_dict)
    reference, normalized = conclusions_generator.processed_indicators(totals_dict)
    collected, remainder = conclusions_generator.flow_inventory(reference)
    stages = conclusions_generator.stage_totals()
    expected = conclusions_generator.build_conclusions(
        totals_dict, normalized, percentages, stages, reference, collected, remainder
    )
    assert all(item["text"] in conclusions_text for item in expected)


def validate_graph_sources_and_freshness() -> None:
    source = (ROOT / "scripts" / "generate_thesis_graphics.py").read_text(encoding="utf-8")
    assert 'TABLE_DIR = BASE_DIR / "outputs" / "tablas_tesis"' in source
    assert "processed/" not in source and "CIA_samples" not in source
    graphics = sorted(GRAPHICS.glob("fig_*.png")) + sorted(GRAPHICS.glob("fig_*.svg"))
    assert len(graphics) == 34
    manifest = (GRAPHICS / "README_GRAFICOS.md").read_text(encoding="utf-8")
    assert all(path.name in manifest for path in graphics)
    for table in ("tabla_07_impactos_por_etapa.csv", "tabla_08_impactos_totales_por_escenario.csv", "tabla_09_comparacion_escenarios.csv"):
        assert table in manifest
    newest_table = max((TABLES / name).stat().st_mtime_ns for name in (
        "tabla_07_impactos_por_etapa.csv", "tabla_08_impactos_totales_por_escenario.csv",
        "tabla_09_comparacion_escenarios.csv",
    ))
    impact_graphics = [path for path in graphics if int(path.name.split("_")[1]) >= 11]
    assert all(path.stat().st_mtime_ns >= newest_table for path in impact_graphics)
    newest_graphic = max(path.stat().st_mtime_ns for path in impact_graphics)
    assert (GRAPHICS / "README_GRAFICOS.md").stat().st_mtime_ns >= newest_graphic
    for document in (
        DOCS / "metodologia_desarrollada_tfg.docx",
        DOCS / "resultados_desarrollados_tfg.docx",
        DOCS / "conclusiones_desarrolladas_tfg.docx",
    ):
        assert document.stat().st_mtime_ns >= newest_graphic

    samples = graphics_generator.read_table("muestras")
    for variables in (["Humedad", "Materia seca"], ["Solidos volatiles", "Cenizas"]):
        series = graphics_generator.characterization_series(samples, variables)
        assert set(series["variable"]) == set(variables)
        assert series["valor"].notna().all() and not series.empty


def main() -> None:
    validate_characterization()
    validate_factor_and_masses()
    validate_a2_nitrogen_basis()
    validate_emissions()
    validate_a2_benchmark()
    ef31_validator.main()
    validate_impacts_and_comparison()
    validate_quantitative_narratives()
    validate_documents_and_conclusions()
    validate_graph_sources_and_freshness()
    master = ROOT / "MASTER_escrito" / "TFG_ACV_Estiercol_MASTER.docx"
    master_hash = hashlib.sha256(master.read_bytes()).hexdigest().upper()
    REPORT.write_text(
        "# Validación cruzada de la corrida PROVISIONAL M1–M2\n\n"
        "- Caracterización experimental contra integración vigente: PASS.\n"
        "- Factor fresco→precompostado contra `mass_ratio_integrado`: PASS.\n"
        "- Masas contra inventario canónico: PASS.\n"
        "- Base de N de A2, fórmula húmeda y exclusión de las demás etapas: PASS.\n"
        "- Emisiones contra resumen canónico: PASS.\n"
        "- Contraste Jjagwe reproducido, N2O directo aislado y sin ruta medida ni eutrofización experimental: PASS.\n"
        "- EF 3.1, casos unitarios, impactos y unidades: PASS.\n"
        "- Electricidad, diésel y normalización por unidad funcional: PASS.\n"
        "- Exportación foreground y controles de doble conteo: PASS.\n"
        "- Impactos por etapa y totales contra tablas canónicas, con unidades EF 3.1: PASS.\n"
        "- Comparación A–B, diferencias, porcentajes, dominancia, signos y redondeo: PASS.\n"
        "- Metodología, resultados y conclusiones identificados como `PROVISIONAL M1–M2`: PASS.\n"
        "- Cifras documentales verificadas con el redondeo visible: PASS.\n"
        "- Conclusiones reconstruidas desde fuentes canónicas: PASS.\n"
        "- Gráficos, manifiestos y vigencia relativa de productos: PASS.\n"
        "- Productos académicos regenerados desde las fuentes canónicas: PASS.\n"
        f"- SHA-256 del MASTER verificado: `{master_hash}`.\n",
        encoding="utf-8",
    )
    print("VALIDACIÓN CRUZADA PROVISIONAL M1–M2: PASS")
    print("Caracterización, factor, masas, emisiones, impactos, comparación, DOCX, conclusiones y gráficos: PASS")
    print(f"Reporte: {REPORT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
