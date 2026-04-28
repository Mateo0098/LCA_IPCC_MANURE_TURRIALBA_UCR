from __future__ import annotations

from pathlib import Path

import latex2mathml.converter
import mathml2omml
from docx import Document
from docx.oxml import parse_xml

MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def add_equation_from_latex(doc: Document, latex_expr: str) -> None:
    mathml = latex2mathml.converter.convert(latex_expr)
    omml = mathml2omml.convert(mathml)
    p = doc.add_paragraph()
    p._p.append(parse_xml(f'<m:oMathPara xmlns:m="{MATH_NS}">{omml}</m:oMathPara>'))


def add_variable(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def main() -> None:
    project_root = Path(__file__).resolve().parent.parent
    output_path = project_root / "Ecuaciones_boniga_fresco_precompostado.docx"

    doc = Document()
    doc.add_heading(
        "Ecuaciones para humedad, materia seca, cenizas, solidos volatiles y perdida de peso",
        level=1,
    )

    doc.add_heading("1) Contenido de humedad", level=2)
    add_equation_from_latex(
        doc,
        r"H = \frac{m_{\mathrm{humeda}} - m_{\mathrm{seca}}}{m_{\mathrm{humeda}}}\times 100\;[\%]",
    )
    add_variable(doc, r"m_humeda [g]: masa de muestra humeda.")
    add_variable(doc, r"m_seca [g]: masa de muestra seca despues del secado.")
    add_variable(doc, r"H [%]: contenido de humedad.")

    doc.add_heading("2) Materia seca", level=2)
    add_equation_from_latex(
        doc, r"MS = \frac{m_{\mathrm{seca}}}{m_{\mathrm{humeda}}}\times 100\;[\%]"
    )
    add_variable(doc, r"MS [%]: contenido de materia seca.")

    doc.add_heading("3) Contenido de cenizas (base seca)", level=2)
    add_equation_from_latex(
        doc,
        r"C = \frac{m_{\mathrm{ceniza}}}{m_{\mathrm{seca,inc}}}\times 100\;[\%\ \mathrm{base\ seca}]",
    )
    add_variable(doc, r"m_seca_inc [g]: masa seca usada en prueba de cenizas.")
    add_variable(doc, r"m_ceniza [g]: masa de cenizas tras incineracion.")
    add_variable(doc, r"C [% base seca]: contenido de cenizas.")

    doc.add_heading("4) Solidos volatiles (base seca)", level=2)
    add_equation_from_latex(doc, r"SV = 100 - C\;[\%\ \mathrm{base\ seca}]")
    add_variable(doc, r"SV [% base seca]: solidos volatiles.")

    doc.add_heading("5) Perdida de peso entre dos estados (general)", level=2)
    add_equation_from_latex(
        doc, r"\%P = \left(1-\frac{M_2}{M_1}\right)\times 100\;[\%]"
    )
    add_variable(doc, r"M_1 [kg o g]: masa humeda en estado inicial (fresco).")
    add_variable(doc, r"M_2 [kg o g]: masa humeda en estado final (precompostado).")
    add_variable(doc, r"%P [%]: porcentaje de perdida de peso.")

    doc.add_heading(
        "6) Perdida de peso incluyendo perdida de materia seca (trazador cenizas)",
        level=2,
    )
    add_equation_from_latex(doc, r"\frac{MS_2}{MS_1} = \frac{C_1}{C_2}\;[-]")
    add_equation_from_latex(
        doc,
        r"\frac{M_2}{M_1}=\left(\frac{MS_2}{MS_1}\right)\left(\frac{MS\%_1}{MS\%_2}\right)=\left(\frac{C_1}{C_2}\right)\left(\frac{MS\%_1}{MS\%_2}\right)\;[-]",
    )
    add_equation_from_latex(
        doc, r"\%P = \left(1-\frac{M_2}{M_1}\right)\times 100\;[\%]"
    )
    add_variable(doc, r"C_1, C_2 [% base seca]: cenizas en estado 1 (fresco) y 2 (precompostado).")
    add_variable(
        doc,
        r"MS%_1, MS%_2 [% base humeda]: porcentaje de materia seca en estado 1 y 2.",
    )
    add_variable(doc, r"MS_1, MS_2 [kg o g]: masa seca total en estado 1 y 2.")
    add_variable(doc, r"M_1, M_2 [kg o g]: masa humeda total en estado 1 y 2.")

    doc.add_paragraph(
        "Nota metodologica: bajo el supuesto A->B como mismo lote, se usa la ceniza "
        "como trazador para estimar retencion de materia seca y luego retencion de masa humeda."
    )

    doc.save(output_path)
    print(f"Generated file: {output_path}")


if __name__ == "__main__":
    main()
