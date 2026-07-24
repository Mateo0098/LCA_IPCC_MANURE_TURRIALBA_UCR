from __future__ import annotations

from dataclasses import asdict, dataclass
from functools import lru_cache
from pathlib import Path
from statistics import median
from zipfile import ZipFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


@dataclass(frozen=True)
class MasterFormatProfile:
    font_name: str
    body_size_pt: float
    title_size_pt: float
    heading_size_pt: float
    caption_size_pt: float
    table_size_pt: float
    body_alignment: str
    body_line_spacing: float
    body_left_indent_pt: float
    body_first_line_indent_pt: float
    top_margin_pt: float
    bottom_margin_pt: float
    left_margin_pt: float
    right_margin_pt: float


def _pt(value, default: float = 0.0) -> float:
    return round(value.pt, 2) if value is not None else default


def _median(values: list[float], default: float) -> float:
    return round(float(median(values)), 2) if values else default


def _default_font_from_styles(reference_docx: Path) -> str:
    with ZipFile(reference_docx) as archive:
        xml = archive.read("word/styles.xml").decode("utf-8", errors="ignore")
    marker = '<w:rFonts w:ascii="'
    if marker in xml:
        return xml.split(marker, 1)[1].split('"', 1)[0]
    return "Times New Roman"


@lru_cache(maxsize=4)
def analyze_master_format(reference_docx: Path) -> MasterFormatProfile:
    document = Document(str(reference_docx))
    font_name = _default_font_from_styles(reference_docx)

    body_paragraphs = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip()
        and paragraph.style.name == "Normal"
        and paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
        and not paragraph.text.strip().lower().startswith(("figura ", "tabla "))
    ]
    body_sizes = [
        _pt(run.font.size)
        for paragraph in body_paragraphs
        for run in paragraph.runs
        if run.text.strip() and run.font.size is not None
    ]
    first_line_indents = [
        _pt(paragraph.paragraph_format.first_line_indent)
        for paragraph in body_paragraphs
        if paragraph.paragraph_format.first_line_indent is not None
        and _pt(paragraph.paragraph_format.first_line_indent) > 0
    ]
    left_indents = [
        _pt(paragraph.paragraph_format.left_indent)
        for paragraph in body_paragraphs
        if paragraph.paragraph_format.left_indent is not None
        and _pt(paragraph.paragraph_format.left_indent) >= 0
    ]

    title_candidates = [
        run
        for paragraph in document.paragraphs[:30]
        if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs
        if run.text.strip() and run.font.italic and run.font.size is not None
    ]
    title_size = max((_pt(run.font.size) for run in title_candidates), default=16.0)

    heading_runs = [
        run
        for paragraph in document.paragraphs
        if paragraph.style.name in {"Heading 1", "Heading 3"}
        for run in paragraph.runs
        if run.text.strip()
    ]
    heading_sizes = [_pt(run.font.size) for run in heading_runs if run.font.size]

    caption_runs = [
        run
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip().lower().startswith(("figura ", "tabla "))
        for candidate in document.paragraphs[index : index + 2]
        for run in candidate.runs
        if run.text.strip() and run.font.size is not None
    ]
    caption_sizes = [_pt(run.font.size) for run in caption_runs]

    table_sizes = [
        _pt(run.font.size)
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        for run in paragraph.runs
        if run.text.strip() and run.font.size is not None
    ]
    section = document.sections[0]
    return MasterFormatProfile(
        font_name=font_name,
        body_size_pt=_median(body_sizes, 12.0),
        title_size_pt=title_size,
        heading_size_pt=_median(heading_sizes, 12.0),
        caption_size_pt=_median(caption_sizes, 12.0),
        table_size_pt=_median(table_sizes, 12.0),
        body_alignment="justificada",
        body_line_spacing=1.0,
        # The MASTER also contains bibliography/list paragraphs with larger
        # direct indents; use the smallest positive body indent as the prose
        # reference and deliberately exclude numbering-related indentation.
        body_left_indent_pt=min(left_indents, default=5.0),
        body_first_line_indent_pt=_median(first_line_indents, 35.95),
        top_margin_pt=_pt(section.top_margin),
        bottom_margin_pt=_pt(section.bottom_margin),
        left_margin_pt=_pt(section.left_margin),
        right_margin_pt=_pt(section.right_margin),
    )


def _set_style_font(style, profile: MasterFormatProfile, size_pt: float) -> None:
    style.font.name = profile.font_name
    style.font.size = Pt(size_pt)
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), profile.font_name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), profile.font_name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), profile.font_name)


def _get_or_add_paragraph_style(document: Document, name: str):
    try:
        return document.styles[name]
    except KeyError:
        return document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def apply_master_format(document: Document, reference_docx: Path) -> MasterFormatProfile:
    profile = analyze_master_format(reference_docx)
    section = document.sections[0]
    section.top_margin = Pt(profile.top_margin_pt)
    section.bottom_margin = Pt(profile.bottom_margin_pt)
    section.left_margin = Pt(profile.left_margin_pt)
    section.right_margin = Pt(profile.right_margin_pt)

    normal = document.styles["Normal"]
    _set_style_font(normal, profile, profile.body_size_pt)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = profile.body_line_spacing
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.left_indent = Pt(profile.body_left_indent_pt)
    normal.paragraph_format.first_line_indent = Pt(profile.body_first_line_indent_pt)

    title = document.styles["Title"]
    _set_style_font(title, profile, profile.title_size_pt)
    title.font.bold = False
    title.font.italic = True
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.left_indent = Pt(0)
    title.paragraph_format.first_line_indent = Pt(0)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        _set_style_font(style, profile, profile.heading_size_pt)
        style.font.bold = True
        style.font.italic = False
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.left_indent = Pt(0)
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.keep_with_next = True

    label_style = _get_or_add_paragraph_style(document, "Rótulo académico")
    _set_style_font(label_style, profile, profile.caption_size_pt)
    label_style.font.bold = True
    label_style.font.italic = False
    label_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    label_style.paragraph_format.space_before = Pt(6)
    label_style.paragraph_format.space_after = Pt(0)
    label_style.paragraph_format.left_indent = Pt(0)
    label_style.paragraph_format.first_line_indent = Pt(0)

    caption_style = _get_or_add_paragraph_style(document, "Descripción académica")
    _set_style_font(caption_style, profile, profile.caption_size_pt)
    caption_style.font.bold = False
    caption_style.font.italic = True
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    caption_style.paragraph_format.space_before = Pt(0)
    caption_style.paragraph_format.space_after = Pt(6)
    caption_style.paragraph_format.left_indent = Pt(0)
    caption_style.paragraph_format.first_line_indent = Pt(0)
    return profile


def add_master_caption(document: Document, caption: str) -> None:
    label, separator, description = caption.partition(". ")
    label_paragraph = document.add_paragraph(style="Rótulo académico")
    label_paragraph.add_run(label.rstrip(".")).bold = True
    if separator and description:
        description_paragraph = document.add_paragraph(style="Descripción académica")
        description_paragraph.add_run(description).italic = True


def format_table_like_master(table, profile: MasterFormatProfile) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for edge in ("top", "left", "bottom", "right"):
                node = tc_mar.find(qn(f"w:{edge}"))
                if node is None:
                    node = OxmlElement(f"w:{edge}")
                    tc_mar.append(node)
                node.set(qn("w:w"), "72")
                node.set(qn("w:type"), "dxa")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.left_indent = Pt(0)
                paragraph.paragraph_format.first_line_indent = Pt(0)
                for run in paragraph.runs:
                    run.font.name = profile.font_name
                    run.font.size = Pt(profile.table_size_pt)
                    run.bold = True if row_index == 0 else False


def finalize_document_format(
    document: Document, profile: MasterFormatProfile
) -> None:
    for paragraph in document.paragraphs:
        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            continue
        if any(run.font.name == "Cambria Math" for run in paragraph.runs):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            continue
        if paragraph.style.name == "Normal":
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.line_spacing = profile.body_line_spacing
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.left_indent = Pt(profile.body_left_indent_pt)
            paragraph.paragraph_format.first_line_indent = Pt(
                profile.body_first_line_indent_pt
            )
        for run in paragraph.runs:
            if run.font.name != "Cambria Math":
                run.font.name = profile.font_name
            if paragraph.style.name == "Normal" and run.font.size is None:
                run.font.size = Pt(profile.body_size_pt)


def profile_markdown(profile: MasterFormatProfile) -> str:
    values = asdict(profile)
    return "\n".join(
        [
            f"- Fuente predominante: {values['font_name']}.",
            f"- Texto normal: {values['body_size_pt']:.0f} pt, alineación {values['body_alignment']}, interlineado {values['body_line_spacing']:.2f}.",
            f"- Sangría izquierda del texto normal: {values['body_left_indent_pt']:.2f} pt.",
            f"- Sangría de primera línea: {values['body_first_line_indent_pt']:.2f} pt.",
            f"- Título principal detectado: {values['title_size_pt']:.0f} pt, centrado y en cursiva.",
            f"- Títulos y subtítulos detectados: {values['heading_size_pt']:.0f} pt y negrita.",
            f"- Rótulos y descripciones de tablas/figuras: {values['caption_size_pt']:.0f} pt; rótulo en negrita y descripción en cursiva.",
            f"- Texto de tablas: {values['table_size_pt']:.0f} pt y alineación centrada.",
            f"- Márgenes (superior, inferior, izquierdo, derecho): {values['top_margin_pt']:.0f}, {values['bottom_margin_pt']:.0f}, {values['left_margin_pt']:.0f}, {values['right_margin_pt']:.0f} pt.",
        ]
    )
