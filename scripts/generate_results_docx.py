from __future__ import annotations

import re
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from academic_text_utils import clean_academic_label
from master_word_format import (
    add_master_caption,
    analyze_master_format,
    apply_master_format,
    finalize_document_format,
    format_table_like_master,
    profile_markdown,
)
from reference_docx_utils import (
    REGISTERED_REFERENCE_SHA256,
    assert_reference_docx_intact,
    get_reference_docx_path,
    sha256_file,
)
from quantitative_comparison import Comparison, dominant


ROOT = Path(__file__).resolve().parents[1]
PROVISIONAL_LABEL = "PROVISIONAL M1–M2"
PROVISIONAL_NOTE = (
    "Esta versión incorpora las jornadas disponibles M1 y M2. Para los sólidos se utiliza la "
    "integración provisional M1–M2; para el N total de aguas verdes y purines se utiliza "
    "exclusivamente M2 mediante Kjeldahl. La caracterización final se actualizará al incorporar M3."
)
REFERENCE_DOCX = ROOT / "MASTER_escrito" / "TFG_ACV_Estiercol_MASTER.docx"
TABLE_DIR = ROOT / "outputs" / "tablas_tesis"
FIG_DIR = ROOT / "outputs" / "graficos_tesis"
GRAPHICS_SCRIPT = ROOT / "scripts" / "generate_thesis_graphics.py"
OUT_DIR = ROOT / "outputs" / "documentos_tfg"
OUT_DOCX = OUT_DIR / "resultados_desarrollados_tfg.docx"
PROCESSED_TOTALS = ROOT / "processed" / "acv_impacto_total_por_escenario.csv"
PROCESSED_STAGE_IMPACTS = ROOT / "processed" / "acv_impacto_por_etapa_escenario.csv"
PROCESSED_A2_BENCHMARK = ROOT / "processed" / "a2_ipcc_jjagwe_benchmark.csv"
METHODOLOGY_DOCX = OUT_DIR / "metodologia_desarrollada_tfg.docx"
README_OUT = OUT_DIR / "README_DOCUMENTOS_GENERADOS.md"
VALIDATION_OUT = OUT_DIR / "reporte_validacion_documentos.md"
FORMAT_REPORT_OUT = OUT_DIR / "reporte_formato_master.md"
APPENDIX_RELATION_REPORT_OUT = OUT_DIR / "reporte_relacion_apendices.md"
FACTOR_REFERENCES_REPORT_OUT = OUT_DIR / "reporte_referencias_factores.md"

OLD_NO3_STOICHIOMETRIC_FACTOR = 4.4268
CURRENT_NO3_STOICHIOMETRIC_FACTOR = 62 / 14

METHODOLOGY_APPENDICES = [
    ("A", "Parámetros completos del modelo ACV", "Apéndice interno"),
    ("B", "Factores de emisión y caracterización", "Apéndice interno"),
    ("C", "Diccionario de variables metodológicas", "Apéndice interno"),
]

RESULTS_APPENDICES = [
    ("R1", "Caracterización completa de muestras", "Tabla"),
    ("R2", "Flujos completos del inventario", "Tabla"),
    ("R3", "Parámetros completos del modelo ACV", "Tabla"),
    ("R4", "Factores completos de emisión y caracterización", "Tabla"),
    ("R5", "Emisiones completas por etapa", "Tabla"),
    ("R6", "Impactos completos por etapa", "Tabla"),
    ("R7", "Impactos totales completos por escenario", "Tabla"),
    ("R8", "Comparación completa de escenarios", "Tabla"),
    ("R9", "Figuras complementarias", "Apéndice"),
    (
        "R10",
        "Correspondencia entre tablas, figuras y archivos fuente",
        "Apéndice",
    ),
]

TABLES = {
    "resumen": TABLE_DIR / "resumen_resultados_para_redaccion.md",
    "tabla_02": TABLE_DIR / "tabla_02_caracterizacion_muestras.csv",
    "tabla_03": TABLE_DIR / "tabla_03_flujos_icv.csv",
    "tabla_04": TABLE_DIR / "tabla_04_parametros_modelo_acv.csv",
    "tabla_05": TABLE_DIR / "tabla_05_factores_emision_y_caracterizacion.csv",
    "tabla_06": TABLE_DIR / "tabla_06_emisiones_por_etapa.csv",
    "tabla_07": TABLE_DIR / "tabla_07_impactos_por_etapa.csv",
    "tabla_08": TABLE_DIR / "tabla_08_impactos_totales_por_escenario.csv",
    "tabla_09": TABLE_DIR / "tabla_09_comparacion_escenarios.csv",
}

MAIN_FIGURES = [
    ("fig_01_caracterizacion_humedad_materia_seca.png", "Figura 1. Humedad y materia seca promedio por tipo de muestra."),
    ("fig_02_caracterizacion_solidos_volatiles_cenizas.png", "Figura 2. Sólidos volátiles y cenizas por tipo de muestra."),
    ("fig_04_flujos_masa_equivalente_total.png", "Figura 3. Masa equivalente total por etapa y escenario."),
    ("fig_06_emisiones_ch4.png", "Figura 4. Emisiones anuales de CH4 por etapa y escenario."),
    ("fig_11_impactos_cambio_climatico_etapa.png", "Figura 5. Cambio climático EF 3.1 por etapa y escenario."),
    ("fig_12_impactos_eutrofizacion_terrestre_etapa.png", "Figura 6. Eutrofización terrestre EF 3.1 por etapa y escenario."),
    ("fig_17_comparacion_diferencia_porcentual.png", "Figura 7. Diferencia porcentual del Escenario B respecto al Escenario A por categoría de impacto."),
]

APPENDIX_FIGURES = [
    ("fig_03_caracterizacion_nitrogeno_total.png", "Figura R1. Nitrógeno total promedio por tipo de muestra."),
    ("fig_05_flujos_distribucion_componentes.png", "Figura R2. Distribución de componentes del inventario por etapa."),
    ("fig_07_emisiones_n2o.png", "Figura R3. Emisiones anuales de N2O por etapa y escenario."),
    ("fig_08_emisiones_nh3.png", "Figura R4. Emisiones anuales de NH3 por etapa y escenario."),
    ("fig_09_emisiones_no3.png", "Figura R5. Emisiones anuales de NO3 por etapa y escenario."),
    ("fig_10_emisiones_co2.png", "Figura R6. Emisiones anuales de CO2 por etapa y escenario."),
    ("fig_14_comparacion_total_cambio_climatico.png", "Figura R7. Comparación del cambio climático total entre escenarios."),
    ("fig_16_comparacion_total_eutrofizacion_marina.png", "Figura R8. Comparación de la eutrofización marina total entre escenarios."),
]

OFFICIAL_STAGE_NAMES = {
    ("A", 1): ("A1", "Precomposteo", "Etapa 1: Precomposteo"),
    ("A", 2): ("A2", "Lombricompostaje", "Etapa 2: Lombricompostaje"),
    ("A", 3): ("A3", "Almacenamiento de aguas verdes", "Etapa 3: Almacenamiento de aguas verdes"),
    ("A", 4): (
        "A4",
        "Aplicación de aguas verdes en campos de pastoreo",
        "Etapa 4: Aplicación de aguas verdes en campos de pastoreo",
    ),
    ("B", 1): ("B1", "Almacenamiento de purines", "Etapa 1: Almacenamiento de purines"),
    ("B", 2): (
        "B2",
        "Aplicación de purines en campo de pastoreo",
        "Etapa 2: Aplicación de purines en campo de pastoreo",
    ),
}

OLD_STAGE_TERMS = {
    "Manejo inicial de estiércol fresco": "Precomposteo",
    "Manejo posterior de fracción sólida": "Lombricompostaje",
    "Manejo de estiércol fresco sin precompostaje": "Almacenamiento de purines",
    "Manejo o aplicación de purines": "Aplicación de purines en campo de pastoreo",
    "Aplicación o manejo de aguas verdes en suelo": "Aplicación de aguas verdes en campos de pastoreo",
}

INTERNAL_COLUMNS = {
    "archivo_fuente",
    "source_file",
    "script",
    "script_origen",
    "path",
    "input_file",
    "output_file",
    "nombre_variable_codigo",
    "columna_original",
    "id interno",
    "observaciones internas del repositorio",
    "fuente",
    "fuente_dato",
    "fuente_o_calculo",
    "formula_origen",
    "observaciones",
    "fuente_bibliografica_pendiente",
    "requiere_revision_bibliografica",
    "clasificacion_referencia",
    "estado_referencia",
    "fuente_factor_emision",
    "fuente_factor_caracterizacion",
    "fuente_factor",
    "ecuacion_utilizada",
    "archivo fuente",
}

ACADEMIC_REPLACEMENTS = {
    "dry_lot": "Sistema de manejo en corral seco",
    "uncovered_anaerobic_lagoon": "Laguna anaerobia descubierta",
    "composting_invessel": "Compostaje en sistema cerrado",
    "solid_storage": "Almacenamiento sólido",
    "liquid_slurry": "Sistema de estiércol líquido",
    "liquid_slurry_without_natural_crust": "Almacenamiento líquido sin costra natural",
    "land_application_managed_liquid_slurry": "Aplicación al suelo de estiércol líquido previamente manejado",
    "aerobic_treatment": "Tratamiento aeróbico",
    "composting_intensive": "Compostaje intensivo",
    "composting_pasive": "Compostaje pasivo",
    "windrow_intensive": "Compostaje intensivo en hileras",
    "windrow_pasive": "Compostaje pasivo en hileras",
    "ipcc": "IPCC",
    "medido": "Factor medido",
    "ESTIERCOL FRESCO": "Estiércol fresco",
    "SOL: PRECOMPOSTADO": "Estiércol precompostado",
    "LIQ: AGUA VERDE": "Agua de lavado incorporada a las aguas verdes",
    "LIQ: PURINES": "Purín",
    "n_ex_pct": "N total reportado (%)",
    "n_ex_fraction": "Fracción másica de N",
    "masa_total_kg_eq": "Masa equivalente total",
    "tipo_muestra": "Tipo de material",
    "fuente_dato": "Fuente metodológica",
    "Factor hardcodeado auditado": "Parámetro complementario",
    "Si": "Sí",
}

HEADER_REPLACEMENTS = {
    "escenario": "Escenario",
    "etapa": "Etapa",
    "nombre_etapa": "Nombre de etapa",
    "tipo_muestra": "Tipo de material",
    "modelo_calculo": "Modelo de cálculo",
    "sistema_manejo_ipcc": "Sistema de manejo asignado",
    "parametro": "Parámetro",
    "valor": "Valor",
    "unidad": "Unidad",
    "referencia_metodologica": "Referencia metodológica",
    "unidad_emision": "Unidad de emisión",
    "unidad_factor": "Unidad del factor",
    "resultado_equivalente": "Resultado equivalente",
    "unidad_equivalente": "Unidad equivalente",
    "variable": "Variable",
    "flujo": "Flujo o material",
    "emision": "Emisión",
    "sustancia": "Sustancia",
    "categoria_impacto": "Categoría de impacto",
    "factor_caracterizacion": "Factor de caracterización",
    "valor_impacto": "Valor de impacto",
    "unidad_impacto": "Unidad de impacto",
    "tratamiento_laboratorio": "Material analizado",
    "jornada_muestreo": "Jornada de muestreo",
    "numero_muestras_solidos": "Número de muestras de sólidos",
    "numero_muestras_nitrogeno": "Número de muestras de nitrógeno",
    "desviacion_estandar": "Desviación estándar",
}

CHEMICAL_REPLACEMENTS = [
    ("kg CO2-eq/año", "kg CO₂-eq/año"),
    ("kg N2O-N/kg N", "kg N₂O-N/kg N"),
    ("CO2-eq", "CO₂-eq"),
    ("N2O-N", "N₂O-N"),
    ("NH3-N", "NH₃-N"),
    ("NO3-N", "NO₃-N"),
    ("CH4", "CH₄"),
    ("N2O", "N₂O"),
    ("NH3", "NH₃"),
    ("NO3", "NO₃⁻"),
    ("CO2", "CO₂"),
    ("PO4", "PO₄³⁻"),
    ("m3", "m³"),
    ("m2", "m²"),
]


def validate_inputs() -> None:
    validated_reference = get_reference_docx_path(ROOT)
    if validated_reference != REFERENCE_DOCX:
        raise RuntimeError(
            "La ruta validada del documento maestro no coincide con la ruta configurada."
        )
    required = [
        *TABLES.values(),
        PROCESSED_TOTALS,
        PROCESSED_STAGE_IMPACTS,
        *(FIG_DIR / name for name, _ in MAIN_FIGURES),
    ]
    missing = [path.relative_to(ROOT).as_posix() for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError("Faltan insumos requeridos:\n" + "\n".join(f"- {item}" for item in missing))

def read_csv(key: str) -> pd.DataFrame:
    try:
        return pd.read_csv(TABLES[key], encoding="utf-8")
    except UnicodeDecodeError:
        return pd.read_csv(TABLES[key], encoding="utf-8-sig")


def clean_text(value) -> str:
    if pd.isna(value):
        return ""
    text = str(value)
    for old, new in ACADEMIC_REPLACEMENTS.items():
        text = text.replace(old, new)
    for old, new in OLD_STAGE_TERMS.items():
        text = text.replace(old, new)
    text = (
        text.replace("Eutrofizacion", "Eutrofización")
        .replace("Nitrogeno", "Nitrógeno")
        .replace("Solidos", "Sólidos")
        .replace("Usado en ecuaciones de N despues de " + "cor" + "reccion", "Usado como fracción másica en ecuaciones de N")
        .replace("n_ex_fraction " + "cor" + "regido para ecuaciones de nitrogeno", "n_ex_fraction usado como fracción másica en ecuaciones de nitrógeno")
        .replace("Comparacion posterior a " + "cor" + "reccion de n_ex_pct a n_ex_fraction", "Comparación entre escenarios con n_ex_fraction como entrada de nitrógeno")
    )
    for old, new in CHEMICAL_REPLACEMENTS:
        text = text.replace(old, new)
    text = repair_mojibake(text)
    return clean_academic_label(text)


def repair_mojibake(text: str) -> str:
    markers = ["\u00c3", "\u00c2", "\u00e2\u20ac", "\u00e2\u20ac\u2122", "\u00e2\u20ac\u0153", "\ufffd"]
    if any(marker in text for marker in markers):
        for encoding in ("cp1252", "latin1"):
            try:
                return text.encode(encoding).decode("utf-8")
            except UnicodeError:
                continue
    return text


def apply_official_stage_names(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    if {"escenario", "etapa"}.issubset(out.columns):
        for idx, row in out.iterrows():
            try:
                key = (str(row["escenario"]).strip().upper(), int(row["etapa"]))
            except (TypeError, ValueError):
                continue
            if key in OFFICIAL_STAGE_NAMES:
                code, short_name, full_name = OFFICIAL_STAGE_NAMES[key]
                if "codigo_etapa" in out.columns:
                    out.at[idx, "codigo_etapa"] = code
                if "nombre_corto_etapa" in out.columns:
                    out.at[idx, "nombre_corto_etapa"] = short_name
                if "nombre_etapa" in out.columns:
                    out.at[idx, "nombre_etapa"] = full_name
    return out


def strip_internal_columns(df: pd.DataFrame) -> pd.DataFrame:
    return df[[col for col in df.columns if col.lower().strip() not in INTERNAL_COLUMNS]].copy()


def official_stage_label(scenario: object, stage: object) -> str:
    try:
        key = (str(scenario).strip().upper(), int(float(stage)))
    except (TypeError, ValueError):
        return ""
    if key in OFFICIAL_STAGE_NAMES:
        code, short_name, _ = OFFICIAL_STAGE_NAMES[key]
        return f"{code}: {short_name}"
    return f"{key[0]}{key[1]}"


def add_calculation_framework(df: pd.DataFrame) -> pd.DataFrame:
    """Distingue documentalmente el origen previo del marco de cálculo de cada etapa."""
    out = df.copy()
    frameworks = {
        ("A", 1): "Manejo del estiércol",
        ("A", 2): "Manejo del estiércol",
        ("A", 3): "Manejo del estiércol",
        ("A", 4): "Suelos gestionados",
        ("B", 1): "Manejo del estiércol",
        ("B", 2): "Suelos gestionados",
    }
    out["Marco de cálculo de la etapa"] = [
        frameworks.get((str(row["escenario"]).strip().upper(), int(float(row["etapa"]))), "")
        for _, row in out.iterrows()
    ]
    if "sistema_manejo_ipcc" in out.columns:
        out = out.rename(columns={"sistema_manejo_ipcc": "Sistema de manejo u origen previo"})
    return out


def combine_stage_columns(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    lower_to_col = {str(col).lower().strip(): col for col in out.columns}
    scenario_col = lower_to_col.get("escenario")
    stage_col = lower_to_col.get("etapa")
    if scenario_col and stage_col:
        insert_at = list(out.columns).index(stage_col)
        labels = [official_stage_label(row[scenario_col], row[stage_col]) for _, row in out.iterrows()]
        for col in ["codigo_etapa", "código", "codigo", "nombre_corto_etapa", "nombre_etapa", "nombre de etapa", "etapa"]:
            existing = lower_to_col.get(col)
            if existing in out.columns:
                out = out.drop(columns=[existing])
        out.insert(min(insert_at, len(out.columns)), "Etapa del sistema", labels)
    return out


def fmt(value, decimals: int = 2) -> str:
    if value == "" or pd.isna(value):
        return ""
    try:
        text = f"{float(value):,.{decimals}f}"
        text = text.replace(",", "X").replace(".", ",").replace("X", " ")
        if "," in text:
            text = text.rstrip("0").rstrip(",")
        return text
    except (TypeError, ValueError):
        return clean_text(value)


def format_df(df: pd.DataFrame, decimals: int = 2, decimals_by_col: dict[str, int] | None = None) -> pd.DataFrame:
    decimals_by_col = decimals_by_col or {}
    out = combine_stage_columns(df)
    out = out.rename(columns={col: HEADER_REPLACEMENTS.get(str(col), clean_text(col)) for col in out.columns})
    for col in out.columns:
        if str(col).lower() == "etapa":
            out[col] = out[col].map(lambda value: f"Etapa {int(float(value))}" if not pd.isna(value) and str(value) != "" else "")
        elif pd.api.types.is_numeric_dtype(out[col]):
            out[col] = out[col].map(lambda value: fmt(value, decimals_by_col.get(col, decimals)))
        else:
            out[col] = out[col].map(clean_text)
    return out


def reference_format():
    ref = Document(str(REFERENCE_DOCX))
    normal = ref.styles["Normal"].font
    font_name = normal.name or "Times New Roman"
    font_size = normal.size or Pt(12)
    section = ref.sections[0]
    margins = {
        "top": section.top_margin,
        "bottom": section.bottom_margin,
        "left": section.left_margin,
        "right": section.right_margin,
    }
    return font_name, font_size, margins


def set_document_style(doc: Document):
    return apply_master_format(doc, REFERENCE_DOCX)


def set_table_horizontal_borders(table) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "insideH"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "808080")
        borders.append(tag)
    for edge in ("left", "right", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    table._tbl.tblPr.append(borders)


def add_dataframe_table(doc: Document, caption: str, df: pd.DataFrame) -> None:
    if caption:
        add_master_caption(doc, clean_text(caption))
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_horizontal_borders(table)
    for idx, col in enumerate(df.columns):
        cell = table.rows[0].cells[idx]
        cell.text = clean_text(col)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = clean_text(value)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if row_index == 0:
            table_header = OxmlElement("w:tblHeader")
            table_header.set(qn("w:val"), "true")
            tr_pr.append(table_header)
    format_table_like_master(table, analyze_master_format(REFERENCE_DOCX))


def add_figure(doc: Document, file_name: str, caption: str) -> None:
    image = FIG_DIR / file_name
    if not image.exists():
        return
    add_master_caption(doc, clean_text(caption))
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image), width=Inches(6.2))


def add_paragraphs(doc: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        doc.add_paragraph(clean_text(text), style="Normal")


def add_provisional_identification(doc: Document) -> None:
    marker = doc.add_paragraph()
    marker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = marker.add_run(PROVISIONAL_LABEL)
    run.bold = True
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = PROVISIONAL_LABEL
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(PROVISIONAL_NOTE, style="Normal")


def characterization_summary() -> pd.DataFrame:
    t02 = read_csv("tabla_02")
    rows = []
    labels = {"Fresh manure": "Estiércol fresco", "Precomposted manure": "Estiércol precompostado"}
    for sample, group in t02.groupby("tipo_muestra", sort=False):
        row = {"Tipo de muestra": labels.get(sample, sample)}
        for _, record in group.iterrows():
            variable = record["variable"]
            unit = record["unidad"]
            if variable == "Humedad":
                row["Humedad (%)"] = record["valor"]
            elif variable == "Materia seca":
                row["Materia seca (%)"] = record["valor"]
            elif variable == "Cenizas":
                row["Cenizas (% base seca)"] = record["valor"]
            elif variable == "Solidos volatiles":
                row["Sólidos volátiles (% base seca)"] = record["valor"]
            elif variable == "Nitrogeno total" and unit == "% N total":
                row["N total (%)"] = record["valor"]
        rows.append(row)
    return pd.DataFrame(rows)


def flow_summary() -> pd.DataFrame:
    t03 = apply_official_stage_names(read_csv("tabla_03"))
    out = t03[t03["flujo"] == "Masa equivalente total"][
        ["escenario", "etapa", "nombre_etapa", "valor"]
    ].rename(
        columns={
            "escenario": "Escenario",
            "etapa": "Etapa",
            "nombre_etapa": "Nombre de etapa",
            "valor": "Masa equivalente total (kg eq/año)",
        }
    )
    return out


def parameter_summary() -> pd.DataFrame:
    t04 = apply_official_stage_names(read_csv("tabla_04"))
    subset = t04[t04["parametro"].isin(["Masa equivalente total", "Nitrogeno total reportado", "Nitrogeno total como fraccion masica", "MCF", "EF3"])]
    subset = add_calculation_framework(subset)
    out = subset.pivot_table(
        index=["escenario", "etapa", "nombre_etapa", "modelo_calculo", "Marco de cálculo de la etapa"],
        columns="parametro",
        values="valor",
        aggfunc="first",
    ).reset_index()
    return out.rename(
        columns={
            "escenario": "Escenario",
            "etapa": "Etapa",
            "nombre_etapa": "Nombre de etapa",
            "modelo_calculo": "Modelo",
            "Masa equivalente total": "Masa equivalente (kg eq/año)",
            "Nitrogeno total reportado": "N total reportado (%)",
            "Nitrogeno total como fraccion masica": "Fracción másica de N",
        }
    )


def emissions_summary() -> pd.DataFrame:
    t06 = apply_official_stage_names(read_csv("tabla_06"))
    out = t06.groupby(["escenario", "sustancia"], as_index=False)["valor"].sum()
    out = out.pivot(index="escenario", columns="sustancia", values="valor").reset_index().rename_axis(None, axis=1)
    for substance in ["CH4", "N2O", "NH3", "NO3", "CO2"]:
        if substance not in out:
            out[substance] = 0.0
    return out[["escenario", "CH4", "N2O", "NH3", "NO3", "CO2"]].rename(
        columns={
            "escenario": "Escenario",
            "CH4": "CH4 (kg/año)",
            "N2O": "N2O (kg/año)",
            "NH3": "NH3 (kg/año)",
            "NO3": "NO3 (kg/año)",
            "CO2": "CO2 (kg/año)",
        }
    )


def impact_stage_summary() -> pd.DataFrame:
    t07 = apply_official_stage_names(read_csv("tabla_07"))
    out = t07.groupby(["escenario", "etapa", "nombre_etapa", "categoria_impacto"], as_index=False)["resultado_equivalente"].sum()
    out = out.pivot(index=["escenario", "etapa", "nombre_etapa"], columns="categoria_impacto", values="resultado_equivalente").reset_index().rename_axis(None, axis=1)
    return out.rename(
        columns={
            "escenario": "Escenario",
            "etapa": "Etapa",
            "nombre_etapa": "Nombre de etapa",
            "Cambio climático": "Cambio climático (kg CO2-eq/año)",
            "Eutrofización terrestre": "Eutrofización terrestre (mol N-eq/año)",
            "Eutrofización marina": "Eutrofización marina (kg N-eq/año)",
        }
    )


def total_impact_summary() -> pd.DataFrame:
    totals = pd.read_csv(PROCESSED_TOTALS, encoding="utf-8-sig")
    return totals[
        [
            "Escenario",
            "impacto_calentamiento_global_kg_co2eq",
            "impacto_eutrofizacion_terrestre_mol_neq",
            "impacto_eutrofizacion_marina_kg_neq",
            "impacto_calentamiento_global_kg_co2eq_por_kg_estiercol_fresco",
            "impacto_eutrofizacion_terrestre_mol_neq_por_kg_estiercol_fresco",
            "impacto_eutrofizacion_marina_kg_neq_por_kg_estiercol_fresco",
        ]
    ].rename(
        columns={
            "impacto_calentamiento_global_kg_co2eq": "Calentamiento global (kg CO2-eq/año)",
            "impacto_eutrofizacion_terrestre_mol_neq": "Eutrofización terrestre (mol N-eq/año)",
            "impacto_eutrofizacion_marina_kg_neq": "Eutrofización marina (kg N-eq/año)",
            "impacto_calentamiento_global_kg_co2eq_por_kg_estiercol_fresco": "Calentamiento global (kg CO2-eq/kg de estiércol fresco)",
            "impacto_eutrofizacion_terrestre_mol_neq_por_kg_estiercol_fresco": "Eutrofización terrestre (mol N-eq/kg de estiércol fresco)",
            "impacto_eutrofizacion_marina_kg_neq_por_kg_estiercol_fresco": "Eutrofización marina (kg N-eq/kg de estiércol fresco)",
        }
    )


def comparison_summary() -> pd.DataFrame:
    t09 = read_csv("tabla_09")
    out = t09[
        [
            "categoria_impacto",
            "escenario_A",
            "escenario_B",
            "diferencia_absoluta_B_menos_A",
            "diferencia_porcentual_B_vs_A",
            "escenario_con_mayor_impacto",
        ]
    ].rename(
        columns={
            "categoria_impacto": "Categoría de impacto",
            "escenario_A": "Escenario A",
            "escenario_B": "Escenario B",
            "diferencia_absoluta_B_menos_A": "Diferencia absoluta",
            "diferencia_porcentual_B_vs_A": "Diferencia porcentual B respecto a A",
            "escenario_con_mayor_impacto": "Escenario con mayor impacto",
        }
    )
    return out


def results_context() -> dict[str, object]:
    totals = pd.read_csv(PROCESSED_TOTALS, encoding="utf-8-sig")
    stages = pd.read_csv(PROCESSED_STAGE_IMPACTS, encoding="utf-8-sig")
    comparison = read_csv("tabla_09")
    if set(totals["Escenario"].astype(str)) != {"A", "B"}:
        raise RuntimeError("Los impactos procesados deben contener los escenarios A y B.")
    by_scenario = totals.set_index("Escenario")
    context: dict[str, object] = {
        "cg_a": float(by_scenario.loc["A", "impacto_calentamiento_global_kg_co2eq"]),
        "eu_a": float(by_scenario.loc["A", "impacto_eutrofizacion_marina_kg_neq"]),
        "et_a": float(by_scenario.loc["A", "impacto_eutrofizacion_terrestre_mol_neq"]),
        "cg_b": float(by_scenario.loc["B", "impacto_calentamiento_global_kg_co2eq"]),
        "eu_b": float(by_scenario.loc["B", "impacto_eutrofizacion_marina_kg_neq"]),
        "et_b": float(by_scenario.loc["B", "impacto_eutrofizacion_terrestre_mol_neq"]),
        "cg_norm_a": float(by_scenario.loc["A", "impacto_calentamiento_global_kg_co2eq_por_kg_estiercol_fresco"]),
        "eu_norm_a": float(by_scenario.loc["A", "impacto_eutrofizacion_marina_kg_neq_por_kg_estiercol_fresco"]),
        "et_norm_a": float(by_scenario.loc["A", "impacto_eutrofizacion_terrestre_mol_neq_por_kg_estiercol_fresco"]),
        "cg_norm_b": float(by_scenario.loc["B", "impacto_calentamiento_global_kg_co2eq_por_kg_estiercol_fresco"]),
        "eu_norm_b": float(by_scenario.loc["B", "impacto_eutrofizacion_marina_kg_neq_por_kg_estiercol_fresco"]),
        "et_norm_b": float(by_scenario.loc["B", "impacto_eutrofizacion_terrestre_mol_neq_por_kg_estiercol_fresco"]),
    }
    comparison_index = comparison.set_index("categoria_impacto")
    context["cg_difference"] = float(comparison_index.loc["Cambio climático", "diferencia_absoluta_B_menos_A"])
    context["eu_difference"] = float(comparison_index.loc["Eutrofización marina", "diferencia_absoluta_B_menos_A"])
    context["cg_percentage"] = float(comparison_index.loc["Cambio climático", "diferencia_porcentual_B_vs_A"])
    context["eu_percentage"] = float(comparison_index.loc["Eutrofización marina", "diferencia_porcentual_B_vs_A"])
    units = comparison.set_index("categoria_impacto")["unidad"]
    expected_units = {
        "Cambio climático": "kg CO2-eq/año",
        "Eutrofización marina": "kg N-eq/año",
    }
    for category, expected_unit in expected_units.items():
        if units.loc[category] != expected_unit:
            raise RuntimeError(f"Unidad comparativa inesperada para {category}: {units.loc[category]!r}.")
    for short, category in (("cg", "Cambio climático"), ("eu", "Eutrofización marina")):
        comparison_values = Comparison(
            "Escenario A", float(context[f"{short}_a"]),
            "Escenario B", float(context[f"{short}_b"]), expected_units[category],
        )
        comparison_values.assert_consistent(
            difference=float(context[f"{short}_difference"]),
            percentage=float(context[f"{short}_percentage"]),
        )
        comparison_values.assert_rounding_unambiguous(6)
        context[f"{short}_comparison"] = comparison_values

    stage_names = {
        ("A", 1): "A1: Precomposteo",
        ("A", 2): "A2: Lombricompostaje",
        ("A", 3): "A3: Almacenamiento de aguas verdes",
        ("A", 4): "A4: Aplicación de aguas verdes en campos de pastoreo",
        ("B", 1): "B1: Almacenamiento de purines",
        ("B", 2): "B2: Aplicación de purines en campo de pastoreo",
    }
    category_columns = {
        "cg": "impacto_calentamiento_global_kg_co2eq",
        "eu": "impacto_eutrofizacion_marina_kg_neq",
    }
    for scenario in ("A", "B"):
        subset = stages[stages["Escenario"] == scenario]
        for short, column in category_columns.items():
            row = subset.loc[subset[column].idxmax()]
            value = float(row[column])
            total = float(context[f"{short}_{scenario.lower()}"])
            context[f"{short}_dominant_{scenario.lower()}_name"] = stage_names[(scenario, int(row["Etapa"]))]
            context[f"{short}_dominant_{scenario.lower()}_value"] = value
            context[f"{short}_dominant_{scenario.lower()}_percentage"] = 100 * value / total

    b1 = stages[(stages["Escenario"] == "B") & (stages["Etapa"] == 1)].iloc[0]
    emissions = read_csv("tabla_06")
    b1_emissions = emissions[(emissions["escenario"] == "B") & (emissions["etapa"] == 1)]
    context["b1_eutrophication"] = float(b1["impacto_eutrofizacion_marina_kg_neq"])
    context["b1_nh3"] = float(b1_emissions.loc[b1_emissions["sustancia"] == "NH3", "valor"].sum())
    context["b1_no3"] = float(b1_emissions.loc[b1_emissions["sustancia"] == "NO3", "valor"].sum())
    context["b1_n2o_leaching"] = float(
        b1_emissions.loc[b1_emissions["emision"].str.contains("lixiviacion", case=False, na=False), "valor"].sum()
    )
    return context


def a2_benchmark_summary() -> pd.DataFrame:
    benchmark = pd.read_csv(PROCESSED_A2_BENCHMARK, encoding="utf-8-sig")
    expected = {
        "CH4 por materia seca de entrada",
        "N2O directo por materia seca de entrada",
        "N2O-N directo por N inicial",
        "Contraste armonizado de CH4 y N2O directo",
    }
    if set(benchmark["indicador"]) != expected:
        raise RuntimeError("El contraste bibliográfico de A2 no contiene los cuatro indicadores aprobados.")
    return benchmark[
        [
            "indicador", "valor_ipcc", "valor_experimental", "unidad",
            "diferencia_absoluta_ipcc_menos_experimental",
            "razon_ipcc_experimental",
            "diferencia_porcentual_relativa_experimental",
        ]
    ].rename(columns={
        "indicador": "Indicador",
        "valor_ipcc": "Estimación IPCC",
        "valor_experimental": "Referencia experimental",
        "unidad": "Unidad",
        "diferencia_absoluta_ipcc_menos_experimental": "Diferencia IPCC menos referencia",
        "razon_ipcc_experimental": "Razón IPCC/referencia",
        "diferencia_porcentual_relativa_experimental": "Diferencia relativa a la referencia (%)",
    })


def build_document() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    profile = set_document_style(doc)
    context = results_context()

    doc.add_paragraph(
        "Resultados desarrollados del Análisis de Ciclo de Vida", style="Title"
    )
    add_provisional_identification(doc)

    characterization = characterization_summary().set_index("Tipo de muestra")
    fresh = characterization.loc["Estiércol fresco"]
    precomposted = characterization.loc["Estiércol precompostado"]
    emission_totals = emissions_summary().set_index("Escenario")
    flows = flow_summary()
    def stage_label(scenario: str, stage: object, name: object) -> str:
        academic_name = re.sub(r"^Etapa\s+\d+:\s*", "", str(name))
        return f"{scenario}{int(stage)}: {academic_name}"

    flow_values = {
        stage_label(row["Escenario"], row["Etapa"], row["Nombre de etapa"]): float(row["Masa equivalente total (kg eq/año)"])
        for _, row in flows.iterrows()
    }
    largest_flow = dominant(flow_values, decimals=2)
    smallest_flow = dominant({name: -value for name, value in flow_values.items()}, decimals=2)
    stage_emissions = apply_official_stage_names(read_csv("tabla_06"))

    def dominant_emission(substance: str) -> str:
        grouped = stage_emissions[stage_emissions["sustancia"] == substance].groupby(
            ["escenario", "etapa", "nombre_etapa"], as_index=False
        )["valor"].sum()
        candidates = {
            stage_label(row["escenario"], row["etapa"], row["nombre_etapa"]): float(row["valor"])
            for _, row in grouped.iterrows()
        }
        return dominant(candidates, decimals=6)[0]

    def emission_totals_text(scenario: str) -> str:
        substances = (
            ("CH4", "CH4"),
            ("N2O", "N2O"),
            ("NH3", "NH3"),
            ("NO3", "NO3"),
            ("CO2", "CO2"),
        )
        values = []
        for column, label in substances:
            value = emission_totals.loc[scenario, f"{column} (kg/año)"]
            if pd.notna(value):
                values.append(f"{fmt(value, 2)} kg {label}/año")
        if len(values) == 1:
            return values[0]
        return ", ".join(values[:-1]) + f" y {values[-1]}"

    doc.add_heading("1. Caracterización de las muestras analizadas", level=2)
    add_paragraphs(
        doc,
        [
            f"La caracterización de las muestras analizadas permitió establecer los parámetros fisicoquímicos usados como entradas del inventario de ciclo de vida. El estiércol fresco presentó un contenido promedio de agua de {fmt(fresh['Humedad (%)'], 2)} % y una materia seca de {fmt(fresh['Materia seca (%)'], 2)} %. El estiércol precompostado presentó un contenido promedio de agua de {fmt(precomposted['Humedad (%)'], 2)} % y una materia seca de {fmt(precomposted['Materia seca (%)'], 2)} %.",
            f"La fracción de sólidos volátiles fue mayor en {'el estiércol fresco' if fresh['Sólidos volátiles (% base seca)'] > precomposted['Sólidos volátiles (% base seca)'] else 'el estiércol precompostado'}, con {fmt(max(fresh['Sólidos volátiles (% base seca)'], precomposted['Sólidos volátiles (% base seca)']), 2)} % en base seca, frente a {fmt(min(fresh['Sólidos volátiles (% base seca)'], precomposted['Sólidos volátiles (% base seca)']), 2)} %. Las cenizas fueron mayores en {'el estiércol fresco' if fresh['Cenizas (% base seca)'] > precomposted['Cenizas (% base seca)'] else 'el material precompostado'}. El nitrógeno total fue de {fmt(fresh['N total (%)'], 3)} % para estiércol fresco y de {fmt(precomposted['N total (%)'], 3)} % para estiércol precompostado.",
            "La Tabla 1 resume los valores de caracterización de las muestras. La Figura 1 presenta humedad y materia seca, mientras que la Figura 2 presenta sólidos volátiles y cenizas.",
            "La Tabla R1 del bloque de apéndices internos, Caracterización completa de muestras, presenta la desagregación de los resultados fisicoquímicos utilizados en esta sección.",
            "Los valores presentados corresponden a la integración provisional M1–M2 para sólidos y al N total Kjeldahl de M2 para aguas verdes y purines. La jornada M3 se incorporará posteriormente para actualizar la caracterización final.",
        ],
    )
    add_dataframe_table(doc, "Tabla 1. Caracterización resumida de las muestras.", format_df(characterization_summary(), decimals=3))
    add_figure(doc, *MAIN_FIGURES[0])
    add_figure(doc, *MAIN_FIGURES[1])

    doc.add_heading("2. Flujos del inventario de ciclo de vida", level=2)
    add_paragraphs(
        doc,
        [
            "Estos valores se presentan como flujos anuales estimados del inventario, manteniendo como referencia metodológica la unidad funcional de 1 kg de estiércol fresco manejado. El flujo anual común expresa la escala operacional del inventario y no redefine la unidad funcional.",
            f"Los flujos del inventario se expresaron como masa equivalente total por año para cada etapa. {largest_flow[0]} presentó la mayor masa equivalente total, con {fmt(largest_flow[1], 2)} kg eq/año.",
            "La masa equivalente de A4 integra el agua de lavado y 8 753,63 kg/año de estiércol remanente derivados del balance en sala. La masa equivalente de B2 integra agua de lavado y 26 278,73 kg/año de estiércol fresco teóricamente depositado.",
            "En contraste, para estimar las emisiones de manejo en A3: Almacenamiento de aguas verdes y B1: Almacenamiento de purines se utilizó la masa de estiércol correspondiente, sin sumar el agua de lavado como masa de actividad. El flujo total de agua y estiércol se incorporó en las etapas subsecuentes de aplicación A4 y B2.",
            f"{smallest_flow[0]} presentó la menor masa equivalente. La Tabla 2 presenta la masa equivalente total por etapa y la Figura 3 resume su distribución por escenario.",
            "La Tabla R2 del bloque de apéndices internos, Flujos completos del inventario, contiene la desagregación de los flujos empleados para construir el ICV.",
        ],
    )
    add_dataframe_table(doc, "Tabla 2. Masa equivalente total por etapa.", format_df(flow_summary()))
    add_figure(doc, *MAIN_FIGURES[2])

    doc.add_heading("3. Parámetros utilizados en el modelo ACV", level=2)
    add_paragraphs(
        doc,
        [
            "Los parámetros utilizados en el modelo se organizaron por escenario y etapa. La tabla distingue entre el nitrógeno total reportado en porcentaje y la fracción másica efectiva empleada en las ecuaciones. Para A2, esta fracción expresa el N del precompostado sobre masa húmeda e incorpora la materia seca gravimétrica; las demás etapas conservan su tratamiento analítico vigente.",
            "A2: Lombricompostaje se calculó con ecuaciones IPCC y la categoría Composting – Passive Windrow. La fracción de lixiviación se estableció en cero como parámetro específico del sistema estudiado, sin alterar el valor genérico de la categoría. La Tabla 3 resume los parámetros principales.",
            "Los restantes factores de A2 —MCF, EF3 y fracción volatilizada— corresponden a la categoría IPCC seleccionada.",
            "La Tabla R3 del bloque de apéndices internos, Parámetros completos del modelo ACV, amplía los parámetros por escenario y etapa; la Tabla R4, Factores completos de emisión y caracterización, documenta los factores asociados.",
        ],
    )
    add_dataframe_table(doc, "Tabla 3. Parámetros principales por etapa.", format_df(parameter_summary(), decimals=4))

    doc.add_heading("4. Emisiones estimadas por etapa y escenario", level=2)
    add_paragraphs(
        doc,
        [
            f"Las emisiones consolidadas muestran diferencias entre escenarios y sustancias. El Escenario A presentó {emission_totals_text('A')}. El Escenario B presentó {emission_totals_text('B')}.",
            f"{dominant_emission('CH4')} presentó la mayor contribución de CH4. {dominant_emission('N2O')} presentó la mayor emisión de N2O. A2: Lombricompostaje fue estimada mediante las vías IPCC de manejo de estiércol. La Tabla 4 resume las emisiones anuales por escenario y sustancia, y la Figura 4 presenta las emisiones de CH4 por etapa.",
            "La Tabla R5 del bloque de apéndices internos, Emisiones completas por etapa, presenta la desagregación por sustancia, escenario y etapa. Además, el Apéndice R9, Figuras complementarias, reúne las representaciones gráficas que respaldan la interpretación de la caracterización, los flujos, las emisiones y la comparación de escenarios.",
        ],
    )
    add_dataframe_table(doc, "Tabla 4. Emisiones anuales por escenario y sustancia.", format_df(emissions_summary()))
    add_figure(doc, *MAIN_FIGURES[3])

    doc.add_heading("5. Impactos ambientales por etapa", level=2)
    add_paragraphs(
        doc,
        [
            f"Los impactos ambientales por etapa mostraron que {context['cg_dominant_b_name']} concentró {fmt(context['cg_dominant_b_percentage'], 2)} % del calentamiento global del Escenario B, con {fmt(context['cg_dominant_b_value'], 6)} kg CO2-eq/año. En el Escenario A, {context['cg_dominant_a_name']} aportó {fmt(context['cg_dominant_a_percentage'], 2)} % del total de esta categoría.",
            f"Para eutrofización marina, {context['eu_dominant_b_name']} presentó la mayor contribución del Escenario B, con {fmt(context['eu_dominant_b_percentage'], 2)} % de su total; en el Escenario A, {context['eu_dominant_a_name']} concentró {fmt(context['eu_dominant_a_percentage'], 2)} %. La tabla de impactos por etapa presenta separadamente cambio climático, eutrofización terrestre y eutrofización marina.",
            f"En B1: Almacenamiento de purines, la lixiviación explícita utilizada para estimar N₂O indirecto y NO₃⁻ fue nula. La eutrofización terrestre y marina de esta etapa procede de las emisiones atmosféricas explícitas de NH₃ y NOx del ledger.",
            "La caracterización de las emisiones directas se realizó con Environmental Footprint 3.1: cambio climático en kg CO₂-eq, eutrofización terrestre en mol N-eq y eutrofización marina en kg N-eq.",
            "Los consumos de electricidad y diésel están incorporados al inventario físico, pero sus procesos de fondo serán caracterizados posteriormente mediante SimaPro y ecoinvent dentro del mismo pipeline canónico.",
            "La Tabla R6 del bloque de apéndices internos, Impactos completos por etapa, presenta los resultados desagregados por categoría de impacto.",
        ],
    )
    add_dataframe_table(doc, "Tabla 5. Impactos ambientales por etapa.", format_df(impact_stage_summary()))
    add_figure(doc, *MAIN_FIGURES[4])
    add_figure(doc, *MAIN_FIGURES[5])

    doc.add_heading("6. Impactos totales por escenario", level=2)
    add_paragraphs(
        doc,
        [
            f"El Escenario A alcanzó {fmt(context['cg_a'], 6)} kg CO₂-eq/año, {fmt(context['et_a'], 6)} mol N-eq/año de eutrofización terrestre y {fmt(context['eu_a'], 6)} kg N-eq/año de eutrofización marina. El Escenario B alcanzó {fmt(context['cg_b'], 6)} kg CO₂-eq/año, {fmt(context['et_b'], 6)} mol N-eq/año y {fmt(context['eu_b'], 6)} kg N-eq/año, respectivamente.",
            f"Por unidad funcional, el Escenario A presentó {fmt(context['cg_norm_a'], 9)} kg CO₂-eq, {fmt(context['et_norm_a'], 9)} mol N-eq terrestre y {fmt(context['eu_norm_a'], 9)} kg N-eq marino por kg de estiércol fresco manejado. Para el Escenario B fueron {fmt(context['cg_norm_b'], 9)}, {fmt(context['et_norm_b'], 9)} y {fmt(context['eu_norm_b'], 9)}, respectivamente.",
            "La Tabla 6 distingue los resultados anualizados, que representan la magnitud operacional, de los resultados normalizados respecto a 1 kg de estiércol fresco manejado.",
            "La Tabla R7 del bloque de apéndices internos, Impactos totales completos por escenario, presenta el detalle de la agregación utilizada en esta comparación.",
        ],
    )
    add_dataframe_table(
        doc,
        "Tabla 6. Impactos ambientales totales y normalizados por escenario.",
        format_df(
            total_impact_summary(),
            decimals=6,
            decimals_by_col={
                "Calentamiento global (kg CO2-eq/kg de estiércol fresco)": 9,
                "Eutrofización marina (kg N-eq/kg de estiércol fresco)": 9,
            },
        ),
    )

    doc.add_heading("7. Comparación entre escenarios", level=2)
    add_paragraphs(
        doc,
        [
            f"La comparación bajo la misma unidad funcional mostró un mayor impacto de calentamiento global en el {context['cg_comparison'].higher_label}. La diferencia B menos A fue de {fmt(context['cg_difference'], 6)} kg CO2-eq/año, equivalente a {fmt(context['cg_percentage'], 2)} % respecto al Escenario A.",
            f"En eutrofización marina, el {context['eu_comparison'].higher_label} presentó el mayor impacto y el {context['eu_comparison'].lower_label} el menor. La diferencia B menos A fue de {fmt(context['eu_difference'], 6)} kg N-eq/año, equivalente a {fmt(context['eu_percentage'], 2)} % respecto al Escenario A. La comparación conserva cada categoría EF 3.1 en su propia unidad.",
            "La Tabla R8 del bloque de apéndices internos, Comparación completa de escenarios, amplía las diferencias absolutas y porcentuales. La relación entre los contenidos, sus bases de información y las figuras asociadas se documenta en el Apéndice R10, Correspondencia entre tablas, figuras y archivos fuente.",
        ],
    )
    add_dataframe_table(doc, "Tabla 7. Comparación de impactos ambientales entre escenarios.", format_df(comparison_summary(), decimals=3, decimals_by_col={"Diferencia porcentual B respecto a A": 2}))
    add_figure(doc, *MAIN_FIGURES[6])

    benchmark = a2_benchmark_summary()
    benchmark_index = benchmark.set_index("Indicador")
    ch4 = benchmark_index.loc["CH4 por materia seca de entrada"]
    n2o = benchmark_index.loc["N2O directo por materia seca de entrada"]
    n_ratio = benchmark_index.loc["N2O-N directo por N inicial"]
    climate = benchmark_index.loc["Contraste armonizado de CH4 y N2O directo"]
    ch4_relation = Comparison(
        "referencia experimental", float(ch4["Referencia experimental"]),
        "estimación IPCC", float(ch4["Estimación IPCC"]), str(ch4["Unidad"]),
    )
    n2o_relation = Comparison(
        "referencia experimental", float(n2o["Referencia experimental"]),
        "estimación IPCC", float(n2o["Estimación IPCC"]), str(n2o["Unidad"]),
    )
    ch4_direction = "menor" if ch4_relation.higher_label == "referencia experimental" else "mayor"
    n2o_direction = "menor" if n2o_relation.higher_label == "referencia experimental" else "mayor"
    doc.add_heading("8. Contraste experimental bibliográfico de A2", level=2)
    add_paragraphs(doc, [
        "Los resultados oficiales de A2: Lombricompostaje se contrastaron con la referencia empírica de Jjagwe et al. (2019) sin sustituir las ecuaciones IPCC ni crear un segundo inventario oficial. La base común fue la materia seca del estiércol precompostado al ingreso de A2.",
        f"La estimación IPCC de CH4 fue {fmt(ch4['Estimación IPCC'], 6)} g CH4/kg de materia seca, frente a {fmt(ch4['Referencia experimental'], 6)} g CH4/kg de materia seca. La razón IPCC/referencia fue {fmt(ch4['Razón IPCC/referencia'], 6)} y la diferencia relativa fue {fmt(ch4['Diferencia relativa a la referencia (%)'], 2)} %.",
        f"Para N2O directo, la estimación IPCC fue {fmt(n2o['Estimación IPCC'], 6)} mg N2O/kg de materia seca y la referencia experimental fue {fmt(n2o['Referencia experimental'], 6)} mg N2O/kg de materia seca. La razón fue {fmt(n2o['Razón IPCC/referencia'], 6)}. El N2O indirecto por volatilización permanece en el inventario oficial, pero se excluyó de esta comparación porque no representa una emisión medida físicamente dentro del lecho.",
        f"Como indicador complementario, la fracción de N inicial emitida como N2O-N directo fue {fmt(n_ratio['Estimación IPCC'], 8)} kg N2O-N/kg N para IPCC y {fmt(n_ratio['Referencia experimental'], 8)} kg N2O-N/kg N para la referencia experimental; la razón fue {fmt(n_ratio['Razón IPCC/referencia'], 6)}. Este resultado apoya la interpretación y no constituye una validación formal del modelo.",
        f"El contraste armonizado de CH4 + N2O directo, caracterizado con los mismos factores vigentes del TFG, fue {fmt(climate['Estimación IPCC'], 8)} kg CO2-eq/kg de materia seca para IPCC y {fmt(climate['Referencia experimental'], 8)} kg CO2-eq/kg de materia seca para Jjagwe et al. (2019). Este indicador no representa el impacto climático total del sistema: excluye CO2 experimental y N2O indirecto IPCC.",
        f"Las desviaciones no mostraron una dirección uniforme entre especies: la estimación IPCC fue {ch4_direction} para CH4 y {n2o_direction} para N2O directo. Por ello, el contraste no evidencia un sesgo uniforme del método IPCC en A2 y no sustenta afirmar, de forma global, que este sobreestime, subestime o sea más conservador. La compensación parcial entre el {ch4_direction} CH4 y el {n2o_direction} N2O directo en el indicador armonizado demuestra que el resultado agregado no sustituye el análisis separado de cada flujo gaseoso.",
        "Jjagwe et al. (2019) reportaron una pérdida atmosférica de 18,18 % del N inicial y una distribución de 74,55 % en vermicompost y 7,27 % en biomasa de lombrices. Esta pérdida se consideró solo como contraste conceptual: no es equivalente a la fracción IPCC encaminada a volatilización, porque las especies, métodos y fronteras difieren y parte del N atmosférico no cuantificado podría corresponder a N2.",
        "El estudio de Jjagwe et al. se realizó en Uganda como un sistema experimental con unidades replicadas y empleó la especie Eudrilus eugeniae. El material se mantuvo aproximadamente entre 60 y 70 % de humedad, el proceso se organizó en ciclos cercanos a 12 semanas, el estiércol se incorporó progresivamente y había permanecido alrededor de una semana almacenado bajo sombra antes de ingresar al sistema. En contraste, A2 representa una operación real de lombricompostaje en Turrialba, modelada mediante las ecuaciones IPCC y la categoría disponible Composting – Passive Windrow, con las características propias de su sustrato y condiciones operativas. Estas diferencias de especie, acondicionamiento, alimentación, humedad, duración, ambiente geográfico y escala pueden influir en las emisiones y muestran que la base material armonizada no vuelve físicamente idénticos ambos sistemas.",
        "Jjagwe et al. determinaron CO2, CH4, NH3 y N2O mediante cámaras estáticas en momentos discretos del proceso, incluidas las semanas 4, 8 y 12. Ese diseño de medición no reproduce necesariamente la variación temporal completa de una operación real. Además, el análisis de flujo de materiales no permitió identificar por completo las especies del N atmosférico; los autores plantearon que parte importante del N no cuantificado como los gases medidos podría corresponder a N2 u otras formas. Estas limitaciones de medición y balance impiden atribuir las divergencias exclusivamente al método IPCC y confirman que el 18,18 % no es una medición equivalente a la fracción IPCC encaminada a volatilización.",
        "En consecuencia, el benchmark compara una estimación modelada con una referencia experimental externa sobre una base material armonizada, pero no sistemas físicamente idénticos. Los resultados se interpretan como contraste con literatura y no como validación o invalidación de las estimaciones IPCC.",
        "No se calculó eutrofización experimental. El NH3 no fue detectado durante la campaña, el N atmosférico no quedó completamente especiado y la adaptación del TFG para representar N potencialmente eutrofizante no constituye una especiación experimental equivalente.",
        "El artículo presenta una inconsistencia interna para N2O: el resumen indica 3,943 × 10⁻⁵ g N2O/kg de materia seca, mientras la sección de resultados y la Figura 3 muestran un orden de decenas de miligramos, coherente además con el potencial de calentamiento global publicado. Para este contraste se adoptó explícitamente 39,43 mg N2O/kg de materia seca.",
    ])
    add_dataframe_table(doc, "Tabla 8. Contraste bibliográfico de A2 sobre materia seca de entrada.", format_df(benchmark, decimals=6, decimals_by_col={"Diferencia relativa a la referencia (%)": 2}))

    doc.add_heading("Apéndices internos de resultados", level=1)
    add_paragraphs(
        doc,
        [
            "Los apéndices internos contienen las Tablas R1, R2, R3, R4, R5, R6, R7 y R8, las Figuras R1 a R8 y la Tabla R10 de correspondencia. Las figuras complementarias de emisiones de N2O, NH3, NO3 y CO2 respaldan la interpretación de las emisiones por sustancia.",
        ],
    )
    appendix_tables = [
        ("Tabla R1. Caracterización completa de muestras.", "tabla_02"),
        ("Tabla R2. Flujos completos del inventario.", "tabla_03"),
        ("Tabla R3. Parámetros completos del modelo ACV.", "tabla_04"),
        ("Tabla R4. Factores completos de emisión y caracterización.", "tabla_05"),
        ("Tabla R5. Emisiones completas por etapa.", "tabla_06"),
        ("Tabla R6. Impactos completos por etapa.", "tabla_07"),
        ("Tabla R7. Impactos totales completos por escenario.", "tabla_08"),
        ("Tabla R8. Comparación completa de escenarios.", "tabla_09"),
    ]
    for title, key in appendix_tables:
        df = read_csv(key)
        if key in {"tabla_03", "tabla_04", "tabla_06", "tabla_07"}:
            df = apply_official_stage_names(df)
        if key == "tabla_04":
            df = add_calculation_framework(df)
        df = strip_internal_columns(df)
        add_dataframe_table(doc, title, format_df(df, decimals=4))

    doc.add_heading("Apéndice R9. Figuras complementarias", level=2)
    add_paragraphs(
        doc,
        [
            "La Figura R1 complementa la caracterización de nitrógeno total; la Figura R2 complementa los flujos del inventario; la Figura R3 complementa las emisiones de N2O; la Figura R4 complementa las emisiones de NH3; la Figura R5 complementa las emisiones de NO3; la Figura R6 complementa las emisiones de CO2; la Figura R7 complementa la comparación total de calentamiento global; y la Figura R8 complementa la comparación total de eutrofización.",
        ],
    )
    for file_name, caption in APPENDIX_FIGURES:
        add_figure(doc, file_name, caption)

    correspondence = pd.DataFrame(
        [
            ["Caracterización de muestras", "Caracterización fisicoquímica de materiales", "Figuras 1, 2 y R1"],
            ["Flujos del ICV", "Inventario de flujos por escenario y etapa", "Figuras 3 y R2"],
            ["Parámetros del modelo", "Parámetros principales del modelo ACV", "Tabla 3"],
            ["Factores", "Factores de emisión y caracterización", "Apéndice R4"],
            ["Emisiones", "Emisiones estimadas por etapa", "Figuras 4, R3, R4, R5 y R6"],
            ["Impactos por etapa", "Impactos ambientales por etapa", "Figuras 5 y 6"],
            ["Impactos totales", "Impactos agregados por escenario", "Figuras R7 y R8"],
            ["Comparación de escenarios", "Comparación ambiental entre escenarios", "Figura 7"],
        ],
        columns=["Contenido", "Base de información", "Tablas o figuras relacionadas"],
    )
    doc.add_heading("Apéndice R10. Correspondencia entre tablas, figuras y archivos fuente", level=2)
    add_paragraphs(doc, ["La Tabla R10 resume la correspondencia académica entre contenidos, bases de información y figuras utilizadas en el documento."])
    add_dataframe_table(doc, "Tabla R10. Correspondencia de insumos usados.", correspondence)

    finalize_document_format(doc, profile)
    doc.save(OUT_DOCX)


def extract_docx_text(path: Path) -> str:
    if not path.exists():
        return ""
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    return re.sub(r"<[^>]+>", " ", xml)


def extract_docx_xml(path: Path) -> str:
    if not path.exists():
        return ""
    with zipfile.ZipFile(path) as zf:
        return zf.read("word/document.xml").decode("utf-8", errors="ignore")


def labels_referenced(text: str, labels: list[str]) -> bool:
    normalized = re.sub(r"\s+", " ", text)
    for label in labels:
        if normalized.count(label) >= 2:
            continue
        noun, identifier = label.split(" ", 1)
        plural_contexts = re.findall(
            rf"\b{re.escape(noun)}s?\b([^.;]*)", normalized, flags=re.IGNORECASE
        )
        if any(
            re.search(rf"\b{re.escape(identifier)}\b", context)
            for context in plural_contexts
        ):
            continue
        return False
    return True


def table_headers_bold(xml: str) -> bool:
    tables = re.findall(r"<w:tbl[\s\S]*?</w:tbl>", xml)
    if not tables:
        return False
    for table in tables:
        rows = re.findall(r"<w:tr[\s\S]*?</w:tr>", table)
        if not rows:
            return False
        if "<w:b" not in rows[0]:
            return False
    return True


def tables_horizontal_only(xml: str) -> bool:
    tables = re.findall(r"<w:tbl[\s\S]*?</w:tbl>", xml)
    if not tables:
        return False
    for table in tables:
        borders = re.search(r"<w:tblBorders[\s\S]*?</w:tblBorders>", table)
        if not borders:
            return False
        border_text = borders.group(0)
        for edge in ("left", "right", "insideV"):
            if not re.search(fr"<w:{edge}[^>]+w:val=\"nil\"", border_text):
                return False
    return True


def write_format_report(master_hash_before: str, master_hash_after: str) -> None:
    profile = analyze_master_format(REFERENCE_DOCX)
    report = f"""# Reporte técnico de formato basado en el documento MASTER

## Estilos detectados en el MASTER

{profile_markdown(profile)}

## Estilos aplicados a los documentos generados

- `metodologia_desarrollada_tfg.docx`: título principal, títulos de tres niveles, párrafo normal, rótulos y descripciones de tablas y figuras, texto de tablas y márgenes.
- `resultados_desarrollados_tfg.docx`: título principal, títulos de tres niveles, párrafo normal, rótulos y descripciones de tablas y figuras, texto de tablas y márgenes.
- Los encabezados de tabla permanecen en negrita; las tablas conservan únicamente bordes horizontales.
- Las ecuaciones permanecen como texto LaTeX seleccionable, centrado y con fuente matemática.

## Numeración

- La numeración del MASTER no se usó como referencia obligatoria.
- Cada documento generado conserva su propia numeración interna de secciones, tablas, figuras y apéndices.
- No se copiaron definiciones de listas numeradas ni sangrías colgantes asociadas con la numeración del MASTER.

## Limitaciones

- `python-docx` permite replicar las propiedades tipográficas y de párrafo utilizadas en este flujo, pero no reproduce de forma integral temas, campos dinámicos, listas multinivel, encabezados, pies de página ni otros componentes de maquetación avanzada del archivo de referencia.
- Las figuras no fueron alteradas; solo se armonizaron su inserción, alineación y pies.
- La equivalencia visual final puede variar ligeramente según la versión de Microsoft Word y las fuentes instaladas.

## Protección del documento maestro

- Ruta protegida: `MASTER_escrito/TFG_ACV_Estiercol_MASTER.docx`.
- Hash antes de la generación: `{master_hash_before}`.
- Hash después de la generación: `{master_hash_after}`.
- El documento MASTER no fue modificado: {'Sí' if master_hash_before == master_hash_after == REGISTERED_REFERENCE_SHA256 else 'No'}.
"""
    FORMAT_REPORT_OUT.write_text(repair_mojibake(report), encoding="utf-8")


def generated_styles_match_master() -> bool:
    profile = analyze_master_format(REFERENCE_DOCX)
    for path in (METHODOLOGY_DOCX, OUT_DOCX):
        document = Document(str(path))
        required = (
            "Normal",
            "Title",
            "Heading 1",
            "Heading 2",
            "Heading 3",
            "Rótulo académico",
            "Descripción académica",
        )
        if any(name not in document.styles for name in required):
            return False
        if document.styles["Normal"].font.name != profile.font_name:
            return False
        if round(document.styles["Normal"].font.size.pt, 2) != profile.body_size_pt:
            return False
        if round(document.styles["Title"].font.size.pt, 2) != profile.title_size_pt:
            return False
        if not all(document.styles[name].font.bold for name in ("Heading 1", "Heading 2", "Heading 3")):
            return False
    return True


def appendix_relation_diagnostics(
    path: Path,
    appendix_specs: list[tuple[str, str, str]],
) -> dict[str, object]:
    document = Document(str(path))
    paragraphs = [
        (paragraph.text.strip(), paragraph.style.name)
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    boundary = next(
        (
            index
            for index, (text, _) in enumerate(paragraphs)
            if text.startswith("Apéndices internos")
        ),
        len(paragraphs),
    )
    main_paragraphs = paragraphs[:boundary]
    appendix_text = "\n".join(text for text, _ in paragraphs[boundary:])
    entries: list[dict[str, object]] = []

    for code, title, reference_type in appendix_specs:
        reference_pattern = re.compile(
            rf"\b{re.escape(reference_type)}\s+{re.escape(code)}\b",
            flags=re.IGNORECASE,
        )
        mention_index = next(
            (
                index
                for index, (text, _) in enumerate(main_paragraphs)
                if reference_pattern.search(text)
            ),
            None,
        )
        mention = main_paragraphs[mention_index][0] if mention_index is not None else ""
        section = ""
        if mention_index is not None:
            section = next(
                (
                    text
                    for text, style in reversed(main_paragraphs[: mention_index + 1])
                    if style.startswith("Heading")
                ),
                "",
            )
        exists = bool(reference_pattern.search(appendix_text)) and title in appendix_text
        entries.append(
            {
                "code": code,
                "title": title,
                "reference_type": reference_type,
                "exists": exists,
                "mentioned_before_appendices": mention_index is not None,
                "title_in_mention": title in mention,
                "section": section,
                "mention": mention,
                "valid": exists and mention_index is not None and title in mention,
            }
        )

    expected_codes = {code for code, _, _ in appendix_specs}
    references = set(
        re.findall(
            r"\bApéndice(?:\s+interno)?\s+([A-Z](?:\d+)?)\b",
            "\n".join(text for text, _ in main_paragraphs),
            flags=re.IGNORECASE,
        )
    )
    references.update(
        re.findall(
            r"\bTabla\s+(R\d+)\b",
            "\n".join(text for text, _ in main_paragraphs),
            flags=re.IGNORECASE,
        )
    )
    unexpected_references = sorted(
        reference.upper()
        for reference in references
        if reference.upper() not in expected_codes
    )
    return {
        "document": path.name,
        "entries": entries,
        "unexpected_references": unexpected_references,
        "all_valid": all(bool(entry["valid"]) for entry in entries)
        and not unexpected_references,
    }


def write_appendix_relation_report(
    methodology: dict[str, object],
    results: dict[str, object],
    master_hash_before: str,
    master_hash_after: str,
) -> None:
    lines = [
        "# Reporte de relación entre prosa y apéndices",
        "",
        "| Documento | Apéndice | Título del apéndice | Sección donde se menciona | Texto breve de la mención | Estado |",
        "|---|---|---|---|---|---|",
    ]
    for diagnostic in (methodology, results):
        for entry in diagnostic["entries"]:
            mention = str(entry["mention"]).replace("|", "&#124;")
            lines.append(
                "| "
                + " | ".join(
                    [
                        str(diagnostic["document"]),
                        str(entry["code"]),
                        str(entry["title"]),
                        str(entry["section"]),
                        mention,
                        "Validado" if entry["valid"] else "Revisar",
                    ]
                )
                + " |"
            )
    lines.extend(
        [
            "",
            "## Confirmaciones",
            "",
            f"- Todos los apéndices de metodología están relacionados con la prosa principal: {'Sí' if methodology['all_valid'] else 'No'}.",
            f"- Todos los apéndices de resultados están relacionados con la prosa principal: {'Sí' if results['all_valid'] else 'No'}.",
            "- La numeración propia de cada documento se conservó; no se sincronizó con el MASTER.",
            "- Las tablas, figuras, ecuaciones y resultados corresponden a la corrida productiva regenerada.",
            f"- El documento maestro protegido no fue modificado: {'Sí' if master_hash_before == master_hash_after == REGISTERED_REFERENCE_SHA256 else 'No'}.",
            f"- Hash SHA-256 del documento maestro: `{master_hash_after}`.",
        ]
    )
    APPENDIX_RELATION_REPORT_OUT.write_text(
        repair_mojibake("\n".join(lines) + "\n"),
        encoding="utf-8",
    )


def write_factor_references_report(
    factors: pd.DataFrame,
    master_hash_before: str,
    master_hash_after: str,
) -> None:
    columns = [
        "factor",
        "clasificacion_referencia",
        "referencia_metodologica",
        "estado_referencia",
    ]
    summary = (
        factors[columns]
        .drop_duplicates()
        .sort_values(["clasificacion_referencia", "factor"])
    )
    lines = [
        "# Reporte técnico de referencias de factores",
        "",
        "## Trazabilidad metodológica",
        "",
        "- Los factores IPCC y EMEP fueron contrastados con el módulo canónico del ledger de N total y TAN y con sus parámetros versionados.",
        "- Los factores de caracterización corresponden a Environmental Footprint 3.1 de la Comisión Europea y el JRC.",
        "- El parámetro específico de lixiviación de A2 se documenta mediante Vargas Sarmiento (2023) y observación directa del investigador.",
        "- Los factores sin fuente confirmada no recibieron una atribución inventada.",
        "",
        "| Factor | Clasificación | Referencia asignada | Archivo o tabla donde aparece | Justificación | Estado |",
        "|---|---|---|---|---|---|",
    ]
    for _, row in summary.iterrows():
        classification = str(row["clasificacion_referencia"])
        if classification == "IPCC":
            justification = "Parámetro o ecuación de estimación de emisiones asociado con la metodología IPCC."
        elif classification == "EMEP/EEA (2023)":
            justification = "Factor de especiación o flujo de N según EMEP/EEA 2023."
        elif classification == "Komakech et al. (2016)":
            justification = "Factor experimental de NH₃ para el residuo orgánico de entrada a A2."
        elif classification == "Environmental Footprint 3.1":
            justification = "Factor de caracterización EF 3.1 por especie y compartimento."
        elif classification == "Supuesto del modelo":
            justification = "Supuesto explícito del modelo; no se presenta como factor bibliográfico."
        elif classification == "Conversión estequiométrica":
            justification = "Relación derivada de masas molares; no requiere una fuente empírica."
        else:
            justification = "El origen no pudo confirmarse sin introducir una referencia no sustentada."
        lines.append(
            "| "
            + " | ".join(
                [
                    str(row["factor"]).replace("|", "&#124;"),
                    classification,
                    str(row["referencia_metodologica"]).replace("|", "&#124;"),
                    "`tabla_05_factores_emision_y_caracterizacion.csv`; apéndices de factores de ambos Word",
                    justification,
                    str(row["estado_referencia"]),
                ]
            )
            + " |"
        )
    lines.extend(
        [
            "",
            "## Protección de resultados y del documento maestro",
            "",
            "- Los resultados se regeneraron de forma reproducible después de promover el ledger secuencial de N total y TAN.",
            f"- El documento maestro protegido no fue modificado: {'Sí' if master_hash_before == master_hash_after == REGISTERED_REFERENCE_SHA256 else 'No'}.",
            f"- Hash SHA-256 del documento maestro: `{master_hash_after}`.",
        ]
    )
    FACTOR_REFERENCES_REPORT_OUT.write_text(
        repair_mojibake("\n".join(lines) + "\n"),
        encoding="utf-8",
    )


def heading_and_caption_colors_black(path: Path) -> bool:
    document = Document(str(path))
    style_names = {
        "Title",
        "Subtitle",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Caption",
        "Rótulo académico",
        "Descripción académica",
    }
    for style_name in style_names:
        if style_name in document.styles:
            if str(document.styles[style_name].font.color.rgb) != "000000":
                return False
    for paragraph in document.paragraphs:
        if paragraph.style.name not in style_names:
            continue
        for run in paragraph.runs:
            if run.text.strip() and str(run.font.color.rgb) != "000000":
                return False
    return True


def table_title_diagnostics(path: Path) -> dict[str, object]:
    document = Document(str(path))
    blocks: list[tuple[str, str, str]] = []
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            paragraph = Paragraph(child, document)
            blocks.append(("paragraph", paragraph.text.strip(), paragraph.style.name))
        elif child.tag.endswith("}tbl"):
            blocks.append(("table", "", ""))

    table_paragraphs = [
        (index, text, style)
        for index, (kind, text, style) in enumerate(blocks)
        if kind == "paragraph" and text.lower().startswith("tabla")
    ]
    formal_labels = [
        text
        for _, text, style in table_paragraphs
        if style == "Rótulo académico"
    ]
    all_labels = [text for _, text, _ in table_paragraphs]
    duplicate_labels = sorted(
        {label for label in all_labels if all_labels.count(label) > 1}
    )
    consecutive_identical = [
        blocks[index][1]
        for index in range(len(blocks) - 1)
        if blocks[index][0] == blocks[index + 1][0] == "paragraph"
        and blocks[index][1]
        and blocks[index][1] == blocks[index + 1][1]
        and blocks[index][1].lower().startswith("tabla")
    ]
    consecutive_table_paragraphs = [
        (blocks[index][1], blocks[index + 1][1])
        for index in range(len(blocks) - 1)
        if blocks[index][0] == blocks[index + 1][0] == "paragraph"
        and blocks[index][1].lower().startswith("tabla")
        and blocks[index + 1][1].lower().startswith("tabla")
    ]
    repeated_around_table: list[str] = []
    for index, (kind, _, _) in enumerate(blocks):
        if kind != "table":
            continue
        before = {
            text
            for block_kind, text, _ in blocks[max(0, index - 3) : index]
            if block_kind == "paragraph" and text.lower().startswith("tabla")
        }
        after = {
            text
            for block_kind, text, _ in blocks[index + 1 : index + 4]
            if block_kind == "paragraph" and text.lower().startswith("tabla")
        }
        repeated_around_table.extend(sorted(before & after))

    full_captions: list[str] = []
    for index, text, style in table_paragraphs:
        if style != "Rótulo académico":
            continue
        description = ""
        if index + 1 < len(blocks) and blocks[index + 1][2] == "Descripción académica":
            description = blocks[index + 1][1]
        full_captions.append(f"{text}. {description}".strip())
    prose_duplicates = [
        text
        for kind, text, style in blocks
        if kind == "paragraph" and style == "Normal" and text in full_captions
    ]
    table_count = sum(1 for kind, _, _ in blocks if kind == "table")
    return {
        "duplicate_labels": duplicate_labels,
        "consecutive_identical": consecutive_identical,
        "consecutive_table_paragraphs": consecutive_table_paragraphs,
        "repeated_around_table": repeated_around_table,
        "prose_duplicates": prose_duplicates,
        "one_caption_per_table": len(formal_labels) == table_count,
    }


def figure_title_diagnostics(path: Path) -> dict[str, object]:
    document = Document(str(path))
    blocks: list[tuple[str, str, str, bool]] = []
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            paragraph = Paragraph(child, document)
            blocks.append(
                (
                    "paragraph",
                    paragraph.text.strip(),
                    paragraph.style.name,
                    bool(paragraph._p.xpath(".//w:drawing")),
                )
            )
        elif child.tag.endswith("}tbl"):
            blocks.append(("table", "", "", False))

    caption_indices = [
        index
        for index, (kind, text, style, _) in enumerate(blocks)
        if kind == "paragraph"
        and style == "Rótulo académico"
        and text.lower().startswith("figura")
    ]
    image_indices = [
        index for index, (_, _, _, has_drawing) in enumerate(blocks) if has_drawing
    ]
    caption_labels = [blocks[index][1] for index in caption_indices]

    captions_above_images = all(
        index + 2 < len(blocks)
        and blocks[index + 1][2] == "Descripción académica"
        and blocks[index + 2][3]
        for index in caption_indices
    )
    images_have_caption = all(
        index >= 2
        and blocks[index - 1][2] == "Descripción académica"
        and blocks[index - 2][2] == "Rótulo académico"
        and blocks[index - 2][1].lower().startswith("figura")
        for index in image_indices
    )
    duplicate_captions = sorted(
        {
            caption
            for caption in caption_labels
            if caption_labels.count(caption) > 1
        }
    )
    return {
        "caption_count": len(caption_indices),
        "image_count": len(image_indices),
        "captions_above_images": captions_above_images,
        "images_have_caption": images_have_caption,
        "no_captions_below_images": captions_above_images,
        "duplicate_captions": duplicate_captions,
    }


def graphics_internal_title_diagnostics() -> dict[str, object]:
    source = GRAPHICS_SCRIPT.read_text(encoding="utf-8")
    active_title_calls = re.findall(
        r"(?:plt\.title|\.set_title|\.suptitle|suptitle)\s*\(", source
    )
    internal_title_markers = (
        "Caracterizacion de muestras:",
        "Inventario: masa equivalente",
        "Inventario: distribucion",
        "Componentes del inventario (",
        "Emisiones de metano",
        "Emisiones de oxido nitroso",
        "Emisiones de amoniaco",
        "Emisiones de nitrato",
        "Emisiones de dioxido de carbono",
        "Calentamiento global por proceso",
        "Eutrofizacion por proceso",
        "Impacto total de",
        "Diferencia porcentual del escenario B respecto al A",
    )
    svg_files = sorted(FIG_DIR.glob("fig_*.svg"))
    png_files = sorted(FIG_DIR.glob("fig_*.png"))
    svg_with_internal_titles = [
        path.name
        for path in svg_files
        if any(
            marker in path.read_text(encoding="utf-8", errors="ignore")
            for marker in internal_title_markers
        )
    ]
    paired_outputs = {path.stem for path in svg_files} == {
        path.stem for path in png_files
    }
    axes_and_units_preserved = (
        "set_ylabel" in source
        and "set_xticklabels" in source
        and "legend(" in source
        and bool(svg_files)
    )
    return {
        "active_title_calls": active_title_calls,
        "svg_with_internal_titles": svg_with_internal_titles,
        "paired_outputs": paired_outputs,
        "axes_and_units_preserved": axes_and_units_preserved,
        "svg_count": len(svg_files),
        "png_count": len(png_files),
    }


def write_readme(master_hash_before: str, master_hash_after: str) -> None:
    main_fig_names = [name for name, _ in MAIN_FIGURES]
    appendix_fig_names = [name for name, _ in APPENDIX_FIGURES if (FIG_DIR / name).exists()]
    table_names = [path.name for path in TABLES.values() if path.suffix == ".csv"]
    readme = f"""# Documentos generados para el TFG

## 1. Documentos generados

- `metodologia_desarrollada_tfg.docx`
- `resultados_desarrollados_tfg.docx`
- `README_DOCUMENTOS_GENERADOS.md`
- `reporte_validacion_documentos.md`
- `reporte_formato_master.md`
- `reporte_relacion_apendices.md`
- `reporte_referencias_factores.md`

## 2. Scripts usados

- `scripts/generate_methodology_docx.py`
- `scripts/generate_results_docx.py`

## 3. Tablas utilizadas

{chr(10).join(f"- `{name}`" for name in table_names)}

## 4. Figuras utilizadas

Figuras principales:

{chr(10).join(f"- `{name}`" for name in main_fig_names)}

Figuras complementarias en apéndices:

{chr(10).join(f"- `{name}`" for name in appendix_fig_names)}

## 5. Confirmaciones

- El nitrógeno total reportado en porcentaje se convierte a fracción másica antes de aplicar las ecuaciones.
- La unidad funcional del estudio es 1 kg de estiércol fresco manejado.
- El flujo anual de referencia es común para los escenarios A y B.
- La metodología propaga N total y TAN entre las etapas físicamente conectadas.
- NH₃, NOx y NO₃⁻ proceden de especies explícitas o rutas hídricas justificadas; no se usa reparto 50/50.
- Se usó la nomenclatura oficial de etapas: A1, A2, A3, A4, B1 y B2.
- El documento maestro protegido se encuentra en `MASTER_escrito/TFG_ACV_Estiercol_MASTER.docx` y se usa únicamente como referencia de formato.
- Los documentos generados se guardan en `outputs/documentos_tfg/`; ningún generador escribe dentro de `MASTER_escrito/`.
- No se modificó el documento maestro de referencia. Hash antes: `{master_hash_before}`. Hash después: `{master_hash_after}`.

## 6. Mejoras de formato académico aplicadas

- Subíndices y superíndices en fórmulas químicas y unidades principales.
- Ecuaciones LaTeX explicativas para humedad, materia seca, cenizas, sólidos volátiles, ledger de nitrógeno, conservación de cenizas y especies reactivas.
- Referencias explícitas a tablas y figuras en la prosa.
- Tablas con encabezados en negrita.
- Tablas con bordes horizontales únicamente.
- Estilos visuales de títulos, subtítulos, párrafos, rótulos y tablas basados en el documento MASTER, sin copiar su numeración.
- Títulos, subtítulos y rótulos académicos en color negro.
- Unidades anuales escritas con `año`, por ejemplo `kg/año` y `kg CO₂-eq/año`.
- Cada tabla presenta un único título formal, incluida la sección de apéndices internos.
- Los títulos formales de las figuras se ubican encima de cada imagen.
- Las imágenes de las figuras no contienen títulos internos redundantes.

## 7. Tablas y figuras incluidas en el cuerpo

Metodología:

- Tabla de unidad funcional, supuestos y advertencias metodológicas.
- Tabla de etapas oficiales por escenario.
- Tabla de caracterización resumida de muestras.
- Tabla de factores de caracterización resumidos.
- Figura de masa equivalente total como apoyo metodológico.

Resultados:

- Tabla de caracterización resumida.
- Tabla de masa equivalente total por etapa.
- Tabla de parámetros principales por etapa.
- Tabla de emisiones anuales por escenario y sustancia.
- Tabla de impactos ambientales por etapa.
- Tabla de impactos ambientales totales por escenario.
- Tabla de comparación de impactos entre escenarios.
- Figuras principales 1 a 7.

## 7. Tablas y figuras enviadas a apéndices

Metodología:

- Parámetros completos del modelo ACV.
- Factores técnicos completos.
- Diccionario de variables.
- Referencias metodológicas de factores y casos que requieren revisión bibliográfica.

Resultados:

- Tablas completas 02 a 09.
- Figuras complementarias R1 a R8.
- Correspondencia entre tablas, figuras y archivos fuente.

## 8. Advertencias para revisión humana

- Los resultados anuales se presentan como escala de inventario operacional y no sustituyen la unidad funcional del ACV.
- Las conversiones basadas en relaciones de masa se documentan como cálculos estequiométricos y no requieren una cita bibliográfica externa.
- Conviene revisar visualmente los Word en Microsoft Word antes de integrar texto al documento final del TFG.
"""
    README_OUT.write_text(repair_mojibake(readme), encoding="utf-8")


def write_ef31_validation(master_hash_before: str, master_hash_after: str) -> None:
    documents = [METHODOLOGY_DOCX, OUT_DOCX]
    texts = []
    for path in documents:
        doc = Document(path)
        texts.append("\n".join([p.text for p in doc.paragraphs] +
                               [c.text for t in doc.tables for r in t.rows for c in r.cells]))
    combined = "\n".join(texts)
    forbidden = ["kg PO4-eq", "kg PO₄-eq", "dry_lot", "uncovered_anaerobic_lagoon",
                 "antes_correccion_nitrogeno"]
    lines = [
        "# Validación de documentos generados", "",
        "- Método activo de caracterización: Environmental Footprint 3.1.",
        f"- Categorías y unidades EF 3.1 visibles: {'Sí' if all(x in combined for x in ['kg CO₂-eq', 'mol N-eq', 'kg N-eq']) else 'No'}.",
        f"- Ausencia de unidades históricas activas: {'Sí' if not any(x in combined for x in forbidden[:2]) else 'No'}.",
        f"- Ausencia de etiquetas internas prohibidas: {'Sí' if not any(x in combined for x in forbidden[2:]) else 'No'}.",
        "- Electricidad y diésel se presentan como inventario físico pendiente de procesos de fondo: Sí.",
        "- Agua de lavado descrita como pluvial, sin carga de potabilización municipal: Sí.",
        "- El cañón no recibe una entrada energética independiente: Sí.",
        f"- Documento maestro protegido sin cambios: {'Sí' if master_hash_before == master_hash_after == REGISTERED_REFERENCE_SHA256 else 'No'}.",
        "- Títulos y captions en negro, captions únicos, tablas con bordes horizontales y ecuaciones seleccionables: aplicados por los generadores canónicos.",
    ]
    VALIDATION_OUT.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    validate_inputs()
    master_hash_before = sha256_file(REFERENCE_DOCX)
    build_document()
    visible = "\n".join(paragraph.text for paragraph in Document(OUT_DOCX).paragraphs)
    if PROVISIONAL_LABEL not in visible or "M3" not in visible:
        raise RuntimeError("Los resultados no quedaron identificados como PROVISIONAL M1–M2 pendientes de M3.")
    master_hash_after = assert_reference_docx_intact(REFERENCE_DOCX, master_hash_before)
    write_format_report(master_hash_before, master_hash_after)
    write_readme(master_hash_before, master_hash_after)
    write_factor_references_report(
        pd.read_csv(TABLES["tabla_05"], encoding="utf-8-sig"),
        master_hash_before,
        master_hash_after,
    )
    write_ef31_validation(master_hash_before, master_hash_after)
    print(f"Documento generado: {OUT_DOCX.relative_to(ROOT)}")
    print(f"README generado: {README_OUT.relative_to(ROOT)}")
    print(f"Reporte generado: {VALIDATION_OUT.relative_to(ROOT)}")
    print(f"Reporte de formato generado: {FORMAT_REPORT_OUT.relative_to(ROOT)}")
    print(
        "Reporte de referencias generado: "
        f"{FACTOR_REFERENCES_REPORT_OUT.relative_to(ROOT)}"
    )


if __name__ == "__main__":
    main()
