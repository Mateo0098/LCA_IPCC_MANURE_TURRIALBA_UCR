from __future__ import annotations

import csv
import math
import re
import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))

from master_word_format import apply_master_format, finalize_document_format  # noqa: E402
from reference_docx_utils import (  # noqa: E402
    assert_reference_docx_intact,
    get_reference_docx_path,
    sha256_file,
)
from quantitative_comparison import Comparison, dominant  # noqa: E402


MASTER = ROOT / "MASTER_escrito" / "TFG_ACV_Estiercol_MASTER.docx"
TOTALS = ROOT / "outputs" / "tablas_tesis" / "tabla_08_impactos_totales_por_escenario.csv"
COMPARISON = ROOT / "outputs" / "tablas_tesis" / "tabla_09_comparacion_escenarios.csv"
STAGES = ROOT / "outputs" / "tablas_tesis" / "tabla_07_impactos_por_etapa.csv"
PROCESSED_TOTALS = ROOT / "processed" / "acv_impacto_total_por_escenario.csv"
PROCESSED_MASS = ROOT / "processed" / "masa_total_escenario_etapa.csv"
OUT_DIR = ROOT / "outputs" / "documentos_tfg"
OUT_DOCX = OUT_DIR / "conclusiones_desarrolladas_tfg.docx"
OUT_TRACE = OUT_DIR / "trazabilidad_conclusiones_tfg.md"

GENERAL_PREFIX = "Desarrollar un Análisis de Ciclo de Vida"
OE1_PREFIX = "Realizar el inventario para el Análisis de Ciclo de Vida"
OE2_PREFIX = "Evaluar el impacto ambiental del estiércol bovino"
EXPECTED_SCENARIOS = {"A", "B"}
EXPECTED_CATEGORIES = {"Cambio climático", "Eutrofización terrestre", "Eutrofización marina"}
PROVISIONAL_LABEL = "PROVISIONAL M1–M2"

OFFICIAL_STAGES = {
    ("A", 1): "A1: Precomposteo",
    ("A", 2): "A2: Lombricompostaje",
    ("A", 3): "A3: Almacenamiento de aguas verdes",
    ("A", 4): "A4: Aplicación de aguas verdes en campos de pastoreo",
    ("B", 1): "B1: Almacenamiento de purines",
    ("B", 2): "B2: Aplicación de purines en campo de pastoreo",
}


def read_rows(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as source:
        return list(csv.DictReader(source))


def extract_objectives() -> dict[str, str]:
    paragraphs = [paragraph.text.strip() for paragraph in Document(MASTER).paragraphs]

    def find(prefix: str) -> str:
        matches = [text for text in paragraphs if text.startswith(prefix)]
        if len(matches) != 1:
            raise RuntimeError(f"No se encontró un objetivo único con el inicio: {prefix}")
        return matches[0]

    return {"OG": find(GENERAL_PREFIX), "OE1": find(OE1_PREFIX), "OE2": find(OE2_PREFIX)}


def finite_number(value: str, label: str, *, nonnegative: bool = False) -> float:
    try:
        number_value = float(value)
    except (TypeError, ValueError) as exc:
        raise RuntimeError(f"{label} no es numérico: {value!r}") from exc
    if not math.isfinite(number_value):
        raise RuntimeError(f"{label} debe ser finito: {number_value}")
    if nonnegative and number_value < 0:
        raise RuntimeError(f"{label} debe ser no negativo: {number_value}")
    return number_value


def impact_totals() -> dict[tuple[str, str], float]:
    values = {}
    for row in read_rows(TOTALS):
        scenario = row["escenario"].strip().upper()
        category = row["categoria_impacto"].strip()
        key = (scenario, category)
        if key in values:
            raise RuntimeError(f"El impacto total está duplicado para {key}.")
        values[key] = finite_number(row["resultado_total"], f"Impacto total {key}", nonnegative=True)
    expected_keys = {
        (scenario, category)
        for scenario in EXPECTED_SCENARIOS
        for category in EXPECTED_CATEGORIES
    }
    if values.keys() != expected_keys:
        raise RuntimeError("Las categorías o escenarios de impactos totales no son los esperados.")
    return values


def comparisons(totals: dict[tuple[str, str], float]) -> dict[str, float]:
    values = {}
    for row in read_rows(COMPARISON):
        category = row["categoria_impacto"]
        if category not in EXPECTED_CATEGORIES:
            raise RuntimeError(f"Categoría comparativa inesperada: {category}")
        if category in values:
            raise RuntimeError(f"Categoría comparativa duplicada: {category}")
        current = finite_number(
            row["diferencia_porcentual_B_vs_A"],
            f"Diferencia porcentual de {category}",
        )
        if totals[("A", category)] == 0:
            raise RuntimeError(f"No puede calcularse la diferencia relativa de {category} con A igual a cero.")
        calculated = 100 * (
            totals[("B", category)] - totals[("A", category)]
        ) / totals[("A", category)]
        if not math.isclose(current, calculated, rel_tol=1e-12):
            raise RuntimeError(f"La diferencia porcentual no cierra para {category}.")
        values[category] = current
    if values.keys() != EXPECTED_CATEGORIES:
        raise RuntimeError("La comparación no contiene exactamente las categorías esperadas.")
    return values


def processed_indicators(
    totals: dict[tuple[str, str], float],
) -> tuple[float, dict[tuple[str, str], float]]:
    rows = read_rows(PROCESSED_TOTALS)
    if {row["Escenario"].strip().upper() for row in rows} != EXPECTED_SCENARIOS:
        raise RuntimeError("La salida procesada debe contener exactamente los escenarios A y B.")
    references: dict[str, float] = {}
    normalized_values: dict[tuple[str, str], float] = {}
    columns = {
        "Cambio climático": (
            "impacto_calentamiento_global_kg_co2eq",
            "impacto_calentamiento_global_kg_co2eq_por_kg_estiercol_fresco",
        ),
        "Eutrofización terrestre": (
            "impacto_eutrofizacion_terrestre_mol_neq",
            "impacto_eutrofizacion_terrestre_mol_neq_por_kg_estiercol_fresco",
        ),
        "Eutrofización marina": (
            "impacto_eutrofizacion_marina_kg_neq",
            "impacto_eutrofizacion_marina_kg_neq_por_kg_estiercol_fresco",
        ),
    }
    for row in rows:
        scenario = row["Escenario"].strip().upper()
        reference = finite_number(
            row["referencia_funcional_estiercol_fresco_kg"],
            f"Flujo de referencia de {scenario}",
            nonnegative=True,
        )
        if reference == 0:
            raise RuntimeError(f"El flujo de referencia de {scenario} debe ser positivo.")
        references[scenario] = reference
        for category, (annual_column, normalized_column) in columns.items():
            annual = finite_number(row[annual_column], f"Impacto procesado de {scenario}, {category}", nonnegative=True)
            indicator = finite_number(row[normalized_column], f"Indicador normalizado de {scenario}, {category}", nonnegative=True)
            if not math.isclose(annual, totals[(scenario, category)], rel_tol=1e-12, abs_tol=1e-12):
                raise RuntimeError(f"El impacto procesado y la tabla de tesis no coinciden para {(scenario, category)}.")
            if not math.isclose(indicator * reference, annual, rel_tol=1e-12, abs_tol=1e-9):
                raise RuntimeError(f"La normalización no reconstruye el resultado anual para {(scenario, category)}.")
            normalized_values[(scenario, category)] = indicator
    if not math.isclose(references["A"], references["B"], rel_tol=1e-12, abs_tol=1e-9):
        raise RuntimeError("Los escenarios A y B no comparten el mismo flujo anual de referencia.")
    return references["A"], normalized_values


def stage_totals() -> dict[tuple[str, int, str], float]:
    values: dict[tuple[str, int, str], float] = {}
    for row in read_rows(STAGES):
        key = (row["escenario"], int(float(row["etapa"])), row["categoria_impacto"])
        if key[0] not in EXPECTED_SCENARIOS or key[2] not in EXPECTED_CATEGORIES:
            raise RuntimeError(f"Escenario o categoría inesperada en impactos por etapa: {key}")
        contribution = finite_number(
            row["resultado_equivalente"],
            f"Impacto por etapa {key}",
            nonnegative=True,
        )
        values[key] = values.get(key, 0.0) + contribution
    return values


def validate_stage_sums(
    stages: dict[tuple[str, int, str], float],
    totals: dict[tuple[str, str], float],
) -> None:
    for scenario in EXPECTED_SCENARIOS:
        for category in EXPECTED_CATEGORIES:
            subtotal = sum(
                value
                for (row_scenario, _, row_category), value in stages.items()
                if row_scenario == scenario and row_category == category
            )
            if not math.isclose(subtotal, totals[(scenario, category)], rel_tol=1e-12, abs_tol=1e-9):
                raise RuntimeError(f"Los impactos por etapa no suman el total para {(scenario, category)}.")


def dominant_stage(
    scenario: str,
    category: str,
    stages: dict[tuple[str, int, str], float],
    totals: dict[tuple[str, str], float],
) -> tuple[str, float, float]:
    candidates = {OFFICIAL_STAGES[(key[0], key[1])]: value for key, value in stages.items() if key[0] == scenario and key[2] == category}
    stage_name, value = dominant(candidates, decimals=9)
    return stage_name, value, 100 * value / totals[(scenario, category)]


def flow_inventory(reference_flow: float) -> tuple[float, float]:
    rows = read_rows(PROCESSED_MASS)
    references_by_scenario: dict[str, set[float]] = {"A": set(), "B": set()}
    collected = remainder = None
    for row in rows:
        scenario = row["escenario"].strip().upper()
        if scenario not in references_by_scenario:
            continue
        references_by_scenario[scenario].add(
            finite_number(row["flujo_referencia_anual_kg"], f"Flujo de referencia de {scenario}", nonnegative=True)
        )
        if scenario == "A":
            collected = finite_number(row["fraccion_recolectada"], "Fracción recolectada", nonnegative=True)
            remainder = finite_number(row["fraccion_remanente"], "Fracción remanente", nonnegative=True)
    if any(len(values) != 1 for values in references_by_scenario.values()):
        raise RuntimeError("Cada escenario debe declarar un único flujo anual de referencia.")
    reference_a = next(iter(references_by_scenario["A"]))
    reference_b = next(iter(references_by_scenario["B"]))
    if not math.isclose(reference_a, reference_b, rel_tol=1e-12, abs_tol=1e-9):
        raise RuntimeError("Los flujos de referencia declarados para A y B son diferentes.")
    if not math.isclose(reference_a, reference_flow, rel_tol=1e-12, abs_tol=1e-9):
        raise RuntimeError("Las salidas procesadas de masa e impactos declaran referencias diferentes.")
    if collected is None or remainder is None or not math.isclose(collected + remainder, 1.0, rel_tol=1e-12, abs_tol=1e-12):
        raise RuntimeError("Las fracciones recolectada y remanente del Escenario A no suman uno.")
    return collected, remainder


def number(value: float, decimals: int) -> str:
    text = f"{value:,.{decimals}f}"
    return text.replace(",", "X").replace(".", ",").replace("X", " ")


def build_conclusions(
    totals: dict[tuple[str, str], float],
    normalized_values: dict[tuple[str, str], float],
    percentage: dict[str, float],
    stages: dict[tuple[str, int, str], float],
    reference_flow: float,
    collected_fraction: float,
    remainder_fraction: float,
) -> list[dict[str, str]]:
    a_cg = dominant_stage("A", "Cambio climático", stages, totals)
    b_cg = dominant_stage("B", "Cambio climático", stages, totals)
    a_eu = dominant_stage("A", "Eutrofización marina", stages, totals)
    b_eu = dominant_stage("B", "Eutrofización marina", stages, totals)
    comparisons_by_category = {}
    units = {
        "Cambio climático": "kg CO₂-eq/kg de estiércol fresco manejado",
        "Eutrofización terrestre": "mol N-eq/kg de estiércol fresco manejado",
        "Eutrofización marina": "kg N-eq/kg de estiércol fresco manejado",
    }
    decimals = {"Cambio climático": 3, "Eutrofización terrestre": 6, "Eutrofización marina": 6}
    for category in EXPECTED_CATEGORIES:
        item = Comparison(
            "Escenario A", normalized_values[("A", category)],
            "Escenario B", normalized_values[("B", category)], units[category],
        )
        annual_item = Comparison(
            "Escenario A", totals[("A", category)],
            "Escenario B", totals[("B", category)], units[category].split("/kg")[0] + "/año",
        )
        annual_item.assert_consistent(
            difference=annual_item.difference_right_minus_left,
            percentage=percentage[category],
        )
        if item.higher_label != annual_item.higher_label:
            raise RuntimeError(f"La normalización invierte la comparación de {category} entre escenarios.")
        item.assert_rounding_unambiguous(decimals[category])
        comparisons_by_category[category] = item
    cg_comparison = comparisons_by_category["Cambio climático"]
    terrestrial_comparison = comparisons_by_category["Eutrofización terrestre"]
    eu_comparison = comparisons_by_category["Eutrofización marina"]
    return [
        {
            "id": "C1",
            "objective": "OE1",
            "text": (
                "El inventario de ciclo de vida representó y cuantificó las rutas diferenciadas de manejo del estiércol en "
                f"los dos escenarios, a partir de un flujo anual común de {number(reference_flow, 6)} kg de estiércol fresco. "
                f"En el Escenario A, {number(100 * collected_fraction, 2)} % ingresó a la ruta sólida mediante A1: Precomposteo; "
                "posteriormente, la masa resultante continuó hacia A2: Lombricompostaje, mientras que "
                f"{number(100 * remainder_fraction, 2)} % se incorporó a las aguas verdes. En el Escenario B, el "
                "100 % ingresó al sistema de purines. De este modo, el inventario caracterizó las trayectorias de manejo y "
                "los flujos asociados a cada alternativa estudiada. El inventario incorporó además 53,23 kWh/año de electricidad en A3 o B1 y 182,50 L/año de diésel en A4 o B2; sus procesos de fondo permanecen pendientes."
            ),
            "evidence": (
                f"Flujo común: {reference_flow:.9f} kg/año; fracción recolectada de A: {collected_fraction:.12f}; "
                f"fracción remanente de A: {remainder_fraction:.12f}; B recibe el flujo completo."
            ),
            "source": "Tabla 3 de flujos del ICV; auditoría del balance del Escenario B y la unidad funcional.",
        },
        {
            "id": "C2",
            "objective": "OE2",
            "text": (
                "Las cargas ambientales se concentraron en etapas específicas, que variaron según el escenario y la categoría. "
                "En calentamiento global, "
                f"{a_cg[0]} aportó {number(a_cg[2], 2)} % del total del Escenario A y {b_cg[0]} aportó "
                f"{number(b_cg[2], 2)} % del total del Escenario B. Para eutrofización, las mayores contribuciones "
                f"correspondieron a {a_eu[0]} ({number(a_eu[2], 2)} % del total de A) y {b_eu[0]} "
                f"({number(b_eu[2], 2)} % del total de B), considerando el supuesto de representación del nitrógeno "
                "potencialmente eutrofizante adoptado en el estudio. La evaluación por etapa identificó así patrones de "
                "concentración diferenciados entre las alternativas estudiadas."
            ),
            "evidence": (
                f"CG: {a_cg[0]} = {a_cg[1]:.9f}; {b_cg[0]} = {b_cg[1]:.9f}. "
                f"Eutrofización: {a_eu[0]} = {a_eu[1]:.9f}; {b_eu[0]} = {b_eu[1]:.9f}."
            ),
            "source": "Tabla 7 de impactos por etapa y Tabla 8 de impactos totales.",
        },
        {
            "id": "C3",
            "objective": "OE2",
            "text": (
                f"Para calentamiento global, el {cg_comparison.higher_label} presentó el mayor impacto por unidad funcional: "
                f"{number(max(cg_comparison.left, cg_comparison.right), 3)} kg CO₂-eq/kg de estiércol fresco manejado, frente a "
                f"{number(min(cg_comparison.left, cg_comparison.right), 3)} kg CO₂-eq/kg de estiércol fresco manejado en el {cg_comparison.lower_label}. La diferencia "
                f"B menos A fue de {number(percentage['Cambio climático'], 2)} % respecto a A, lo que evidenció un menor indicador "
                f"para el {cg_comparison.lower_label} bajo las condiciones estudiadas."
            ),
            "evidence": (
                f"A: {normalized_values[('A', 'Cambio climático')]:.12f} kg CO₂-eq/kg y "
                f"{totals[('A', 'Cambio climático')]:.12f} kg CO₂-eq/año; B: "
                f"{normalized_values[('B', 'Cambio climático')]:.12f} kg CO₂-eq/kg y "
                f"{totals[('B', 'Cambio climático')]:.12f} kg CO₂-eq/año; B − A: "
                f"{percentage['Cambio climático']:.12f} %."
            ),
            "source": "Tablas 8 y 9 de impactos totales y comparación de escenarios.",
        },
        {
            "id": "C4",
            "objective": "OE2",
            "text": (
                f"Para eutrofización marina EF 3.1, el {eu_comparison.higher_label} presentó el mayor impacto por unidad funcional: "
                f"{number(max(eu_comparison.left, eu_comparison.right), 6)} kg N-eq/kg de estiércol fresco manejado, frente a "
                f"{number(min(eu_comparison.left, eu_comparison.right), 6)} kg N-eq/kg en el {eu_comparison.lower_label}. La diferencia "
                f"B menos A fue de {number(percentage['Eutrofización marina'], 2)} % respecto a A. Para eutrofización terrestre, "
                f"el {terrestrial_comparison.higher_label} también presentó el mayor impacto por unidad funcional: "
                f"{number(max(terrestrial_comparison.left, terrestrial_comparison.right), 6)} mol N-eq/kg de estiércol fresco manejado, "
                f"frente a {number(min(terrestrial_comparison.left, terrestrial_comparison.right), 6)} mol N-eq/kg en el "
                f"{terrestrial_comparison.lower_label}; la diferencia B menos A fue de "
                f"{number(percentage['Eutrofización terrestre'], 2)} % respecto a A."
            ),
            "evidence": (
                f"A: {normalized_values[('A', 'Eutrofización marina')]:.12f} kg N-eq/kg y "
                f"{totals[('A', 'Eutrofización marina')]:.12f} kg N-eq/año; B: "
                f"{normalized_values[('B', 'Eutrofización marina')]:.12f} kg N-eq/kg y "
                f"{totals[('B', 'Eutrofización marina')]:.12f} kg N-eq/año; B − A: "
                f"{percentage['Eutrofización marina']:.12f} %."
                f" Eutrofización terrestre A: {normalized_values[('A', 'Eutrofización terrestre')]:.12f} mol N-eq/kg; "
                f"B: {normalized_values[('B', 'Eutrofización terrestre')]:.12f} mol N-eq/kg; B − A: "
                f"{percentage['Eutrofización terrestre']:.12f} %."
            ),
            "source": "Tablas 8 y 9 de impactos totales y comparación de escenarios.",
        },
        {
            "id": "C5",
            "objective": "Objetivo general",
            "text": (
                "En conjunto, el ACV permitió estimar el desempeño ambiental de las dos alternativas de manejo en la lechería "
                "estudiada. Bajo la unidad funcional común de 1 kg de estiércol fresco manejado y las condiciones modeladas, "
                f"el {cg_comparison.lower_label} presentó el menor indicador de calentamiento global y el {eu_comparison.lower_label} "
                f"presentó el menor indicador de eutrofización marina; el {terrestrial_comparison.lower_label} presentó asimismo "
                "el menor indicador de eutrofización terrestre, expresado separadamente en mol N-eq. Environmental Footprint 3.1 caracterizó las emisiones directas. Este resultado se "
                "circunscribió al sistema evaluado y no implicó la superioridad universal de una alternativa en otras lecherías "
                "o condiciones operativas."
            ),
            "evidence": f"Menor calentamiento global: {cg_comparison.lower_label}; menor eutrofización marina: {eu_comparison.lower_label}; menor eutrofización terrestre: {terrestrial_comparison.lower_label}; referencia común de 1 kg de estiércol fresco manejado.",
            "source": "Tablas 7, 8 y 9; metodología desarrollada; objetivos extraídos del master.",
        },
    ]


def write_docx(conclusions: list[dict[str, str]]) -> None:
    document = Document()
    profile = apply_master_format(document, MASTER)
    document.add_paragraph("Conclusiones", style="Title")
    marker = document.add_paragraph()
    marker.alignment = 1
    run = marker.add_run(PROVISIONAL_LABEL)
    run.bold = True
    for section in document.sections:
        header = section.header.paragraphs[0]
        header.text = PROVISIONAL_LABEL
        header.alignment = 1
    document.add_paragraph(
        "Estas conclusiones se derivan de la corrida provisional M1–M2 vigente. La incorporación de M3 "
        "actualizará la caracterización final y podrá modificar los resultados cuantitativos.",
        style="Normal",
    )
    document.add_heading("Conclusiones del análisis de ciclo de vida", level=1)
    for conclusion in conclusions:
        document.add_paragraph(conclusion["text"], style="Normal")
    finalize_document_format(document, profile)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUT_DOCX)


def escape_markdown(value: str) -> str:
    return value.replace("|", "\\|").replace("\n", " ")


def write_trace(objectives: dict[str, str], conclusions: list[dict[str, str]], master_hash: str) -> None:
    lines = [
        "# Trazabilidad de las conclusiones del TFG",
        "",
        "## Estado documental",
        "",
        "Esta matriz corresponde a la corrida **PROVISIONAL M1–M2** vigente. Para sólidos integra M1 y M2; para N total de aguas verdes y purines utiliza exclusivamente M2 Kjeldahl. Todas las conclusiones permanecen provisionales hasta incorporar M3 y regenerar los documentos.",
        "",
        f"El documento maestro se utilizó únicamente en modo de lectura. SHA-256 verificado: `{master_hash}`.",
        "",
        "## Objetivos extraídos del documento maestro",
        "",
        f"- **Objetivo general:** {objectives['OG']}",
        f"- **OE1:** {objectives['OE1']}",
        f"- **OE2:** {objectives['OE2']}",
        "",
        "## Matriz de trazabilidad",
        "",
        "| Conclusión | Objetivo relacionado | Resultados que la sustentan | Tabla/Figura/Fuente | Estado |",
        "|---|---|---|---|---|",
    ]
    for conclusion in conclusions:
        lines.append(
            "| " + " | ".join(
                escape_markdown(value)
                for value in (
                    f"{conclusion['id']}. {conclusion['text']}",
                    conclusion["objective"],
                    conclusion["evidence"],
                    conclusion["source"],
                    "PROVISIONAL",
                )
            ) + " |"
        )
    lines.extend(
        [
            "",
            "## Criterio para fortalecer las conclusiones",
            "",
            "La clasificación podrá revisarse cuando se incorpore M3 y se regenere la cadena de resultados. Hasta entonces, la integración M1–M2 no se interpreta como evidencia definitiva de toda la campaña experimental.",
            "",
        ]
    )
    OUT_TRACE.write_text("\n".join(lines), encoding="utf-8")


def validate_outputs(conclusions: list[dict[str, str]]) -> None:
    document = Document(OUT_DOCX)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    if len([text for text in paragraphs if text in {item["text"] for item in conclusions}]) != len(conclusions):
        raise RuntimeError("No se encontraron todas las conclusiones una sola vez en el Word.")
    if document.tables or document.inline_shapes:
        raise RuntimeError("El documento de conclusiones no debe contener tablas ni figuras.")
    visible = "\n".join(paragraphs)
    if PROVISIONAL_LABEL not in visible or "M3" not in visible:
        raise RuntimeError("Las conclusiones no quedaron identificadas como PROVISIONAL M1–M2 pendientes de M3.")
    forbidden = ["processed", "outputs", "scripts", ".csv", "n_ex_pct", "n_ex_fraction", "dry_lot", "1,000", "2,000", "3,000", "4,000", "L/ano", "kg/ano", "AnÃ"]
    found = [term for term in forbidden if term in visible]
    if found:
        raise RuntimeError(f"Se detectó lenguaje no permitido en el Word: {found}")
    by_id = {item["id"]: item["text"] for item in conclusions}
    if "tal y como fue recolectado" in visible.lower() or re.search(
        r"unidad funcional.{0,80}(?:estiércol|estiercol).{0,25}recolectad",
        visible,
        flags=re.IGNORECASE | re.DOTALL,
    ):
        raise RuntimeError("La unidad funcional no debe limitarse al estiércol físicamente recolectado.")
    if "precomposteo y lombricompostaje" in by_id["C1"].lower():
        raise RuntimeError("C1 presenta simultáneamente el ingreso a precomposteo y lombricompostaje.")
    if not all(term in by_id["C1"] for term in ["A1: Precomposteo", "posteriormente", "A2: Lombricompostaje"]):
        raise RuntimeError("C1 no conserva la secuencia A1 seguida de A2.")
    unit_denominator = "kg de estiércol fresco manejado"
    if any(unit_denominator not in by_id[item] for item in ["C3", "C4"]):
        raise RuntimeError("C3 y C4 deben expresar el denominador completo de la unidad funcional.")
    if "kg PO₄-eq" in visible or "kg PO4-eq" in visible:
        raise RuntimeError("Las conclusiones conservan una unidad histórica de eutrofización.")
    if "kg N-eq" not in by_id["C4"] or "mol N-eq" not in by_id["C4"]:
        raise RuntimeError("C4 debe distinguir eutrofización marina y terrestre EF 3.1.")
    if re.search(r"B1.{0,160}lixiviación física|lixiviación física.{0,160}B1", visible, flags=re.IGNORECASE | re.DOTALL):
        raise RuntimeError("Las conclusiones no deben atribuir la eutrofización de B1 a lixiviación física.")
    clauses = re.split(r"[.;]", visible)
    clauses_without_b_stages = [
        re.sub(r"\bB[12]:.*$", "", clause, flags=re.IGNORECASE) for clause in clauses
    ]
    if any(
        re.search(r"\bA[1-4]:", clause)
        and re.search(r"\bpur[ií]n(?:es)?\b", clause, flags=re.IGNORECASE)
        for clause in clauses_without_b_stages
    ):
        raise RuntimeError("Se asoció purín con una etapa del Escenario A.")
    trace = OUT_TRACE.read_text(encoding="utf-8")
    if any(f"{item['id']}." not in trace for item in conclusions):
        raise RuntimeError("La matriz no incluye todas las conclusiones.")
    if trace.count("PROVISIONAL") < len(conclusions):
        raise RuntimeError("No todas las conclusiones quedaron marcadas como PROVISIONAL.")


def main() -> None:
    reference = get_reference_docx_path(ROOT)
    if reference != MASTER:
        raise RuntimeError("La ruta validada del documento maestro no coincide con la esperada.")
    before_hash = sha256_file(MASTER)
    objectives = extract_objectives()
    totals = impact_totals()
    percentage = comparisons(totals)
    reference_flow, normalized_values = processed_indicators(totals)
    collected_fraction, remainder_fraction = flow_inventory(reference_flow)
    stages = stage_totals()
    validate_stage_sums(stages, totals)
    conclusions = build_conclusions(
        totals,
        normalized_values,
        percentage,
        stages,
        reference_flow,
        collected_fraction,
        remainder_fraction,
    )
    write_docx(conclusions)
    after_hash = assert_reference_docx_intact(MASTER, before_hash)
    write_trace(objectives, conclusions, after_hash)
    validate_outputs(conclusions)
    print(f"Documento generado: {OUT_DOCX.relative_to(ROOT)}")
    print(f"Trazabilidad generada: {OUT_TRACE.relative_to(ROOT)}")
    print(f"Conclusiones sustantivas: {len(conclusions)}")
    print(f"Hash del master sin cambios: {after_hash}")


if __name__ == "__main__":
    main()
