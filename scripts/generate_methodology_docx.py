from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from academic_text_utils import clean_academic_label
from master_word_format import (
    add_master_caption,
    analyze_master_format,
    apply_master_format,
    finalize_document_format,
    format_table_like_master,
)
from reference_docx_utils import (
    assert_reference_docx_intact,
    get_reference_docx_path,
    sha256_file,
)


ROOT = Path(__file__).resolve().parents[1]
PROVISIONAL_LABEL = "PROVISIONAL M1–M2"
PROVISIONAL_NOTE = (
    "Esta versión incorpora las jornadas disponibles M1 y M2. Para los sólidos se utiliza la "
    "integración provisional M1–M2; para el N total de aguas verdes y purines se utiliza "
    "exclusivamente M2 mediante Kjeldahl. La caracterización final se actualizará al incorporar M3."
)
REFERENCE_DOCX = ROOT / "MASTER_escrito" / "TFG_ACV_Estiercol_MASTER.docx"
TABLE_DIR = ROOT / "outputs" / "tablas_tesis"
FIG_DIR = ROOT / "outputs" / "graficos_tesis"
OUT_DIR = ROOT / "outputs" / "documentos_tfg"
OUT_DOCX = OUT_DIR / "metodologia_desarrollada_tfg.docx"
OPERATIVE_PARAMS = ROOT / "Academic_documents" / "references" / "parametros_operativos_sanchez_2026.csv"
MASS_BALANCE = ROOT / "processed" / "masa_total_escenario_etapa.csv"

TABLES = {
    "unidad": TABLE_DIR / "tabla_00_unidad_funcional_y_supuestos.csv",
    "etapas": TABLE_DIR / "tabla_01_etapas_escenarios.csv",
    "caracterizacion": TABLE_DIR / "tabla_02_caracterizacion_muestras.csv",
    "flujos": TABLE_DIR / "tabla_03_flujos_icv.csv",
    "parametros": TABLE_DIR / "tabla_04_parametros_modelo_acv.csv",
    "factores": TABLE_DIR / "tabla_05_factores_emision_y_caracterizacion.csv",
    "diccionario": TABLE_DIR / "diccionario_variables.csv",
}

OFFICIAL_STAGE_NAMES = {
    ("A", 1): ("A1", "Precomposteo", "Etapa 1: Precomposteo"),
    ("A", 2): ("A2", "Lombricompostaje", "Etapa 2: Lombricompostaje"),
    ("A", 3): ("A3", "Almacenamiento de aguas verdes", "Etapa 3: Almacenamiento de aguas verdes"),
    ("A", 4): (
        "A4",
        "Aplicación de aguas verdes en campos de pastoreo",
        "Etapa 4: Aplicación de aguas verdes en campos de pastoreo",
    ),
    ("B", 1): ("B1", "Almacenamiento de purines", "Etapa 1: Almacenamiento de purines"),
    ("B", 2): (
        "B2",
        "Aplicación de purines en campo de pastoreo",
        "Etapa 2: Aplicación de purines en campo de pastoreo",
    ),
}

OLD_STAGE_TERMS = {
    "Manejo inicial de estiércol fresco": "Precomposteo",
    "Manejo posterior de fracción sólida": "Lombricompostaje",
    "Manejo de estiércol fresco sin precompostaje": "Almacenamiento de purines",
    "Manejo o aplicación de purines": "Aplicación de purines en campo de pastoreo",
    "Aplicación o manejo de aguas verdes en suelo": "Aplicación de aguas verdes en campos de pastoreo",
}

INTERNAL_COLUMNS = {
    "archivo_fuente",
    "source_file",
    "script",
    "script_origen",
    "path",
    "input_file",
    "output_file",
    "nombre_variable_codigo",
    "columna_original",
    "id interno",
    "observaciones internas del repositorio",
    "fuente",
    "fuente_dato",
    "fuente_o_calculo",
    "formula_origen",
    "observaciones",
    "fuente_bibliografica_pendiente",
    "requiere_revision_bibliografica",
    "clasificacion_referencia",
    "estado_referencia",
    "fuente_factor_emision",
    "fuente_factor_caracterizacion",
    "fuente_factor",
    "ecuacion_utilizada",
    "formula",
    "archivo fuente",
}

ACADEMIC_REPLACEMENTS = {
    "dry_lot": "Sistema de manejo en corral seco",
    "uncovered_anaerobic_lagoon": "Laguna anaerobia descubierta",
    "composting_invessel": "Compostaje en sistema cerrado",
    "solid_storage": "Almacenamiento sólido",
    "liquid_slurry": "Sistema líquido tipo purín",
    "aerobic_treatment": "Tratamiento aeróbico",
    "composting_intensive": "Compostaje intensivo",
    "composting_pasive": "Compostaje pasivo",
    "windrow_intensive": "Compostaje intensivo en hileras",
    "windrow_pasive": "Compostaje pasivo en hileras",
    "ipcc": "IPCC",
    "medido": "Factor medido",
    "ESTIERCOL FRESCO": "Estiércol fresco",
    "SOL: PRECOMPOSTADO": "Estiércol precompostado",
    "LIQ: AGUA VERDE": "Agua de lavado incorporada a las aguas verdes",
    "LIQ: PURINES": "Purín",
    "n_ex_pct": "N total reportado (%)",
    "n_ex_fraction": "Fracción másica de N",
    "masa_total_kg_eq": "Masa equivalente total",
    "tipo_muestra": "Tipo de material",
    "fuente_dato": "Fuente metodológica",
    "Factor hardcodeado auditado": "Parámetro o factor auditado",
}

HEADER_REPLACEMENTS = {
    "escenario": "Escenario",
    "etapa": "Etapa",
    "nombre_etapa": "Nombre de etapa",
    "tipo_muestra": "Tipo de material",
    "modelo_calculo": "Modelo de cálculo",
    "sistema_manejo_ipcc": "Sistema de manejo asignado",
    "parametro": "Parámetro",
    "valor": "Valor",
    "unidad": "Unidad",
    "referencia_metodologica": "Referencia metodológica",
    "variable": "Variable",
    "descripcion": "Descripción",
    "uso": "Uso metodológico",
    "flujo": "Flujo o material",
}

CHEMICAL_REPLACEMENTS = [
    ("kg CO2-eq/año", "kg CO\u2082-eq/año"),
    ("kg PO4-eq/año", "kg PO\u2084-eq/año"),
    ("kg N2O-N/kg N", "kg N\u2082O-N/kg N"),
    ("CO2-eq", "CO\u2082-eq"),
    ("PO4-eq", "PO\u2084-eq"),
    ("N2O-N", "N\u2082O-N"),
    ("NH3-N", "NH\u2083-N"),
    ("NO3-N", "NO\u2083-N"),
    ("CH4", "CH\u2084"),
    ("N2O", "N\u2082O"),
    ("NH3", "NH\u2083"),
    ("NO3", "NO\u2083\u207b"),
    ("CO2", "CO\u2082"),
    ("PO4", "PO\u2084\u00b3\u207b"),
    ("m3", "m\u00b3"),
    ("m2", "m\u00b2"),
]


def validate_inputs() -> None:
    validated_reference = get_reference_docx_path(ROOT)
    if validated_reference != REFERENCE_DOCX:
        raise RuntimeError(
            "La ruta validada del documento maestro no coincide con la ruta configurada."
        )
    required = [*TABLES.values(), OPERATIVE_PARAMS, MASS_BALANCE]
    missing = [path.relative_to(ROOT).as_posix() for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError("Faltan insumos requeridos:\n" + "\n".join(f"- {item}" for item in missing))

def read_csv(name: str) -> pd.DataFrame:
    try:
        return pd.read_csv(TABLES[name], encoding="utf-8")
    except UnicodeDecodeError:
        return pd.read_csv(TABLES[name], encoding="utf-8-sig")


def clean_text(value) -> str:
    if pd.isna(value):
        return ""
    text = str(value)
    for old, new in ACADEMIC_REPLACEMENTS.items():
        text = text.replace(old, new)
    for old, new in OLD_STAGE_TERMS.items():
        text = text.replace(old, new)
    text = text.replace("Eutrofizacion", "Eutrofización")
    text = text.replace("Nitrogeno", "Nitrógeno")
    text = text.replace("Solidos", "Sólidos")
    text = text.replace("Usado en ecuaciones de N despues de " + "cor" + "reccion", "Usado como fracción másica en ecuaciones de N")
    text = text.replace("n_ex_fraction " + "cor" + "regido para ecuaciones de nitrógeno", "n_ex_fraction usado como fracción másica en ecuaciones de nitrógeno")
    for old, new in CHEMICAL_REPLACEMENTS:
        text = text.replace(old, new)
    text = repair_mojibake(text)
    return clean_academic_label(text)


def repair_mojibake(text: str) -> str:
    markers = ["\u00c3", "\u00c2", "\u00e2\u20ac", "\u00e2\u20ac\u2122", "\u00e2\u20ac\u0153", "\ufffd"]
    if any(marker in text for marker in markers):
        for encoding in ("cp1252", "latin1"):
            try:
                return text.encode(encoding).decode("utf-8")
            except UnicodeError:
                continue
    return text


def apply_official_stage_names(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    if {"escenario", "etapa"}.issubset(out.columns):
        for idx, row in out.iterrows():
            try:
                key = (str(row["escenario"]).strip().upper(), int(row["etapa"]))
            except (TypeError, ValueError):
                continue
            if key in OFFICIAL_STAGE_NAMES:
                code, short_name, full_name = OFFICIAL_STAGE_NAMES[key]
                if "codigo_etapa" in out.columns:
                    out.at[idx, "codigo_etapa"] = code
                if "nombre_corto_etapa" in out.columns:
                    out.at[idx, "nombre_corto_etapa"] = short_name
                if "nombre_etapa" in out.columns:
                    out.at[idx, "nombre_etapa"] = full_name
    return out


def strip_internal_columns(df: pd.DataFrame) -> pd.DataFrame:
    return df[[col for col in df.columns if col.lower().strip() not in INTERNAL_COLUMNS]].copy()


def official_stage_label(scenario: object, stage: object) -> str:
    try:
        key = (str(scenario).strip().upper(), int(float(stage)))
    except (TypeError, ValueError):
        return ""
    if key in OFFICIAL_STAGE_NAMES:
        code, short_name, _ = OFFICIAL_STAGE_NAMES[key]
        return f"{code}: {short_name}"
    return f"{key[0]}{key[1]}"


def add_calculation_framework(df: pd.DataFrame) -> pd.DataFrame:
    """Distingue documentalmente el origen previo del marco de cálculo de cada etapa."""
    out = df.copy()
    frameworks = {
        ("A", 1): "Manejo del estiércol",
        ("A", 2): "Manejo del estiércol",
        ("A", 3): "Manejo del estiércol",
        ("A", 4): "Suelos gestionados",
        ("B", 1): "Manejo del estiércol",
        ("B", 2): "Suelos gestionados",
    }
    out["Marco de cálculo de la etapa"] = [
        frameworks.get((str(row["escenario"]).strip().upper(), int(float(row["etapa"]))), "")
        for _, row in out.iterrows()
    ]
    if "sistema_manejo_ipcc" in out.columns:
        out = out.rename(columns={"sistema_manejo_ipcc": "Sistema de manejo u origen previo"})
    return out


def combine_stage_columns(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    lower_to_col = {str(col).lower().strip(): col for col in out.columns}
    scenario_col = lower_to_col.get("escenario")
    stage_col = lower_to_col.get("etapa")
    if scenario_col and stage_col:
        insert_at = list(out.columns).index(stage_col)
        labels = [official_stage_label(row[scenario_col], row[stage_col]) for _, row in out.iterrows()]
        for col in ["codigo_etapa", "código", "codigo", "nombre_corto_etapa", "nombre_etapa", "nombre de etapa", "etapa"]:
            existing = lower_to_col.get(col)
            if existing in out.columns:
                out = out.drop(columns=[existing])
        out.insert(min(insert_at, len(out.columns)), "Etapa del sistema", labels)
    return out


def fmt(value, decimals: int = 2) -> str:
    if value == "" or pd.isna(value):
        return ""
    try:
        text = f"{float(value):,.{decimals}f}"
        text = text.replace(",", "X").replace(".", ",").replace("X", " ")
        if "," in text:
            text = text.rstrip("0").rstrip(",")
        return text
    except (TypeError, ValueError):
        return clean_text(value)


def format_df(df: pd.DataFrame, decimals: int = 2) -> pd.DataFrame:
    out = combine_stage_columns(df)
    out = out.rename(columns={col: HEADER_REPLACEMENTS.get(str(col), clean_text(col)) for col in out.columns})
    for col in out.columns:
        if str(col).lower() == "etapa":
            out[col] = out[col].map(lambda value: f"Etapa {int(float(value))}" if not pd.isna(value) and str(value) != "" else "")
        elif pd.api.types.is_numeric_dtype(out[col]):
            out[col] = out[col].map(lambda value: fmt(value, decimals))
        else:
            out[col] = out[col].map(clean_text)
    return out


def reference_format():
    ref = Document(str(REFERENCE_DOCX))
    normal = ref.styles["Normal"].font
    font_name = normal.name or "Times New Roman"
    font_size = normal.size or Pt(12)
    section = ref.sections[0]
    margins = {
        "top": section.top_margin,
        "bottom": section.bottom_margin,
        "left": section.left_margin,
        "right": section.right_margin,
    }
    return font_name, font_size, margins


def set_document_style(doc: Document):
    return apply_master_format(doc, REFERENCE_DOCX)


def set_table_horizontal_borders(table) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "insideH"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "808080")
        borders.append(tag)
    for edge in ("left", "right", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    table._tbl.tblPr.append(borders)


def add_dataframe_table(doc: Document, caption: str, df: pd.DataFrame) -> None:
    if caption:
        add_master_caption(doc, clean_text(caption))
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_horizontal_borders(table)
    for idx, col in enumerate(df.columns):
        cell = table.rows[0].cells[idx]
        cell.text = clean_text(col)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = clean_text(value)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if row_index == 0:
            table_header = OxmlElement("w:tblHeader")
            table_header.set(qn("w:val"), "true")
            tr_pr.append(table_header)
    format_table_like_master(table, analyze_master_format(REFERENCE_DOCX))


def add_figure(doc: Document, image: Path, caption: str) -> None:
    if not image.exists():
        return
    add_master_caption(doc, clean_text(caption))
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image), width=Inches(6.2))


def add_paragraphs(doc: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        doc.add_paragraph(clean_text(text), style="Normal")


def add_provisional_identification(doc: Document) -> None:
    marker = doc.add_paragraph()
    marker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = marker.add_run(PROVISIONAL_LABEL)
    run.bold = True
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = PROVISIONAL_LABEL
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(PROVISIONAL_NOTE, style="Normal")


def add_latex_equation(doc: Document, equation: str, definitions: list[str] | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(equation)
    run.font.name = "Cambria Math"
    run.font.size = Pt(12)
    if definitions:
        for item in definitions:
            doc.add_paragraph(clean_text(item), style="Normal")


def stage_summary() -> pd.DataFrame:
    etapas = apply_official_stage_names(read_csv("etapas"))
    out = etapas[
        ["nombre_escenario", "escenario", "etapa", "tipo_muestra_o_flujo", "sistema_ipcc", "modelo_calculo"]
    ].copy()
    out["Etapa del sistema"] = out.apply(
        lambda row: official_stage_label(row["escenario"], row["etapa"]), axis=1
    )
    return out[
        ["nombre_escenario", "Etapa del sistema", "tipo_muestra_o_flujo", "sistema_ipcc", "modelo_calculo"]
    ].rename(
        columns={
            "nombre_escenario": "Escenario",
            "tipo_muestra_o_flujo": "Material o flujo",
            "sistema_ipcc": "Sistema de manejo asignado",
            "modelo_calculo": "Modelo de estimación",
        }
    )


def model_summary() -> pd.DataFrame:
    params = apply_official_stage_names(read_csv("parametros"))
    subset = params[params["parametro"].isin(["Masa equivalente total", "Nitrogeno total reportado", "Nitrogeno total como fraccion masica", "MCF", "EF3"])]
    pivot = subset.pivot_table(
        index=["escenario", "etapa", "nombre_etapa", "modelo_calculo"],
        columns="parametro",
        values="valor",
        aggfunc="first",
    ).reset_index()
    pivot = pivot.rename(
        columns={
            "escenario": "Escenario",
            "etapa": "Etapa",
            "nombre_etapa": "Nombre de etapa",
            "modelo_calculo": "Modelo",
            "Masa equivalente total": "Masa equivalente (kg eq/año)",
            "Nitrogeno total reportado": "N total reportado (%)",
            "Nitrogeno total como fraccion masica": "Fracción másica de N",
        }
    )
    return pivot


def characterization_summary() -> pd.DataFrame:
    t02 = read_csv("caracterizacion")
    rows = []
    labels = {"Fresh manure": "Estiércol fresco", "Precomposted manure": "Estiércol precompostado"}
    for sample, group in t02.groupby("tipo_muestra", sort=False):
        row = {"Tipo de muestra": labels.get(sample, sample)}
        for _, record in group.iterrows():
            variable = record["variable"]
            unit = record["unidad"]
            if variable == "Humedad":
                row["Humedad (%)"] = record["valor"]
            elif variable == "Materia seca":
                row["Materia seca (%)"] = record["valor"]
            elif variable == "Cenizas":
                row["Cenizas (% base seca)"] = record["valor"]
            elif variable == "Solidos volatiles":
                row["Sólidos volátiles (% base seca)"] = record["valor"]
            elif variable == "Nitrogeno total" and unit == "% N total":
                row["N total (%)"] = record["valor"]
        rows.append(row)
    return pd.DataFrame(rows)


def flow_summary() -> pd.DataFrame:
    flows = apply_official_stage_names(read_csv("flujos"))
    out = flows[~flows["flujo"].isin(["Factor restante fresco a precompostado"])][
        ["escenario", "etapa", "nombre_etapa", "flujo", "valor", "unidad"]
    ].rename(
        columns={
            "escenario": "Escenario",
            "etapa": "Etapa",
            "nombre_etapa": "Nombre de etapa",
            "flujo": "Flujo o material",
            "valor": "Cantidad anual estimada",
            "unidad": "Unidad",
        }
    )
    return out


def parameter_long_summary() -> pd.DataFrame:
    params = apply_official_stage_names(read_csv("parametros"))
    keep = [
        "Nitrogeno total reportado",
        "Nitrogeno total como fraccion masica",
        "Solidos volatiles",
        "Masa equivalente total",
        "MCF",
        "EF3",
    ]
    out = params[params["parametro"].isin(keep)][
        ["escenario", "etapa", "nombre_etapa", "sistema_manejo_ipcc", "parametro", "valor", "unidad"]
    ]
    out = add_calculation_framework(out).rename(
        columns={
            "escenario": "Escenario",
            "etapa": "Etapa",
            "nombre_etapa": "Nombre de etapa",
            "parametro": "Variable principal",
            "valor": "Valor",
            "unidad": "Unidad",
        }
    )
    return out


def characterization_factors() -> pd.DataFrame:
    factors = read_csv("factores")
    cols = [
        col
        for col in [
            "sistema_o_compuesto",
            "factor",
            "valor",
            "unidad",
            "referencia_metodologica",
        ]
        if col in factors.columns
    ]
    out = factors[cols].copy()
    out = out[out["sistema_o_compuesto"].isin(["CH4", "N2O", "CO2", "NH3", "NO3"])].drop_duplicates()
    return out.rename(
        columns={
            "sistema_o_compuesto": "Sustancia",
            "factor": "Factor",
            "valor": "Valor",
            "unidad": "Unidad",
            "referencia_metodologica": "Referencia metodológica",
        }
    )


def methodology_context() -> dict[str, float]:
    operational = pd.read_csv(OPERATIVE_PARAMS, encoding="utf-8-sig")
    values = {
        str(row["parametro"]).strip(): float(row["valor"])
        for _, row in operational.iterrows()
    }
    required_operational = {
        "peso_vivo_promedio",
        "poblacion_media",
        "permanencia_sala",
        "estiercol_recolectado_animal",
        "estiercol_recolectado_anual",
        "agua_lavado_diaria",
        "fraccion_generacion_diaria",
    }
    if not required_operational.issubset(values):
        raise RuntimeError("Faltan parámetros operativos requeridos para documentar la metodología.")

    mass = pd.read_csv(MASS_BALANCE, encoding="utf-8-sig")
    for column in [
        "estiercol_total_depositado_anual_kg",
        "estiercol_remanente_anual_kg",
        "fraccion_recolectada",
        "fraccion_remanente",
        "flujo_referencia_anual_kg",
    ]:
        mass[column] = pd.to_numeric(mass[column], errors="raise")
    references = mass.groupby("escenario")["flujo_referencia_anual_kg"].unique()
    if set(references.index) != {"A", "B"} or any(len(item) != 1 for item in references):
        raise RuntimeError("No existe un flujo anual de referencia único para A y B.")
    reference_a = float(references["A"][0])
    reference_b = float(references["B"][0])
    if abs(reference_a - reference_b) > 1e-6:
        raise RuntimeError("Los escenarios no comparten el mismo flujo anual de referencia.")
    first = mass.iloc[0]
    daily_theoretical = values["peso_vivo_promedio"] * values["fraccion_generacion_diaria"]
    residence_deposit = daily_theoretical * values["permanencia_sala"] / 24
    remainder_animal = residence_deposit - values["estiercol_recolectado_animal"]
    return {
        **values,
        "agua_lavado_anual": values["agua_lavado_diaria"] * 365,
        "generacion_teorica_diaria": daily_theoretical,
        "deposito_durante_permanencia": residence_deposit,
        "remanente_animal": remainder_animal,
        "fraccion_recolectada": float(first["fraccion_recolectada"]),
        "fraccion_remanente": float(first["fraccion_remanente"]),
        "total_depositado_anual": float(first["estiercol_total_depositado_anual_kg"]),
        "remanente_anual": float(first["estiercol_remanente_anual_kg"]),
        "flujo_referencia": reference_a,
    }


def build_document() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    profile = set_document_style(doc)
    context = methodology_context()

    doc.add_paragraph(
        "Metodología desarrollada del Análisis de Ciclo de Vida", style="Title"
    )
    add_provisional_identification(doc)

    doc.add_heading("1. Enfoque metodológico general del ACV", level=2)
    add_paragraphs(doc, [
        "El estudio se desarrolló como un Análisis de Ciclo de Vida aplicado al manejo del estiércol bovino en una lechería especializada de Turrialba, Costa Rica. El procedimiento metodológico comprendió la definición de meta y alcance, la construcción del Inventario de Ciclo de Vida, la estimación de emisiones por etapa y la conversión de dichas emisiones a indicadores de impacto ambiental.",
        "El enfoque permitió comparar dos alternativas de manejo bajo una misma unidad funcional, manteniendo la trazabilidad entre las mediciones de laboratorio, los flujos del inventario y las ecuaciones de estimación de emisiones. Las categorías de impacto consideradas fueron calentamiento global y eutrofización.",
    ])

    doc.add_heading("2. Sitio de estudio", level=2)
    add_paragraphs(doc, [
        "El sitio de estudio correspondió a la lechería de la Sede del Atlántico de la Universidad de Costa Rica, ubicada en Turrialba. La evaluación se concentró en el manejo del estiércol fresco, la fracción sólida precompostada, las aguas verdes y los purines generados durante las actividades de ordeño y lavado de las áreas operativas.",
        "Sánchez-Romero y Brenes-Gamboa (2026) estudiaron esta misma lechería y aportaron la base documental de los parámetros operativos utilizados para cuantificar la generación de residuos durante el ordeño. La información de laboratorio del presente TFG permitió caracterizar los materiales empleados como entradas del inventario.",
    ])

    doc.add_heading("3. Meta, alcance y unidad funcional", level=2)
    add_paragraphs(doc, [
        "La meta del estudio fue comparar el desempeño ambiental de dos escenarios de manejo del estiércol bovino: el Escenario A, basado en lombricompostaje y aplicación de aguas verdes, y el Escenario B, basado en aplicación directa de purines en campo.",
        "Para el análisis se utilizó como unidad funcional 1 kg de estiércol fresco manejado. Esta unidad permitió expresar y comparar los impactos de los escenarios bajo una misma base funcional.",
        f"El flujo anual común asociado con la unidad funcional fue {fmt(context['flujo_referencia'], 6)} kg de estiércol fresco/año. Los flujos y emisiones anuales describen la magnitud operacional del sistema, mientras que los indicadores expresados por kilogramo de estiércol fresco corresponden a la normalización respecto a la unidad funcional. La unidad funcional no se definió como 1 kg/año.",
    ])

    doc.add_heading("4. Escenarios evaluados", level=2)
    add_paragraphs(doc, [
        f"Ambos escenarios se definieron como alternativas de manejo para el mismo flujo anual de referencia de {fmt(context['flujo_referencia'], 6)} kg de estiércol fresco/año.",
        f"En el Escenario A, {fmt(100 * context['fraccion_recolectada'], 2)} % del estiércol depositado se recolectó y entró primero a A1: Precomposteo; la masa resultante de esa transformación continuó posteriormente hacia A2: Lombricompostaje. El {fmt(100 * context['fraccion_remanente'], 2)} % restante permaneció adherido al piso, se incorporó a las aguas verdes y fue representado en A3: Almacenamiento de aguas verdes y A4: Aplicación de aguas verdes en campos de pastoreo.",
        f"En el Escenario B no se efectuó esa separación: el 100 % del estiércol teóricamente depositado, equivalente a {fmt(context['flujo_referencia'], 6)} kg/año, ingresó a B1: Almacenamiento de purines. En B2: Aplicación de purines en campo de pastoreo se integraron el agua de lavado y el total de estiércol depositado.",
    ])

    doc.add_heading("5. Fronteras del sistema", level=2)
    add_paragraphs(doc, [
        "La frontera del sistema incluyó las operaciones directamente asociadas con el manejo del estiércol, aguas verdes y purines dentro de la unidad productiva. Para el Escenario A se consideraron las etapas de precomposteo, lombricompostaje, almacenamiento de aguas verdes y aplicación de aguas verdes. Para el Escenario B se consideraron el almacenamiento de purines y su aplicación directa en campo.",
        "Las masas y volúmenes anuales estimados se emplearon para describir la magnitud operacional de las etapas evaluadas. Los flujos considerados para la construcción del inventario se presentan en la Tabla 2, y su distribución operacional se muestra en la Figura M1.",
    ])
    doc.add_heading("6. Etapas del Escenario A", level=2)
    add_paragraphs(doc, [
        "A1: Precomposteo corresponde al manejo inicial de la fracción sólida fresca antes del lombricompostaje. A2: Lombricompostaje corresponde a la transformación de la fracción sólida precompostada y representa el manejo posterior de dicha fracción.",
        "A3: Almacenamiento de aguas verdes corresponde al almacenamiento del efluente generado por el lavado y el arrastre de residuos. Para estimar las emisiones de manejo en esta etapa se utilizó como masa de actividad únicamente el estiércol remanente sometido al sistema, junto con la caracterización química del estiércol fresco y los factores del sistema IPCC seleccionado; el agua de lavado presente físicamente no se sumó como masa de estiércol en esas ecuaciones.",
        "A4: Aplicación de aguas verdes en campos de pastoreo corresponde a la etapa subsecuente de aplicación del efluente al suelo. Su flujo aplicado integró el agua de lavado y el estiércol remanente incorporado a las aguas verdes, y se empleó la caracterización química específica de las aguas verdes en las ecuaciones asociadas con suelos gestionados.",
    ])

    doc.add_heading("7. Etapas del Escenario B", level=2)
    add_paragraphs(doc, [
        "B1: Almacenamiento de purines corresponde físicamente al almacenamiento previo a la aplicación del efluente. Para estimar las emisiones de manejo se utilizó como masa de actividad la totalidad del estiércol teóricamente depositado, junto con la caracterización química del estiércol fresco y los factores del sistema IPCC seleccionado. Aunque el almacenamiento contiene agua de lavado, ese volumen no se incorporó como masa de estiércol en las ecuaciones de manejo.",
        "B2: Aplicación de purines en campo de pastoreo corresponde a la etapa subsecuente de aplicación del efluente al suelo. El flujo aplicado integró el agua de lavado y la totalidad del estiércol incorporado al purín, y se utilizó la caracterización química específica del purín en las ecuaciones asociadas con suelos gestionados.",
    ])

    doc.add_heading("8. Datos de entrada usados para el ICV", level=2)
    add_paragraphs(doc, [
        "Los datos de entrada del ICV estuvieron compuestos por información primaria de laboratorio, registros operativos de la lechería y factores de emisión seleccionados de fuentes bibliográficas. Se consideraron cuatro tipos principales de materiales: estiércol fresco, estiércol precompostado, aguas verdes y purines.",
        "El estiércol fresco constituyó el flujo de referencia del sistema y fue caracterizado mediante humedad, materia seca, cenizas, sólidos volátiles y nitrógeno total. El estiércol precompostado permitió representar la transformación de la fracción sólida durante el proceso de lombricompostaje.",
        "Las aguas verdes correspondieron al efluente generado durante el lavado de las áreas de manejo animal y fueron consideradas en el Escenario A para las etapas de almacenamiento y aplicación. Los purines representaron la mezcla líquida de estiércol, orina, agua de lavado y otros residuos arrastrados, considerada en el Escenario B para almacenamiento y aplicación directa en campo.",
        "Para cada flujo se organizaron las variables necesarias para estimar emisiones de CH4, N2O, NH3 y NO3, así como los impactos asociados a calentamiento global y eutrofización. Los resultados de la caracterización fisicoquímica de los materiales analizados se resumen en la Tabla 1.",
        f"Los parámetros operativos publicados para la misma lechería fueron: peso vivo promedio de {fmt(context['peso_vivo_promedio'], 1)} kg/animal, población media de {fmt(context['poblacion_media'], 0)} vacas en ordeño, permanencia aproximada de {fmt(context['permanencia_sala'], 1)} h/día, recolección de {fmt(context['estiercol_recolectado_animal'], 1)} kg de estiércol fresco/animal, total recolectado de {fmt(context['estiercol_recolectado_anual'], 1)} kg/año y consumo de agua de lavado de {fmt(context['agua_lavado_diaria'], 1)} L/día (Sánchez-Romero y Brenes-Gamboa, 2026). El consumo anual derivado fue {fmt(context['agua_lavado_anual'], 1)} L/año.",
        f"El valor de {fmt(context['estiercol_recolectado_anual'], 1)} kg/año representa el estiércol efectivamente recolectado durante las actividades de ordeño estudiadas; no corresponde a la excreción fisiológica total diaria de los animales.",
    ])
    add_dataframe_table(doc, "Tabla 1. Caracterización fisicoquímica de los materiales analizados.", format_df(characterization_summary(), decimals=3))
    add_dataframe_table(doc, "Tabla 2. Flujos considerados para el inventario de ciclo de vida.", format_df(flow_summary(), decimals=3))
    add_figure(doc, FIG_DIR / "fig_04_flujos_masa_equivalente_total.png", "Figura M1. Masa equivalente total por etapa y escenario.")

    doc.add_heading("9. Muestreo y análisis de laboratorio", level=2)
    add_paragraphs(doc, [
        "De acuerdo con el diseño previsto en el documento maestro, el muestreo se orientó a caracterizar los materiales que intervienen en las etapas del sistema evaluado mediante jornadas distribuidas durante el periodo de estudio. Para la fracción sólida se consideraron muestras de estiércol fresco y material precompostado, mientras que para la fracción líquida se consideraron aguas verdes y purines.",
        "Esta versión metodológica corresponde a la integración provisional de M1 y M2. La jornada M3 permanece pendiente y su incorporación actualizará la caracterización final.",
        "Las muestras sólidas permitieron determinar contenido de agua, materia seca, cenizas, sólidos volátiles y nitrógeno total; las muestras líquidas se utilizaron principalmente para representar el contenido de nitrógeno total y los flujos asociados al almacenamiento y aplicación en campo. Los resultados resumidos de caracterización se presentan en la Tabla 1 y los flujos utilizados para el inventario en la Tabla 2.",
        "La materia seca y los sólidos volátiles se emplearon para representar la fracción orgánica disponible en las estimaciones de CH4. El nitrógeno total se utilizó en la estimación de N2O, NH3 y NO3, según la alternativa de manejo correspondiente. Los parámetros principales utilizados para la estimación de emisiones se muestran en la Tabla 3.",
    ])

    doc.add_heading("10. Organización y procesamiento de datos", level=2)
    add_paragraphs(doc, [
        "Los datos obtenidos fueron organizados de acuerdo con el escenario de manejo, la etapa del sistema, el tipo de material y la variable medida. Esta organización permitió mantener la trazabilidad entre las mediciones de laboratorio, los flujos del inventario y las ecuaciones de estimación de emisiones.",
        "Posteriormente, las variables de laboratorio fueron transformadas en fracciones o masas equivalentes según la unidad requerida por cada ecuación. Los parámetros principales empleados en este proceso se presentan en la Tabla 3.",
        "La información completa de los parámetros utilizados en el modelo se presenta en el Apéndice interno A, Parámetros completos del modelo ACV.",
    ])
    add_dataframe_table(doc, "Tabla 3. Parámetros principales del modelo de estimación de emisiones.", format_df(parameter_long_summary(), decimals=4))

    doc.add_heading("11. Cálculo de humedad y materia seca", level=2)
    add_paragraphs(doc, ["La materia seca se calculó como la proporción entre la masa posterior al secado y la masa fresca inicial. La humedad se estimó como la fracción de agua removida durante el secado."])
    add_latex_equation(doc, r"MS(\%) = \frac{m_{\mathrm{seca}}}{m_{\mathrm{fresca}}} \times 100", ["Donde: MS = materia seca, %; m_seca = masa posterior al secado, g; m_fresca = masa fresca inicial, g."])
    add_latex_equation(doc, r"Humedad(\%) = \frac{m_{\mathrm{fresca}} - m_{\mathrm{seca}}}{m_{\mathrm{fresca}}} \times 100")

    doc.add_heading("12. Cálculo de cenizas y sólidos volátiles", level=2)
    add_paragraphs(doc, ["Las cenizas se determinaron como la fracción mineral remanente después de la calcinación. Los sólidos volátiles se estimaron como la fracción de materia seca que se volatiliza durante la calcinación."])
    add_latex_equation(doc, r"Cenizas(\%) = \frac{m_{\mathrm{cenizas}}}{m_{\mathrm{seca}}} \times 100", ["Donde: m_cenizas = masa remanente después de calcinación, g."])
    add_latex_equation(doc, r"SV(\%) = \frac{m_{\mathrm{seca}} - m_{\mathrm{cenizas}}}{m_{\mathrm{seca}}} \times 100", ["Donde: SV = sólidos volátiles, %."])

    doc.add_heading("13. Procesamiento de nitrógeno total", level=2)
    add_paragraphs(doc, [
        "El nitrógeno total reportado en porcentaje fue convertido a fracción másica antes de emplearse en las ecuaciones de estimación de emisiones. Esta conversión permitió expresar el contenido de nitrógeno en kg N/kg de muestra.",
        "La masa de nitrógeno asociada a cada flujo se estimó multiplicando la masa del flujo por la fracción másica de nitrógeno correspondiente.",
        "La definición académica de las variables empleadas en esta sección y en las demás ecuaciones se reúne en el Apéndice interno C, Diccionario de variables metodológicas.",
    ])
    add_latex_equation(doc, r"n_{\mathrm{ex,fraction}} = \frac{n_{\mathrm{ex,pct}}}{100}", ["Donde: n_ex,pct = nitrógeno total reportado en porcentaje, %; n_ex,fraction = fracción másica de nitrógeno, kg N/kg muestra."])
    add_latex_equation(doc, r"N_{\mathrm{total}} = m_{\mathrm{flujo}} \times n_{\mathrm{ex,fraction}}", ["Donde: N_total = masa de nitrógeno en el flujo, kg N; m_flujo = masa del flujo correspondiente, kg."])

    doc.add_heading("14. Trazabilidad de masa seca mediante conservación de cenizas", level=2)
    add_paragraphs(doc, [
        "Para mantener la trazabilidad de la fracción sólida entre el estiércol fresco y el material precompostado, se aplicó el supuesto de conservación de cenizas. Bajo este supuesto, la fracción mineral del material se considera una referencia conservativa que permite estimar la masa seca remanente después del proceso de precompostaje.",
        "De esta forma, la variación en el contenido de cenizas en base seca se utilizó para relacionar la masa seca inicial del estiércol fresco con la masa seca equivalente del material precompostado.",
    ])
    add_latex_equation(doc, r"m_{\mathrm{cenizas,fresco}} = m_{\mathrm{MS,fresco}} \times f_{\mathrm{cenizas,fresco}}")
    add_latex_equation(doc, r"m_{\mathrm{MS,precompostado}} = \frac{m_{\mathrm{cenizas,fresco}}}{f_{\mathrm{cenizas,precompostado}}}")
    add_latex_equation(doc, r"F_{\mathrm{MS,remanente}} = \frac{m_{\mathrm{MS,precompostado}}}{m_{\mathrm{MS,fresco}}} = \frac{f_{\mathrm{cenizas,fresco}}}{f_{\mathrm{cenizas,precompostado}}}", ["Donde: m_MS,fresco = masa seca inicial del estiércol fresco; m_MS,precompostado = masa seca equivalente del material precompostado; f_cenizas,fresco = fracción de cenizas del estiércol fresco en base seca; f_cenizas,precompostado = fracción de cenizas del material precompostado en base seca; F_MS,remanente = factor de masa seca remanente."])

    doc.add_heading("15. Construcción de flujos del inventario", level=2)
    add_paragraphs(doc, [
        "Los flujos del inventario se construyeron a partir de los parámetros operativos publicados para la lechería, el supuesto conservador adoptado para estimar el estiércol depositado durante la permanencia en sala y las masas o volúmenes asociados a cada etapa.",
        "Sánchez-Romero y Brenes-Gamboa (2026) citan un intervalo bibliográfico según el cual una vaca lechera puede producir diariamente estiércol equivalente a 7–10 % de su peso vivo. El presente TFG adoptó deliberadamente el límite inferior de 7 % como supuesto conservador, con el fin de evitar sobreestimar una fracción que no fue medida directamente. Por tanto, el 7 % no se atribuyó a una medición realizada por esos autores.",
        f"La generación teórica diaria se estimó como {fmt(context['peso_vivo_promedio'], 1)} × 0,07 = {fmt(context['generacion_teorica_diaria'], 3)} kg/animal/día. Para una permanencia de {fmt(context['permanencia_sala'], 1)} h/día, el depósito teórico fue {fmt(context['generacion_teorica_diaria'], 3)} × ({fmt(context['permanencia_sala'], 1)}/24) = {fmt(context['deposito_durante_permanencia'], 6)} kg/animal. Al descontar los {fmt(context['estiercol_recolectado_animal'], 1)} kg/animal recolectados, el remanente estimado fue {fmt(context['remanente_animal'], 6)} kg/animal.",
        f"La fracción remanente se calculó como {fmt(context['remanente_animal'], 6)}/{fmt(context['deposito_durante_permanencia'], 6)} = {fmt(context['fraccion_remanente'], 12)}, equivalente a {fmt(100 * context['fraccion_remanente'], 2)} %. La fracción recolectada fue {fmt(100 * context['fraccion_recolectada'], 2)} %. El {fmt(100 * context['fraccion_remanente'], 2)} % constituye una estimación derivada mediante balance y no una medición directa.",
        f"A partir del total recolectado publicado y de la fracción recolectada se obtuvo un total teóricamente depositado de {fmt(context['total_depositado_anual'], 6)} kg/año y un remanente de {fmt(context['remanente_anual'], 6)} kg/año. En el Escenario A, el total se dividió entre la ruta sólida y las aguas verdes; en el Escenario B, el total ingresó al sistema de purines.",
        "La equivalencia operativa de 1 L de agua igual a 1 kg equivalente se utilizó para integrar flujos líquidos y sólidos en una base común de inventario. Los flujos anuales estimados se reportan como magnitud operacional y no sustituyen la unidad funcional del ACV.",
        "La masa equivalente total de las etapas de aplicación no corresponde únicamente al volumen del componente líquido expresado como kg equivalente. En estas etapas se integran también las fracciones de estiércol asociadas al flujo aplicado. Por ello, aun cuando se empleó la equivalencia 1 L de agua = 1 kg para el componente líquido, la masa equivalente total representa la suma del componente líquido y la fracción de estiércol correspondiente a la etapa. En A4, esta fracción corresponde al estiércol remanente derivado del balance entre la masa teóricamente depositada en sala y la masa recolectada; en B2, corresponde al total anual de estiércol teóricamente depositado e integrado al purín aplicado en campo.",
    ])

    doc.add_heading("16. Normalización respecto a la unidad funcional", level=2)
    add_paragraphs(doc, [
        "La normalización consistió en expresar los impactos respecto a la unidad funcional del estudio: 1 kg de estiércol fresco manejado. Esta referencia permitió comparar los escenarios A y B bajo una base común.",
        f"Los resultados anualizados describen la magnitud operacional asociada al flujo común de {fmt(context['flujo_referencia'], 6)} kg/año, pero no constituyen la unidad funcional. El indicador normalizado se obtuvo al dividir cada impacto anual entre dicho flujo de referencia.",
    ])

    doc.add_heading("17. Aplicación de ecuaciones IPCC", level=2)
    add_paragraphs(doc, [
        "Las ecuaciones IPCC se aplicaron por etapa según el sistema de manejo asignado. Las vías consideradas incluyeron emisiones de CH4 por manejo de estiércol, N2O directo, pérdidas por volatilización y lixiviación, N2O indirecto y emisiones asociadas con suelos gestionados.",
        "Para representar A2: Lombricompostaje se seleccionó la categoría Composting – Passive Windrow de las Directrices del IPCC, por constituir la aproximación disponible más cercana a las condiciones operativas del sistema estudiado. El proceso se desarrolla con material sólido y no corresponde a un reactor cerrado; tampoco utiliza aireación forzada ni volteo mecánico intensivo diario. La actividad y el desplazamiento de las lombrices, sin embargo, contribuyen a la movilización y aireación no intensiva del sustrato.",
        "En A2 se adoptó una fracción de pérdida de N por lixiviación igual a cero como parámetro específico del sistema, sin alterar el valor genérico de la categoría Composting – Passive Windrow. Vargas Sarmiento (2023), en un ensayo efectuado en el mismo lombricario de la UCR Sede del Atlántico, documentó infraestructura bajo techo y camas construidas completamente con piso y paredes de cemento, sin drenajes; indicó además que la ausencia de drenajes dificultaba el escurrimiento del exceso de agua aplicado durante el riego experimental destinado a mantener la humedad objetivo.",
        "De manera separada, las observaciones directas efectuadas durante las giras de campo del presente TFG confirmaron que las camas se encontraban bajo techo, sin exposición directa a la precipitación, y que la adición regular de agua no constituía una práctica operacional habitual. Con base en la exclusión de lluvia directa, la impermeabilización de las camas, la ausencia de drenaje y la ausencia observada de riego operacional regular, no se modeló una pérdida de N hacia el ambiente por lixiviación durante A2. Este supuesto representa la ausencia de una salida significativa de N por esa vía y no niega que pueda ocurrir movimiento de agua dentro del sustrato.",
        "Para A3: Almacenamiento de aguas verdes y B1: Almacenamiento de purines, las emisiones asociadas con el manejo del estiércol se estimaron a partir de la masa de estiércol sometida al sistema, la caracterización del estiércol fresco y los factores correspondientes al sistema IPCC seleccionado. Aunque durante la operación real estos sistemas contienen agua de lavado, dicho volumen no se incorporó como masa de estiércol en las ecuaciones de manejo.",
        "El componente líquido se incorporó en las etapas subsecuentes de aplicación en campo, A4: Aplicación de aguas verdes en campos de pastoreo y B2: Aplicación de purines en campo de pastoreo. En estas etapas se representó el flujo total aplicado, integrado por el agua de lavado y la fracción de estiércol correspondiente, y se utilizó la caracterización química específica de las aguas verdes o del purín, según el escenario, en las ecuaciones asociadas con suelos gestionados.",
        "A2: Lombricompostaje utilizó las ecuaciones IPCC de manejo del estiércol y los factores genéricos de la categoría Composting – Passive Windrow, excepto la fracción de lixiviación, cuyo valor efectivo de cero se incorporó como parámetro específico del sistema estudiado. Las emisiones de las etapas se estimaron con los factores de manejo y parámetros descritos en la Tabla 3.",
        "Los factores asociados con las ecuaciones de estimación de emisiones se organizaron de acuerdo con la metodología IPCC. En A2 se conservaron los factores de metano, N2O directo y volatilización de la categoría Composting – Passive Windrow; únicamente la fracción de lixiviación se sustituyó mediante el parámetro específico descrito anteriormente.",
        "Los factores de emisión y caracterización empleados en estas estimaciones se detallan en el Apéndice interno B, Factores de emisión y caracterización.",
    ])

    doc.add_heading("18. Estimación de emisiones", level=2)
    add_paragraphs(doc, [
        "Las emisiones se estimaron por escenario, etapa y sustancia. Las sustancias consideradas fueron CO2, CH4, N2O, NH3 y NO3. Las ecuaciones de nitrógeno utilizaron n_ex,fraction como fracción másica de entrada.",
        "Las ecuaciones del IPCC permitieron estimar las pérdidas de N por volatilización y lixiviación, así como la fracción de dichas pérdidas transformada en N₂O-N indirecto. Estas ecuaciones no se utilizaron originalmente para calcular eutrofización; por ello, el presente TFG incorporó una adaptación metodológica para representar el N remanente potencialmente contribuyente a esa categoría.",
        "Se definió Nᴳ como el nitrógeno remanente de la ruta de volatilización después de descontar la fracción transformada en N₂O-N indirecto, y Nᴸ como el nitrógeno remanente de la ruta de lixiviación después del descuento equivalente. El conjunto potencialmente eutrofizante, Nₑᵤₜ, se definió como la suma de Nᴳ y Nᴸ. Las ecuaciones siguientes presentan estas magnitudes mediante subíndices en notación LaTeX.",
        "Para representar la especiación de Nₑᵤₜ dentro del inventario, se adoptó como supuesto metodológico del presente TFG una distribución de 50 % como N asociado a NH₃ y 50 % como N asociado a NO₃⁻. Esta distribución tomó como antecedente el supuesto empleado por Komakech et al. (2016) para el N del estiércol que alcanza cuerpos de agua, sin atribuir a esos autores la aplicación directa de este reparto al N remanente de las rutas IPCC.",
        "En consecuencia, NH₃ y NO₃⁻ se derivaron del conjunto Nₑᵤₜ y no de una correspondencia directa entre Nᴳ y NH₃ o entre Nᴸ y NO₃⁻. Una etapa con Nᴸ igual a cero puede presentar NO₃⁻ equivalente porque una mitad de Nₑᵤₜ se asigna a N asociado a nitrato bajo el supuesto de especiación; este resultado no representa una predicción directa de lixiviación física ni de la especie química final en el ambiente.",
        "Los factores de caracterización empleados para convertir emisiones en indicadores de impacto se presentan en la Tabla 4.",
    ])
    add_latex_equation(doc, r"N_G = N_{volatilizado} \times (1 - EF_4)")
    add_latex_equation(doc, r"N_L = N_{lixiviado} \times (1 - EF_5)")
    add_latex_equation(doc, r"N_{eut} = N_G + N_L")
    add_latex_equation(doc, r"N_{NH_3} = 0,5 \times N_{eut}")
    add_latex_equation(doc, r"N_{NO_3^-} = 0,5 \times N_{eut}")
    add_latex_equation(doc, r"m_{NH_3} = N_{NH_3} \times \frac{17}{14}")
    add_latex_equation(doc, r"m_{NO_3^-} = N_{NO_3^-} \times 4,4268")

    doc.add_heading("19. Evaluación de impacto de ciclo de vida", level=2)
    add_paragraphs(doc, [
        "La evaluación de impacto convirtió las emisiones estimadas en resultados equivalentes mediante factores de caracterización. Las categorías consideradas fueron calentamiento global, expresado como kg CO2-eq/año, y eutrofización, expresada como kg PO4-eq/año. La unidad de eutrofización se expresa como equivalente de fosfato y se relaciona con el ion PO4.",
        "Para la categoría de calentamiento global se emplearon los potenciales reportados por el IMN (2021), expresados en kg CO2-eq/kg de sustancia emitida. Para la categoría de eutrofización se utilizaron factores expresados en kg PO4-eq/kg, con base en Ecobilan (1999, como se citó en Vallejo, 2004).",
        "Los impactos se calcularon primero por etapa y posteriormente se agregaron por escenario. La comparación entre escenarios se realizó con diferencias absolutas y porcentuales entre el Escenario A y el Escenario B.",
    ])
    add_dataframe_table(doc, "Tabla 4. Factores de caracterización para las categorías de impacto.", format_df(characterization_factors(), decimals=4))

    doc.add_heading("20. Supuestos metodológicos", level=2)
    add_paragraphs(doc, [
        "Los supuestos principales fueron la equivalencia 1 L de agua = 1 kg equivalente, la extrapolación anual de flujos operativos, la adopción conservadora de 7 % del peso vivo/día para estimar el depósito teórico durante la permanencia en sala, la conservación de cenizas para estimar cambios de masa seca durante el precomposteo, la asignación de sistemas de manejo por etapa y el uso de factores de caracterización para las categorías de impacto evaluadas.",
        "Para la eutrofización se incluyó además la adaptación que distribuye por partes iguales el conjunto Nₑᵤₜ entre N asociado a NH₃ y N asociado a NO₃⁻. Esta adaptación representa una convención de inventario del presente TFG y no una medición directa de la especiación ambiental.",
        "Estos supuestos permiten mantener coherencia entre la caracterización de materiales, los flujos del inventario y las vías de emisión empleadas en el ACV.",
    ])

    doc.add_heading("21. Limitaciones metodológicas", level=2)
    add_paragraphs(doc, [
        "La estimación depende de la representatividad de las muestras de laboratorio y de la duración del muestreo de agua y estiércol. Las conversiones de N₂O-N a N₂O, de N a NH₃ y de N a NO₃⁻ se documentaron como cálculos estequiométricos derivados de relaciones de masa.",
        "La presentación de flujos anuales y masas equivalentes describe la escala operacional del sistema, pero la comparación metodológica entre escenarios se mantiene referida a 1 kg de estiércol fresco.",
    ])

    doc.add_heading("22. Referencias metodológicas incorporadas", level=2)
    add_paragraphs(doc, [
        "Sánchez-Romero, C. A., & Brenes-Gamboa, S. (2026). Cuantificación y caracterización de residuos generados durante el ordeño de ganado Jersey. Agronomía Mesoamericana, 37, artículo 6135ky76. https://doi.org/10.15517/6135ky76",
        "Komakech, A. J., Zurbrügg, C., Miito, G. J., Wanyama, J., & Vinnerås, B. (2016). Environmental impact from vermicomposting of organic waste in Kampala, Uganda. Journal of Environmental Management, 181, 395–402. https://doi.org/10.1016/j.jenvman.2016.06.028",
    ])

    doc.add_heading("Apéndices internos de metodología", level=1)
    add_paragraphs(doc, ["Los apéndices internos reúnen material técnico de apoyo. La Tabla M1 presenta parámetros completos del modelo ACV, la Tabla M2 presenta factores de emisión y caracterización, y la Tabla M3 presenta el diccionario de variables metodológicas."])

    doc.add_heading("Apéndice interno A. Parámetros completos del modelo ACV", level=2)
    params = add_calculation_framework(strip_internal_columns(apply_official_stage_names(read_csv("parametros"))))
    add_dataframe_table(doc, "Tabla M1. Parámetros completos usados por escenario y etapa.", format_df(params, decimals=4))

    doc.add_heading("Apéndice interno B. Factores de emisión y caracterización", level=2)
    factors = strip_internal_columns(read_csv("factores"))
    add_dataframe_table(doc, "Tabla M2. Factores técnicos usados en el modelo.", format_df(factors, decimals=4))

    doc.add_heading("Apéndice interno C. Diccionario de variables metodológicas", level=2)
    dictionary = strip_internal_columns(read_csv("diccionario"))
    add_dataframe_table(doc, "Tabla M3. Diccionario de variables metodológicas.", format_df(dictionary))

    finalize_document_format(doc, profile)
    doc.save(OUT_DOCX)


def main() -> None:
    validate_inputs()
    master_hash_before = sha256_file(REFERENCE_DOCX)
    build_document()
    visible = "\n".join(paragraph.text for paragraph in Document(OUT_DOCX).paragraphs)
    if PROVISIONAL_LABEL not in visible or "M3" not in visible:
        raise RuntimeError("La metodología no quedó identificada como PROVISIONAL M1–M2 pendiente de M3.")
    assert_reference_docx_intact(REFERENCE_DOCX, master_hash_before)
    print(f"Documento generado: {OUT_DOCX.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
