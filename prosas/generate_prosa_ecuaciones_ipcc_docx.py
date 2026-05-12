import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


OUTPUT = Path(__file__).resolve().parent / "prosa_ecuaciones_ipcc.docx"


EQUATION_TEXTS = [
    "La ecuación 1 estima el metano generado durante el manejo del estiércol. "
    "Relaciona sólidos volátiles, potencial metanogénico, conversión y sistema aplicado.",
    "La ecuación 2 calcula el N2O directo del estiércol gestionado. "
    "Usa nitrógeno excretado, codigestión, EF3 y conversión a N2O.",
    "La ecuación 3 estima el nitrógeno volatilizado desde el estiércol. "
    "Considera excreción, codigestión, sistema aplicado y fracción gaseosa.",
    "La ecuación 4 calcula el nitrógeno perdido por lixiviación. "
    "Parte del nitrógeno disponible y aplica la fracción correspondiente.",
    "La ecuación 5 estima el N2O indirecto por volatilización. "
    "Usa nitrógeno volatilizado, EF4 y conversión final a N2O.",
    "La ecuación 6 estima el N2O indirecto por lixiviación. "
    "Usa nitrógeno lixiviado, EF5 y conversión final a N2O.",
    "La ecuación 7 calcula el nitrógeno volatilizado remanente. "
    "Descuenta la fracción transformada en N2O mediante EF4.",
    "La ecuación 8 calcula el nitrógeno lixiviado remanente. "
    "Descuenta la fracción transformada en N2O mediante EF5.",
    "La ecuación 12 estima el amoníaco del manejo del estiércol. "
    "Promedia nitrógeno remanente y convierte esa cantidad a NH3.",
    "La ecuación 13 estima el nitrato del manejo del estiércol. "
    "Promedia nitrógeno remanente y convierte esa cantidad a NO3.",
    "La ecuación 14 calcula N2O por entradas orgánicas al suelo. "
    "Usa nitrógeno orgánico aplicado, EF1 y conversión a N2O.",
    "La ecuación 15 suma las fuentes orgánicas de nitrógeno. "
    "Incluye estiércol, aguas residuales, compost y otras enmiendas.",
    "La ecuación 16 estima N2O indirecto por deposición atmosférica. "
    "Usa nitrógeno aplicado, fracción volatilizada, EF4 y conversión a N2O.",
    "La ecuación 17 calcula el nitrógeno atmosférico remanente. "
    "Descuenta la fracción transformada en N2O mediante EF4.",
    "La ecuación 18 estima N2O por lixiviación en suelos. "
    "Usa nitrógeno aplicado, fracción lixiviada, EF5 y conversión a N2O.",
    "La ecuación 19 calcula el nitrógeno lixiviado remanente en suelos. "
    "Descuenta la fracción transformada en N2O mediante EF5.",
    "La ecuación 20 estima el amoníaco desde suelos gestionados. "
    "Promedia nitrógeno remanente y convierte esa cantidad a NH3.",
    "La ecuación 21 estima el nitrato desde suelos gestionados. "
    "Promedia nitrógeno remanente y convierte esa cantidad a NO3.",
    "La ecuación 22 calcula el nitrógeno disponible para suelos. "
    "Parte del nitrógeno entrante y descuenta pérdidas totales.",
    "La ecuación 23 calcula la fracción total de nitrógeno perdido. "
    "Suma pérdidas por volatilización, lixiviación, N2 y N2O-N.",
    "La ecuación 24 calcula la fracción perdida como N2. "
    "Usa la relación N2:N2O y el factor EF3.",
]


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)


def add_text_with_subscripts(paragraph, text: str) -> None:
    pattern = re.compile(r"(CH4|N2O|NH3|NO3|N2|EF\d+)")
    index = 0

    for match in pattern.finditer(text):
        if match.start() > index:
            paragraph.add_run(text[index : match.start()])

        for character in match.group(0):
            run = paragraph.add_run(character)
            if character.isdigit():
                run.font.subscript = True

        index = match.end()

    if index < len(text):
        paragraph.add_run(text[index:])


def main() -> None:
    document = Document()
    configure_styles(document)

    intro = document.add_paragraph(
        "Este texto resume las ecuaciones IPCC implementadas en Python. "
        "Cada párrafo describe su propósito dentro del análisis."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    paragraphs = [" ".join(EQUATION_TEXTS[index : index + 3]) for index in range(0, len(EQUATION_TEXTS), 3)]

    for text in paragraphs:
        paragraph = document.add_paragraph()
        add_text_with_subscripts(paragraph, text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for paragraph in document.paragraphs:
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.line_spacing = 1.15

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
