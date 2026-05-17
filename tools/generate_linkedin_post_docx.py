from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


BASE_DIR = Path(__file__).resolve().parents[1]
OUTPUT_PATH = BASE_DIR / "prosas" / "met" / "post_linkedin_tfg_acv_estiercol.docx"


POST_LINES = [
    "Estoy desarrollando mi Trabajo Final de Graduacion en la Escuela de Ingenieria de Biosistemas de la Universidad de Costa Rica, enfocado en el analisis ambiental del manejo de estiercol bovino en una lecheria especializada en Turrialba, Costa Rica.",
    "",
    "El proyecto aplica la metodologia de Analisis de Ciclo de Vida (ACV) para estimar impactos ambientales asociados al manejo de residuos solidos y liquidos generados durante las labores diarias de la lecheria de la Sede del Atlantico de la UCR.",
    "",
    "El estudio compara dos escenarios:",
    "",
    "Escenario A: manejo mediante precomposteo, lombricompostaje, almacenamiento de aguas verdes y aplicacion en campo.",
    "",
    "Escenario B: almacenamiento de purines y aplicacion directa en campos de pastoreo.",
    "",
    "El objetivo es construir un inventario de ciclo de vida y evaluar dos categorias de impacto ambiental relevantes para el sector agropecuario:",
    "",
    "- Potencial de calentamiento global",
    "- Potencial de eutrofizacion",
    "",
    "Para esto se integran datos de campo, analisis de laboratorio, reportes historicos de generacion de estiercol y consumo de agua, junto con ecuaciones y factores de emision basados en la metodologia IPCC.",
    "",
    "Una parte clave del trabajo ha sido el desarrollo de un pipeline reproducible en Python para el procesamiento, analisis y visualizacion de datos. Este flujo permite organizar tablas de entrada, calcular emisiones por etapa y escenario, estimar impactos ambientales y generar salidas tabulares y graficas de forma trazable.",
    "",
    "Tambien estoy trabajando con Codex como apoyo para estructurar, documentar y mejorar el pipeline computacional del proyecto, manteniendo un enfoque reproducible y ordenado para el analisis de datos.",
    "",
    "Ademas, se creo un repositorio en GitHub con el fin de compartir el proyecto, facilitar la revision del codigo y dejar una base abierta para futuras mejoras o adaptaciones metodologicas.",
    "",
    "Este trabajo tambien ha contado con apoyo institucional: se concurso y se ganaron fondos de la Vicerrectoria de Investigacion de la UCR para llevar a cabo los analisis de laboratorio necesarios para caracterizar las muestras y alimentar el modelo.",
    "",
    "Algunos resultados preliminares del modelo muestran:",
    "",
    "- Escenario A: 24,683.74 kg CO2-eq y 363.97 kg PO4-eq",
    "- Escenario B: 11,178.21 kg CO2-eq y 442.54 kg PO4-eq",
    "",
    "Estos resultados son preliminares, ya que aun quedan muestreos pendientes por integrar al analisis. En la siguiente etapa se incorporaran estos datos para considerar mejor la variabilidad temporal del sistema y fortalecer la interpretacion de los resultados.",
    "",
    "Mas alla de comparar escenarios, este trabajo busca identificar las etapas criticas del manejo del estiercol y aportar informacion util para la toma de decisiones en sistemas lecheros, especialmente en contextos donde la sostenibilidad ambiental, la gestion de nutrientes y la reduccion de impactos son cada vez mas importantes.",
    "",
    "Este proyecto conecta investigacion aplicada, datos experimentales, sostenibilidad agropecuaria, herramientas computacionales y colaboracion institucional para apoyar una gestion mas informada de los residuos organicos en la produccion de leche.",
    "",
    "#AnalisisDeCicloDeVida #ACV #LCA #Sostenibilidad #GanaderiaLechera #GestionDeResiduos #CambioClimatico #Eutrofizacion #IPCC #Python #GitHub #Codex #UniversidadDeCostaRica #IngenieriaDeBiosistemas #InvestigacionUCR #AgriculturaSostenible",
]


IMAGE_SUGGESTIONS = [
    "graphics_results/ACV_impacto_total_calentamiento_global_por_escenario.png",
    "graphics_results/ACV_impacto_total_eutrofizacion_por_escenario.png",
    "graphics_results/ACV_impacto_calentamiento_global.png",
]


def set_default_font(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)


def add_paragraphs(document: Document, lines: list[str]) -> None:
    for line in lines:
        if not line:
            document.add_paragraph()
            continue

        paragraph = document.add_paragraph()
        if line.startswith("- "):
            paragraph.style = "List Bullet"
            paragraph.add_run(line[2:])
        else:
            paragraph.add_run(line)


def main() -> None:
    document = Document()
    set_default_font(document)

    title = document.add_heading("Post para LinkedIn - TFG ACV estiercol bovino", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph("Version preparada a partir de la propuesta de TFG y los resultados preliminares del pipeline.")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph()
    add_paragraphs(document, POST_LINES)

    document.add_page_break()
    document.add_heading("Imagenes sugeridas para acompanar el post", level=2)
    for suggestion in IMAGE_SUGGESTIONS:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(suggestion)

    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
